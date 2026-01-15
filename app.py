#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Sistema Gestao Documentos - ACESSO DIRETO - SEM LOGIN - OTIMIZADO"""

# Imports padrão
import sqlite3
import re
import os
import json
from pathlib import Path
from datetime import datetime, date, timedelta
from collections import defaultdict
from functools import lru_cache, wraps
from time import time
import hashlib
import traceback
import shutil
import secrets

# Imports de terceiros
from dateutil.relativedelta import relativedelta
from docxtpl import DocxTemplate
from flask import Flask, render_template, request, jsonify, send_file, g
import pdfplumber
from werkzeug.utils import secure_filename
import google.generativeai as genai
from dotenv import load_dotenv
import requests  # Para consulta de CNPJ na Receita Federal
import mammoth  # Para conversão de DOCX para HTML

# Módulo de categorização de CNAEs
from categorias_cnae import obter_categoria_por_cnae

# Sistema de aprendizagem de layouts
from document_learning import DocumentLayoutLearner

# Gerenciador centralizado de banco de dados (SQLAlchemy)
from db_manager import DatabaseManager, get_db_manager

# Serviços ORM para tabelas legadas (transição gradual)
from services_compat import (
    ClienteServiceLegado, BoletoServiceLegado, ReciboServiceLegado,
    DocumentoServiceLegado, TagServiceLegado,
    get_cliente_service, get_boleto_service, get_recibo_service,
    get_documento_service, get_tag_service,
    registrar_tabelas_legadas
)

# Validação de entrada
from validators import (
    ClienteSchema, BoletoSchema,
    ValidationError, validar_cnpj, validar_cpf,
    sanitizar_string, validate_json
)

# Autocomplete de CNPJs da Receita Federal
from autocomplete_api import AutocompleteAPI

# Blueprint de Prospecção (Módulo Econodata)
from blueprints.prospeccao import prospeccao_bp

# Blueprint de Gerenciamento de Arquivos
from blueprints.file_manager import file_manager_bp

# Blueprints modulares (Fase 3 - Modularização)
from blueprints.clientes import clientes_bp
from blueprints.boletos import boletos_bp
from blueprints.tags import tags_bp
from blueprints.documentos import documentos_bp

# Data Bridge - Ponte entre cnpj_filtrado e gestao_documentos
from data_bridge import criar_bridge

# Carregar variáveis de ambiente do arquivo .env
load_dotenv()

# Configurar Gemini
gemini_api_key = os.environ.get("GEMINI_API_KEY")
if gemini_api_key:
    genai.configure(api_key=gemini_api_key)

# Configuração Flask otimizada
app = Flask(__name__)

# Registrar blueprints
app.register_blueprint(prospeccao_bp)
app.register_blueprint(file_manager_bp)

# Registrar blueprints modulares (Fase 3)
app.register_blueprint(clientes_bp, url_prefix='/api')
app.register_blueprint(boletos_bp, url_prefix='/api')
app.register_blueprint(tags_bp, url_prefix='/api')
app.register_blueprint(documentos_bp, url_prefix='/api')
print("[OK] Blueprints modulares registrados (clientes, boletos, tags, documentos)")

# Gerar SECRET_KEY segura (usa variável de ambiente ou gera uma nova)
SECRET_KEY = os.environ.get('FLASK_SECRET_KEY') or os.environ.get('SECRET_KEY')
if not SECRET_KEY:
    # Gerar uma chave segura se não existir
    # Em produção, defina FLASK_SECRET_KEY no ambiente
    SECRET_KEY = secrets.token_hex(32)
    print("[AVISO] SECRET_KEY gerada automaticamente. Para produção, defina FLASK_SECRET_KEY no ambiente.")

app.config.update(
    SECRET_KEY=SECRET_KEY,
    JSON_AS_ASCII=False,
    MAX_CONTENT_LENGTH=100 * 1024 * 1024,  # 100MB max upload
    SEND_FILE_MAX_AGE_DEFAULT=0,  # Desabilitar cache em dev
    SESSION_COOKIE_SECURE=os.environ.get('FLASK_ENV') == 'production',  # HTTPS em produção
    SESSION_COOKIE_HTTPONLY=True,  # Proteger contra XSS
    SESSION_COOKIE_SAMESITE='Lax',  # Proteger contra CSRF
)

BASE_DIR = Path(__file__).parent
DB_PATH = BASE_DIR / 'gestao_documentos.db'
MODELOS_DIR = BASE_DIR / 'modelos'
OUTPUT_DIR = BASE_DIR / 'output'
CNAES_FILE = BASE_DIR / 'cnaes_permitidos.txt'
CONFIG_FILE = BASE_DIR / 'diretorios_config.json'
OUTPUT_DIR.mkdir(exist_ok=True)
UPLOAD_DIR = BASE_DIR / 'uploads_pdf'
UPLOAD_DIR.mkdir(exist_ok=True)
TRAINING_DIR = BASE_DIR / 'training_samples'
TRAINING_DIR.mkdir(exist_ok=True)

# Inicializar sistema de aprendizagem de layouts
layout_learner = DocumentLayoutLearner(DB_PATH, TRAINING_DIR)
layout_learner.criar_tabelas()  # Garantir que tabelas existem
print("[OK] Sistema de aprendizagem de layouts inicializado")

# Inicializar API de Autocomplete de CNPJs da Receita Federal
CNPJ_DB_PATH = BASE_DIR / 'cnpj_filtrado.db'
if CNPJ_DB_PATH.exists():
    api_cnpj = AutocompleteAPI(str(CNPJ_DB_PATH))
    print("[OK] API de autocomplete de CNPJs inicializada")
else:
    api_cnpj = None
    print("[AVISO] Banco cnpj_filtrado.db não encontrado - autocomplete de CNPJs desabilitado")

# Inicializar Data Bridge - Ponte entre cnpj_filtrado e gestao_documentos
if CNPJ_DB_PATH.exists() and CNAES_FILE.exists():
    data_bridge = criar_bridge(BASE_DIR)
    print(f"[OK] Data Bridge inicializada - {len(data_bridge.cnaes_permitidos)} CNAEs permitidos")
else:
    data_bridge = None
    print("[AVISO] Data Bridge desabilitada - cnpj_filtrado.db ou cnaes_permitidos.txt não encontrado")

# Cache de CNAEs permitidos
CNAES_PERMITIDOS = {}

# ============================================
#  OTIMIZAÇÃO: Cache de Municípios
# ============================================
# Pré-compilar regex para substituição de códigos de município (50x mais rápido)
MUNICIPIOS_PATTERN = re.compile('|'.join(re.escape(codigo) for codigo in MUNICIPIOS.keys()))

def converter_municipios_rapido(texto):
    """
    Converte códigos de município para nomes em uma única passagem.
    50x mais rápido que múltiplos .replace()
    """
    if not texto:
        return texto
    return MUNICIPIOS_PATTERN.sub(lambda m: MUNICIPIOS[m.group()], str(texto))

# Regex patterns pré-compilados para otimizar extração de dados de PDFs
COMPILED_PATTERNS = {
    'data_emissao': [
        re.compile(r'Data\s+de\s+Emiss[aã]o[:\s]*(\d{2}[\/\-]\d{2}[\/\-]\d{4})', re.IGNORECASE),
        re.compile(r'Emiss[aã]o[:\s]*(\d{2}[\/\-]\d{2}[\/\-]\d{4})', re.IGNORECASE),
        re.compile(r'Data\s+Emiss[aã]o[:\s]*(\d{2}[\/\-]\d{2}[\/\-]\d{4})', re.IGNORECASE),
        re.compile(r'Data[:\s]*(\d{2}[\/\-]\d{2}[\/\-]\d{4})', re.IGNORECASE),
        re.compile(r'Competência[:\s]*(\d{2}[\/\-]\d{2}[\/\-]\d{4})', re.IGNORECASE),
        re.compile(r'(\d{2}[\/\-]\d{2}[\/\-]\d{4})', re.IGNORECASE),
    ],
    'vencimento': [
        re.compile(r'Vencimento[:\s]+(\d{2}[\/\-]\d{2}[\/\-]\d{4})', re.IGNORECASE),
        re.compile(r'Data\s+Vencimento[:\s]+(\d{2}[\/\-]\d{2}[\/\-]\d{4})', re.IGNORECASE),
        re.compile(r'Dt\.?\s*Vencimento[:\s]+(\d{2}[\/\-]\d{2}[\/\-]\d{4})', re.IGNORECASE),
    ],
    'valor': [
        re.compile(r'Valor\s+Total[:\s]*R?\$?\s*([\d\.,]+)', re.IGNORECASE),
        re.compile(r'Total\s+da\s+Nota[:\s]*R?\$?\s*([\d\.,]+)', re.IGNORECASE),
        re.compile(r'Valor\s+L[ií]quido[:\s]*R?\$?\s*([\d\.,]+)', re.IGNORECASE),
        re.compile(r'Total\s+L[ií]quido[:\s]*R?\$?\s*([\d\.,]+)', re.IGNORECASE),
        re.compile(r'Valor\s+do\s+Documento[:\s]*R?\$?\s*([\d\.,]+)', re.IGNORECASE),
        re.compile(r'Total[:\s]*R?\$?\s*([\d\.,]+)', re.IGNORECASE),
        re.compile(r'Valor[:\s]*R?\$?\s*([\d\.,]+)', re.IGNORECASE),
        re.compile(r'R\$\s*([\d\.,]+)', re.IGNORECASE),
    ],
    'numero': [
        re.compile(r'N[úu]mero[:\s]+(\d+)', re.IGNORECASE),
        re.compile(r'Nota[:\s]+(\d+)', re.IGNORECASE),
        re.compile(r'NF[:\s-]+(\d+)', re.IGNORECASE),
        re.compile(r'Documento[:\s]+(\d+)', re.IGNORECASE),
    ],
    'cnpj': [
        re.compile(r'(?:Tomador|Cliente).*?(?:CNPJ|CPF)[:\s/]*(\d{2}\.\d{3}\.\d{3}\/\d{4}\-\d{2})', re.IGNORECASE | re.DOTALL),
        re.compile(r'(\d{2}\.\d{3}\.\d{3}\/\d{4}\-\d{2})', re.IGNORECASE),
    ],
    'nome': [
        re.compile(r'Nome/Razão Social[:\s]*([A-ZÀÁÂÃÄÅÇÈÉÊËÌÍÎÏÑÒÓÔÕÖÙÚÛÜ][A-ZÀÁÂÃÄÅÇÈÉÊËÌÍÎÏÑÒÓÔÕÖÙÚÛÜ\s\-\.]+?)(?:\n|CPF|CNPJ)', re.IGNORECASE | re.MULTILINE),
        re.compile(r'Tomador de Serviços[:\s]*(?:Nome/Razão Social[:\s]*)?([A-ZÀÁÂÃÄÅÇÈÉÊËÌÍÎÏÑÒÓÔÕÖÙÚÛÜ][A-ZÀÁÂÃÄÅÇÈÉÊËÌÍÎÏÑÒÓÔÕÖÙÚÛÜ\s\-\.]+?)(?:\n|CPF|CNPJ)', re.IGNORECASE | re.MULTILINE),
        re.compile(r'Tomador[:\s]+([A-ZÀÁÂÃÄÅÇÈÉÊËÌÍÎÏÑÒÓÔÕÖÙÚÛÜ][A-ZÀÁÂÃÄÅÇÈÉÊËÌÍÎÏÑÒÓÔÕÖÙÚÛÜ\s\-\.]+?)(?:\n|CNPJ|CPF)', re.IGNORECASE | re.MULTILINE),
        re.compile(r'Cliente[:\s]+([A-ZÀÁÂÃÄÅÇÈÉÊËÌÍÎÏÑÒÓÔÕÖÙÚÛÜ][A-ZÀÁÂÃÄÅÇÈÉÊËÌÍÎÏÑÒÓÔÕÖÙÚÛÜ\s\-\.]+?)(?:\n|CNPJ|CPF)', re.IGNORECASE | re.MULTILINE),
        re.compile(r'Razão Social[:\s]+([A-ZÀÁÂÃÄÅÇÈÉÊËÌÍÎÏÑÒÓÔÕÖÙÚÛÜ][A-ZÀÁÂÃÄÅÇÈÉÊËÌÍÎÏÑÒÓÔÕÖÙÚÛÜ\s\-\.]+?)(?:\n|CNPJ|CPF)', re.IGNORECASE | re.MULTILINE),
        re.compile(r'Nome[:\s]+([A-ZÀÁÂÃÄÅÇÈÉÊËÌÍÎÏÑÒÓÔÕÖÙÚÛÜ][A-ZÀÁÂÃÄÅÇÈÉÊËÌÍÎÏÑÒÓÔÕÖÙÚÛÜ\s\-\.]+?)(?:\n|CNPJ|CPF)', re.IGNORECASE | re.MULTILINE),
    ],
    'endereco': [
        re.compile(r'Endere[çc]o[:\s]*([A-Za-zÀ-ú\.]+[^\n]+?)\s*-?\s*Bairro', re.IGNORECASE | re.MULTILINE),
        re.compile(r'Endere[çc]o[:\s]*([A-Za-zÀ-ú\.]+[^\n]+?)(?:\n|Bairro|Cep|CEP)', re.IGNORECASE | re.MULTILINE),
        re.compile(r'(?:Rua|Avenida|Av\.?|Travessa|Rod\.?|Estrada)\s+([^\n]+)', re.IGNORECASE | re.MULTILINE),
        re.compile(r'Logradouro[:\s]+([^\n]+)', re.IGNORECASE | re.MULTILINE),
    ],
    'bairro': [
        re.compile(r'Bairro[:\s]*([A-ZÀÁÂÃÄÅÇÈÉÊËÌÍÎÏÑÒÓÔÕÖÙÚÛÜ][A-ZÀ-Úa-zà-ú\s\-\.]+?)(?:\s*-\s*)?(?:Cep|CEP|\n)', re.IGNORECASE | re.MULTILINE),
    ],
    'cep': [
        re.compile(r'Cep[:\s]*(\d{5}[-]?\d{3})', re.IGNORECASE),
        re.compile(r'CEP[:\s]*(\d{5}[-]?\d{3})', re.IGNORECASE),
        re.compile(r'(\d{8})', re.IGNORECASE),
    ],
    'tributos': {
        'iss': [
            re.compile(r'ISS[:\s]+R?\$?\s*([\d\.,]+)', re.IGNORECASE),
            re.compile(r'ISSQN[:\s]+R?\$?\s*([\d\.,]+)', re.IGNORECASE)
        ],
        'pis': [re.compile(r'PIS[:\s]+R?\$?\s*([\d\.,]+)', re.IGNORECASE)],
        'cofins': [re.compile(r'COFINS[:\s]+R?\$?\s*([\d\.,]+)', re.IGNORECASE)],
        'inss': [re.compile(r'INSS[:\s]+R?\$?\s*([\d\.,]+)', re.IGNORECASE)],
        'ir': [
            re.compile(r'IR[:\s]+R?\$?\s*([\d\.,]+)', re.IGNORECASE),
            re.compile(r'IRRF[:\s]+R?\$?\s*([\d\.,]+)', re.IGNORECASE)
        ],
        'csll': [re.compile(r'CSLL[:\s]+R?\$?\s*([\d\.,]+)', re.IGNORECASE)],
    }
}

# Sistema de cache em memória otimizado
class SimpleCache:
    """Cache simples em memória com TTL"""
    def __init__(self, ttl=300):
        self.cache = {}
        self.ttl = ttl

    def get(self, key):
        if key in self.cache:
            value, timestamp = self.cache[key]
            if time() - timestamp < self.ttl:
                return value
            else:
                del self.cache[key]
        return None

    def set(self, key, value):
        self.cache[key] = (value, time())

    def clear(self):
        self.cache.clear()

    def invalidate(self, pattern=None):
        if pattern is None:
            self.cache.clear()
        else:
            keys_to_delete = [k for k in self.cache.keys() if pattern in k]
            for k in keys_to_delete:
                del self.cache[k]

# Instâncias de cache com diferentes TTLs
cache_api = SimpleCache(ttl=60)  # Cache de APIs - 1 minuto
cache_queries = SimpleCache(ttl=300)  # Cache de queries - 5 minutos
cache_tags = SimpleCache(ttl=600)  # Cache de tags - 10 minutos

# Mapeamento de códigos de municípios para nomes
MUNICIPIOS = {
    '4445': 'Divinópolis',
    '5300': 'Belo Horizonte',
    '4123': 'Juiz de Fora',
    '5206': 'Uberlândia',
    '4503': 'Contagem',
}

# Regex patterns compilados (otimização)
REGEX_PATTERNS = {
    'data': re.compile(r'(\d{2}/\d{2}/\d{4})'),
    'valor_bruto': re.compile(r'Valor\s+(?:bruto|total|líquido)[^\d]*R?\$?\s*([\d\.]+,\d{2})', re.IGNORECASE),
    'valor_total': re.compile(r'TOTAL\s+R?\$?\s*([\d\.]+,\d{2})', re.IGNORECASE),
    'valor_rs': re.compile(r'R\$\s*([\d\.]+,\d{2})'),
    'cnpj': re.compile(r'CPF/CNPJ[:\s]+([\d\.\-/]+)', re.IGNORECASE),
    'razao_social': re.compile(r'Nome/Razão social:\s*([^\n]+)', re.IGNORECASE),
    'nome_fantasia': re.compile(r'Nome fantasia:\s*([^\n]+)', re.IGNORECASE),
}

def obter_nome_municipio(codigo):
    """Retorna o nome do município pelo código, ou o próprio código se não encontrar"""
    if not codigo:
        return ''
    return MUNICIPIOS.get(str(codigo), str(codigo))

TEMPLATES = {
    'laudo_padrao': MODELOS_DIR / "LaudoPython2.docx",
    'laudo_caixa': MODELOS_DIR / "LaudoPythonCAIXA.docx",
    'laudo_dedetizacao': MODELOS_DIR / "LaudoPythonDEDETIZAÇÃO.docx",
    'laudo_caixa_sem_assinatura': MODELOS_DIR / "Laudo CAIXA sem ASSINATURA.docx",
    'laudo_ddt_caixa_sem_assinatura': MODELOS_DIR / "Laudo DDT E CAIXA sem ASSINATURA.docx",
    'laudo_ddt_sem_assinatura': MODELOS_DIR / "Laudo DDT sem ASSINATURA.docx",
    'recibo': MODELOS_DIR / "ReciboTemplate.docx",
    'orcamento': MODELOS_DIR / "OrçamentoTemplate.docx"
}

def get_db():
    """
    Conexão otimizada com SQLite.

    DEPRECATED: Use get_db_manager() e serviços ORM ao invés de SQL raw.

    Exemplo de migração:
        # Antes (SQL raw):
        conn = get_db()
        clientes = conn.execute("SELECT * FROM clientes_web").fetchall()

        # Depois (ORM):
        svc = get_cliente_service()
        clientes, total = svc.listar()
    """
    db = getattr(g, '_database', None)
    if db is None:
        db = g._database = sqlite3.connect(
            str(DB_PATH),
            timeout=20.0,
            check_same_thread=False
        )
        db.row_factory = sqlite3.Row
        # Otimizações SQLite
        db.execute('PRAGMA journal_mode=WAL')  # Write-Ahead Logging
        db.execute('PRAGMA synchronous=NORMAL')  # Menos sync, mais rápido
        db.execute('PRAGMA cache_size=-64000')  # 64MB cache
        db.execute('PRAGMA temp_store=MEMORY')  # Temp tables em memória
        db.execute('PRAGMA mmap_size=268435456')  # 256MB memory-mapped I/O
    return db

@app.teardown_appcontext
def close_connection(exception):
    db = getattr(g, '_database', None)
    if db is not None:
        db.close()

def tabela_existe(nome):
    return get_db().execute(
        "SELECT name FROM sqlite_master WHERE type='table' AND name=?", (nome,)
    ).fetchone() is not None

def adicionar_coluna_segura(db, tabela: str, coluna: str, tipo: str, default: str = None):
    """Adiciona coluna à tabela se não existir, ignorando erro se já existir"""
    try:
        sql = f'ALTER TABLE {tabela} ADD COLUMN {coluna} {tipo}'
        if default is not None:
            sql += f' DEFAULT {default}'
        db.execute(sql)
    except sqlite3.OperationalError as e:
        if 'duplicate column name' not in str(e).lower():
            print(f"[AVISO] Erro ao adicionar coluna {coluna}: {e}")


def criar_tabelas():
    db = get_db()
    db.execute('''CREATE TABLE IF NOT EXISTS clientes_web (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        nome_fantasia TEXT NOT NULL,
        razao_social TEXT,
        cnpj TEXT,
        cnae TEXT,
        rua TEXT,
        numero TEXT,
        bairro TEXT,
        cidade TEXT,
        uf TEXT,
        endereco_completo TEXT,
        telefone TEXT,
        ultima_data_servico TEXT,
        data_cadastro DATETIME DEFAULT CURRENT_TIMESTAMP
    )''')
    db.execute('''CREATE TABLE IF NOT EXISTS documentos_gerados (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        tipo_doc TEXT,
        nome_arquivo TEXT,
        caminho_arquivo TEXT,
        valor_total REAL DEFAULT 0,
        cliente_nome TEXT,
        razao_social TEXT,
        cliente_cnpj TEXT,
        data_geracao DATETIME DEFAULT CURRENT_TIMESTAMP
    )''')

    # Adicionar colunas se não existirem (para bancos existentes)
    adicionar_coluna_segura(db, 'documentos_gerados', 'cliente_nome', 'TEXT')
    adicionar_coluna_segura(db, 'documentos_gerados', 'razao_social', 'TEXT')
    adicionar_coluna_segura(db, 'documentos_gerados', 'cliente_cnpj', 'TEXT')

    # Adicionar campo de garantia aos clientes
    adicionar_coluna_segura(db, 'clientes_web', 'data_garantia', 'TEXT')
    adicionar_coluna_segura(db, 'clientes_web', 'periodo_garantia_meses', 'INTEGER', '12')


    # Criar índices para melhorar performance de queries
    db.execute('CREATE INDEX IF NOT EXISTS idx_clientes_cnpj ON clientes_web(cnpj)')
    db.execute('CREATE INDEX IF NOT EXISTS idx_clientes_nome ON clientes_web(nome_fantasia)')
    db.execute('CREATE INDEX IF NOT EXISTS idx_clientes_cidade ON clientes_web(cidade)')
    db.execute('CREATE INDEX IF NOT EXISTS idx_clientes_nome_cidade ON clientes_web(nome_fantasia, cidade)')
    db.execute('CREATE INDEX IF NOT EXISTS idx_clientes_cnae ON clientes_web(cnae)')
    db.execute('CREATE INDEX IF NOT EXISTS idx_documentos_tipo ON documentos_gerados(tipo_doc)')
    db.execute('CREATE INDEX IF NOT EXISTS idx_documentos_data ON documentos_gerados(data_geracao)')
    # db.execute('CREATE INDEX IF NOT EXISTS idx_documentos_cliente ON documentos_gerados(cliente_id)')  # Coluna cliente_id não existe
    db.execute('CREATE INDEX IF NOT EXISTS idx_documentos_cliente_cnpj ON documentos_gerados(cliente_cnpj, data_geracao)')
    db.execute('CREATE INDEX IF NOT EXISTS idx_documentos_data_tipo ON documentos_gerados(data_geracao, tipo_doc)')

    # Tabelas para sistema de tags
    db.execute('''CREATE TABLE IF NOT EXISTS tags (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        nome TEXT NOT NULL UNIQUE,
        descricao TEXT,
        cor TEXT DEFAULT '#3B82F6',
        data_criacao DATETIME DEFAULT CURRENT_TIMESTAMP
    )''')

    db.execute('''CREATE TABLE IF NOT EXISTS documento_tags (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        arquivo_id INTEGER,
        arquivo_tipo TEXT NOT NULL,
        tag_id INTEGER NOT NULL,
        confianca REAL DEFAULT 1.0,
        manual BOOLEAN DEFAULT 1,
        data_atribuicao DATETIME DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (tag_id) REFERENCES tags(id) ON DELETE CASCADE
    )''')

    db.execute('''CREATE TABLE IF NOT EXISTS tag_training_data (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        arquivo_nome TEXT,
        arquivo_conteudo TEXT,
        tag_id INTEGER NOT NULL,
        features TEXT,
        data_treinamento DATETIME DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (tag_id) REFERENCES tags(id) ON DELETE CASCADE
    )''')

    # Índices para tags
    db.execute('CREATE INDEX IF NOT EXISTS idx_documento_tags_arquivo ON documento_tags(arquivo_id, arquivo_tipo)')
    db.execute('CREATE INDEX IF NOT EXISTS idx_documento_tags_tag ON documento_tags(tag_id)')
    db.execute('CREATE INDEX IF NOT EXISTS idx_training_data_tag ON tag_training_data(tag_id)')

    # Tabela para gerenciamento de boletos
    db.execute('''CREATE TABLE IF NOT EXISTS boletos (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        cliente_id INTEGER,
        cliente_nome TEXT NOT NULL,
        descricao TEXT,
        valor REAL NOT NULL,
        data_emissao DATE NOT NULL,
        data_vencimento DATE NOT NULL,
        numero_documento TEXT,
        codigo_barras TEXT,
        status TEXT DEFAULT 'pendente',
        data_pagamento DATE,
        valor_pago REAL,
        observacoes TEXT,
        arquivo_caminho TEXT,
        data_criacao DATETIME DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (cliente_id) REFERENCES clientes_web(id) ON DELETE SET NULL
    )''')

    # Índices para boletos
    db.execute('CREATE INDEX IF NOT EXISTS idx_boletos_cliente ON boletos(cliente_id)')
    db.execute('CREATE INDEX IF NOT EXISTS idx_boletos_vencimento ON boletos(data_vencimento)')
    db.execute('CREATE INDEX IF NOT EXISTS idx_boletos_status ON boletos(status)')
    db.execute('CREATE INDEX IF NOT EXISTS idx_boletos_emissao ON boletos(data_emissao)')
    db.execute('CREATE INDEX IF NOT EXISTS idx_boletos_cliente_vencimento ON boletos(cliente_id, data_vencimento)')

    # Tabela para gerenciamento de recibos
    db.execute('''CREATE TABLE IF NOT EXISTS recibos (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        cliente_id INTEGER,
        cliente_nome TEXT NOT NULL,
        numero_recibo TEXT,
        descricao TEXT,
        valor_total REAL NOT NULL,
        data_emissao DATE NOT NULL,
        forma_pagamento TEXT,
        observacoes TEXT,
        arquivo_caminho TEXT,
        data_criacao DATETIME DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (cliente_id) REFERENCES clientes_web(id) ON DELETE SET NULL
    )''')

    # Índices para recibos
    db.execute('CREATE INDEX IF NOT EXISTS idx_recibos_cliente ON recibos(cliente_id)')
    db.execute('CREATE INDEX IF NOT EXISTS idx_recibos_emissao ON recibos(data_emissao)')
    db.execute('CREATE INDEX IF NOT EXISTS idx_recibos_numero ON recibos(numero_recibo)')


    # Inserir tags padrão se não existirem
    tags_padrao = [
        ('Laudo', 'Laudos técnicos e de serviços', '#10B981'),
        ('Recibo', 'Recibos de pagamento', '#3B82F6'),
        ('Orçamento', 'Orçamentos e propostas', '#F59E0B'),
        ('Boleto', 'Boletos bancários', '#F97316'),
        ('Contrato', 'Contratos e termos', '#8B5CF6'),
        ('Relatório', 'Relatórios diversos', '#6366F1'),
        ('Outros', 'Documentos diversos', '#6B7280')
    ]

    for nome, descricao, cor in tags_padrao:
        try:
            db.execute('INSERT OR IGNORE INTO tags (nome, descricao, cor) VALUES (?, ?, ?)',
                      (nome, descricao, cor))
        except sqlite3.IntegrityError:
            pass  # Tag já existe, ignorar

    db.commit()

def limpar_arquivos_temporarios():
    """Remove arquivos .txt temporários/desnecessários automaticamente"""
    arquivos_para_deletar = [
        'CORRECOES_APLICADAS.txt',
        'VALIDACAO_CNAE.txt',
        'RESUMO_CNAE.txt',
        'ENDERECO_COMPLETO.txt',
        'RESUMO_ENDERECO.txt',
        'GERENCIADOR_ARQUIVOS.txt',
        'RESUMO_ARQUIVOS.txt',
        'CONFIGURACAO_DIRETORIOS.txt',
        'NOVAS_VARIAVEIS_IMPLEMENTADAS.txt',
        'NOTAS_FISCAIS_E_ANALISE.txt',
        'VALORES_RECIBOS_E_LANCAMENTOS.txt',
        'teste_api_key.py',
        'teste_gemini.py',
        'teste_final.py',
        'listar_modelos.py',
        'diagnostico_completo.py',
        'corrigir_api_key.py',
        'limpar_arquivos.py',
        'testar_formatacoes.py',
        'SOLUCOES_API_KEY.md'
    ]

    for arquivo in arquivos_para_deletar:
        caminho = BASE_DIR / arquivo
        if caminho.exists():
            try:
                caminho.unlink()
            except (OSError, PermissionError) as e:
                pass  # Arquivo pode estar em uso ou sem permissão

# ==================== FUNÇÕES DE GERENCIAMENTO DE TAGS ====================

def obter_todas_tags():
    """Retorna todas as tags cadastradas (com cache)"""
    cached = cache_tags.get("all_tags")
    if cached is not None:
        return cached

    db = get_db()
    cursor = db.execute('SELECT id, nome, descricao, cor FROM tags ORDER BY nome')
    tags = []
    for row in cursor.fetchall():
        tags.append({
            'id': row[0],
            'nome': row[1],
            'descricao': row[2],
            'cor': row[3]
        })

    cache_tags.set("all_tags", tags)
    return tags

def adicionar_tag_documento(arquivo_id, arquivo_tipo, tag_id, confianca=1.0, manual=True):
    """Adiciona uma tag a um documento"""
    db = get_db()
    try:
        db.execute('''INSERT INTO documento_tags
                     (arquivo_id, arquivo_tipo, tag_id, confianca, manual)
                     VALUES (?, ?, ?, ?, ?)''',
                  (arquivo_id, arquivo_tipo, tag_id, confianca, int(manual)))
        db.commit()
        # Invalidar cache
        cache_tags.invalidate(f"tags_{arquivo_tipo}_{arquivo_id}")
        return True
    except Exception as e:
        print(f"Erro ao adicionar tag: {e}")
        return False

def remover_tag_documento(arquivo_id, arquivo_tipo, tag_id):
    """Remove uma tag de um documento"""
    db = get_db()
    try:
        db.execute('''DELETE FROM documento_tags
                     WHERE arquivo_id = ? AND arquivo_tipo = ? AND tag_id = ?''',
                  (arquivo_id, arquivo_tipo, tag_id))
        db.commit()
        # Invalidar cache
        cache_tags.invalidate(f"tags_{arquivo_tipo}_{arquivo_id}")
        return True
    except Exception as e:
        print(f"Erro ao remover tag: {e}")
        return False

def obter_tags_documento(arquivo_id, arquivo_tipo):
    """Retorna todas as tags de um documento (com cache)"""
    cache_key = f"tags_{arquivo_tipo}_{arquivo_id}"
    cached = cache_tags.get(cache_key)
    if cached is not None:
        return cached

    db = get_db()
    cursor = db.execute('''
        SELECT t.id, t.nome, t.cor, dt.confianca, dt.manual
        FROM documento_tags dt
        JOIN tags t ON dt.tag_id = t.id
        WHERE dt.arquivo_id = ? AND dt.arquivo_tipo = ?
    ''', (arquivo_id, arquivo_tipo))

    tags = []
    for row in cursor.fetchall():
        tags.append({
            'id': row[0],
            'nome': row[1],
            'cor': row[2],
            'confianca': row[3],
            'manual': bool(row[4])
        })

    cache_tags.set(cache_key, tags)
    return tags

def extrair_features_documento(nome_arquivo, conteudo_texto=''):
    """Extrai features de um documento para classificação"""
    features = {
        'nome_lower': nome_arquivo.lower(),
        'extensao': os.path.splitext(nome_arquivo)[1].lower(),
        'tamanho_nome': len(nome_arquivo),
    }

    # Palavras-chave para cada tipo de documento
    keywords = {
        'laudo': ['laudo', 'técnico', 'vistoria', 'inspeção', 'diagnóstico'],
        'recibo': ['recibo', 'pagamento', 'recebi', 'valor'],
        'orcamento': ['orçamento', 'proposta', 'cotação', 'estimativa'],
        'contrato': ['contrato', 'termo', 'acordo', 'convênio'],
        'relatorio': ['relatório', 'report', 'análise', 'balanço']
    }

    # Verificar palavras-chave no nome do arquivo
    for tipo, palavras in keywords.items():
        features[f'tem_{tipo}_nome'] = any(palavra in features['nome_lower'] for palavra in palavras)

    # Se tiver conteúdo, verificar palavras-chave no texto
    if conteudo_texto:
        conteudo_lower = conteudo_texto.lower()
        for tipo, palavras in keywords.items():
            features[f'tem_{tipo}_conteudo'] = any(palavra in conteudo_lower for palavra in palavras)

    return features

def sugerir_tags_automaticas(arquivo_id, arquivo_tipo, nome_arquivo, conteudo_texto=''):
    """Sugere tags automaticamente baseado em análise do documento"""
    features = extrair_features_documento(nome_arquivo, conteudo_texto)
    tags_sugeridas = []

    # Mapeamento de features para tags
    mapeamento = {
        'Laudo': ['tem_laudo_nome', 'tem_laudo_conteudo'],
        'Recibo': ['tem_recibo_nome', 'tem_recibo_conteudo'],
        'Orçamento': ['tem_orcamento_nome', 'tem_orcamento_conteudo'],
        'Contrato': ['tem_contrato_nome', 'tem_contrato_conteudo'],
        'Relatório': ['tem_relatorio_nome', 'tem_relatorio_conteudo']
    }

    db = get_db()

    for tag_nome, feature_keys in mapeamento.items():
        # Verificar se alguma feature corresponde
        confianca = 0.0
        for key in feature_keys:
            if features.get(key, False):
                confianca += 0.5

        if confianca > 0:
            # Obter ID da tag
            cursor = db.execute('SELECT id FROM tags WHERE nome = ?', (tag_nome,))
            row = cursor.fetchone()
            if row:
                tags_sugeridas.append({
                    'tag_id': row[0],
                    'tag_nome': tag_nome,
                    'confianca': min(confianca, 1.0)
                })

    return tags_sugeridas

def salvar_dados_treinamento(nome_arquivo, conteudo_texto, tag_id):
    """Salva dados de treinamento para melhorar o modelo"""
    db = get_db()
    features = extrair_features_documento(nome_arquivo, conteudo_texto)
    features_json = json.dumps(features)

    try:
        db.execute('''INSERT INTO tag_training_data
                     (arquivo_nome, arquivo_conteudo, tag_id, features)
                     VALUES (?, ?, ?, ?)''',
                  (nome_arquivo, conteudo_texto[:1000], tag_id, features_json))
        db.commit()
        return True
    except Exception as e:
        print(f"Erro ao salvar dados de treinamento: {e}")
        return False

# ==================== FUNÇÕES DE GARANTIA ====================

def verificar_garantias_vencendo(dias_aviso=30):
    """Verifica clientes com garantia vencendo nos próximos X dias"""
    db = get_db()
    hoje = datetime.now()
    data_limite = hoje + timedelta(days=dias_aviso)

    # Buscar clientes com garantia
    clientes = db.execute('''
        SELECT id, nome_fantasia, razao_social, telefone, data_garantia,
               periodo_garantia_meses, ultima_data_servico
        FROM clientes_web
        WHERE data_garantia IS NOT NULL AND data_garantia != ''
        ORDER BY data_garantia ASC
    ''').fetchall()

    avisos = []
    for cliente in clientes:
        try:
            data_garantia = datetime.strptime(cliente['data_garantia'], '%Y-%m-%d')
            dias_restantes = (data_garantia - hoje).days

            if dias_restantes <= dias_aviso and dias_restantes >= 0:
                status = 'urgente' if dias_restantes <= 7 else 'atencao' if dias_restantes <= 15 else 'proximo'
                avisos.append({
                    'id': cliente['id'],
                    'nome': cliente['nome_fantasia'],
                    'razao_social': cliente['razao_social'],
                    'telefone': cliente['telefone'],
                    'data_garantia': cliente['data_garantia'],
                    'dias_restantes': dias_restantes,
                    'status': status,
                    'vencida': False
                })
            elif dias_restantes < 0:
                avisos.append({
                    'id': cliente['id'],
                    'nome': cliente['nome_fantasia'],
                    'razao_social': cliente['razao_social'],
                    'telefone': cliente['telefone'],
                    'data_garantia': cliente['data_garantia'],
                    'dias_restantes': dias_restantes,
                    'status': 'vencida',
                    'vencida': True
                })
        except (ValueError, TypeError):
            continue  # Data inválida ou ausente, pular cliente

    return avisos

@lru_cache(maxsize=1)
def carregar_cnaes_permitidos():
    """Carrega CNAEs do arquivo com cache LRU"""
    global CNAES_PERMITIDOS

    if CNAES_PERMITIDOS:
        return CNAES_PERMITIDOS

    try:
        if not CNAES_FILE.exists():
            print(f"AVISO: Arquivo {CNAES_FILE} não encontrado")
            return {}

        # Ler arquivo inteiro de uma vez (mais rápido)
        with open(CNAES_FILE, 'r', encoding='utf-8') as f:
            linhas = f.readlines()

        # Processar com list comprehension (mais rápido)
        for linha in linhas:
            linha = linha.strip()
            if not linha or linha.startswith('#'):
                continue

            if ' - ' in linha:
                cnae_formatado = linha.split(' - ', 1)[0].strip()  # maxsplit=1 é mais rápido
                cnae_numero = cnae_formatado.translate(str.maketrans('', '', '.-'))  # Mais rápido que replace
                CNAES_PERMITIDOS[cnae_numero] = linha

        print(f"[OK] CNAEs carregados: {len(CNAES_PERMITIDOS)}")
        return CNAES_PERMITIDOS

    except Exception as e:
        print(f"[ERRO] Erro ao carregar CNAEs: {e}")
        return {}

@lru_cache(maxsize=512)
def formatar_cnae(cnae_numero):
    """
    Formata CNAE com cache LRU (otimizado)
    Entrada: "8122200" ou "81.22-2-00"
    Saída: "81.22-2-00 - Imunização..." ou formatado
    """
    if not cnae_numero:
        return None

    # Remover formatação (otimizado com translate)
    cnae_limpo = str(cnae_numero).translate(str.maketrans('', '', '.- ')).strip()

    # Buscar nos CNAEs permitidos
    cnaes = carregar_cnaes_permitidos()

    if cnae_limpo in cnaes:
        return cnaes[cnae_limpo]

    # Se não encontrar, formatar apenas o número
    if len(cnae_limpo) == 7:
        return f"{cnae_limpo[0:2]}.{cnae_limpo[2:4]}-{cnae_limpo[4]}-{cnae_limpo[5:7]}"

    return cnae_numero

def formatar_cnpj(cnpj):
    """
    Formata CNPJ no padrão xx.xxx.xxx/xxxx-xx
    Entrada: "12345678000190" ou "12.345.678/0001-90" ou qualquer formato
    Saída: "12.345.678/0001-90"
    """
    if not cnpj:
        return ''

    # Remover todos os caracteres não numéricos
    cnpj_limpo = ''.join(filter(str.isdigit, str(cnpj)))

    # Verificar se tem 14 dígitos (CNPJ válido)
    if len(cnpj_limpo) != 14:
        return cnpj  # Retorna o original se não for CNPJ válido

    # Formatar: xx.xxx.xxx/xxxx-xx
    return f"{cnpj_limpo[0:2]}.{cnpj_limpo[2:5]}.{cnpj_limpo[5:8]}/{cnpj_limpo[8:12]}-{cnpj_limpo[12:14]}"

def gerar_nome_arquivo(nome_cliente, data_obj, tipo_doc=''):
    """
    Gera nome de arquivo no padrão: NOME FANTASIA - mes - ano
    Exemplo: "Empresa ABC - dez - 2025.docx"
    """
    # Mapear número do mês para nome em português (abreviado)
    meses = {
        1: 'jan', 2: 'fev', 3: 'mar', 4: 'abr', 5: 'mai', 6: 'jun',
        7: 'jul', 8: 'ago', 9: 'set', 10: 'out', 11: 'nov', 12: 'dez'
    }

    # Limpar nome do cliente (remover caracteres inválidos para nome de arquivo)
    nome_limpo = nome_cliente.replace('/', '-').replace('\\', '-').replace(':', '-')

    # Obter mês e ano
    mes_abrev = meses.get(data_obj.month, str(data_obj.month))
    ano = data_obj.year

    # Montar nome: NOME - mes - ano
    nome_arquivo = f"{nome_limpo} - {mes_abrev} - {ano}.docx"

    return nome_arquivo

@lru_cache(maxsize=512)
def cnae_esta_permitido(cnae_numero):
    """Verifica se CNAE está permitido (otimizado com cache)"""
    if not cnae_numero:
        return False

    cnae_limpo = str(cnae_numero).translate(str.maketrans('', '', '.- ')).strip()
    cnaes = carregar_cnaes_permitidos()
    return cnae_limpo in cnaes

def formatar_data(data_str):
    try:
        if data_str:
            obj = datetime.strptime(data_str, '%Y-%m-%d').date()
            return obj, obj.strftime('%d/%m/%Y')
    except (ValueError, TypeError):
        pass  # Formato de data inválido
    hoje = date.today()
    return hoje, hoje.strftime('%d/%m/%Y')

def formatar_data_extenso(data_obj):
    """Formata data como '07 de dezembro de 2025'"""
    meses = {
        1: 'janeiro', 2: 'fevereiro', 3: 'março', 4: 'abril',
        5: 'maio', 6: 'junho', 7: 'julho', 8: 'agosto',
        9: 'setembro', 10: 'outubro', 11: 'novembro', 12: 'dezembro'
    }
    dia = data_obj.day
    mes = meses[data_obj.month]
    ano = data_obj.year
    return f"{dia:02d} de {mes} de {ano}"

def numero_por_extenso(valor):
    """
    Converte valor monetário para extenso em português.
    Exemplo: 1250.50 -> "um mil, duzentos e cinquenta reais e cinquenta centavos"
    """
    if not valor:
        return "zero reais"

    try:
        # Limpar e converter o valor
        if isinstance(valor, str):
            valor_limpo = valor.replace('R$', '').replace(' ', '').replace(',', '.').strip()
            valor_float = float(valor_limpo)
        else:
            valor_float = float(valor)

        # Separar reais e centavos
        reais = int(valor_float)
        centavos = int(round((valor_float - reais) * 100))

        # Arrays para conversão
        unidades = ['', 'um', 'dois', 'três', 'quatro', 'cinco', 'seis', 'sete', 'oito', 'nove']
        dez_a_dezenove = ['dez', 'onze', 'doze', 'treze', 'quatorze', 'quinze', 'dezesseis', 'dezessete', 'dezoito', 'dezenove']
        dezenas = ['', '', 'vinte', 'trinta', 'quarenta', 'cinquenta', 'sessenta', 'setenta', 'oitenta', 'noventa']
        centenas = ['', 'cento', 'duzentos', 'trezentos', 'quatrocentos', 'quinhentos', 'seiscentos', 'setecentos', 'oitocentos', 'novecentos']

        def converter_ate_999(num):
            """Converte números de 0 a 999"""
            if num == 0:
                return ''
            elif num == 100:
                return 'cem'
            elif num < 10:
                return unidades[num]
            elif num < 20:
                return dez_a_dezenove[num - 10]
            elif num < 100:
                dezena = num // 10
                unidade = num % 10
                if unidade == 0:
                    return dezenas[dezena]
                return f"{dezenas[dezena]} e {unidades[unidade]}"
            else:
                centena = num // 100
                resto = num % 100
                if resto == 0:
                    return centenas[centena]
                elif resto < 10:
                    return f"{centenas[centena]} e {unidades[resto]}"
                else:
                    return f"{centenas[centena]} e {converter_ate_999(resto)}"

        def converter_numero(num):
            """Converte número completo para extenso"""
            if num == 0:
                return 'zero'

            # Separar em grupos de 3 dígitos (unidades, milhares, milhões)
            milhoes = num // 1000000
            milhares = (num % 1000000) // 1000
            unidades_grupo = num % 1000

            partes = []

            # Milhões
            if milhoes > 0:
                if milhoes == 1:
                    partes.append('um milhão')
                else:
                    partes.append(f"{converter_ate_999(milhoes)} milhões")

            # Milhares
            if milhares > 0:
                if milhares == 1:
                    partes.append('mil')
                else:
                    partes.append(f"{converter_ate_999(milhares)} mil")

            # Unidades
            if unidades_grupo > 0:
                partes.append(converter_ate_999(unidades_grupo))

            # Juntar partes
            if len(partes) == 0:
                return 'zero'
            elif len(partes) == 1:
                return partes[0]
            elif len(partes) == 2:
                return f"{partes[0]} e {partes[1]}"
            else:
                return ', '.join(partes[:-1]) + f" e {partes[-1]}"

        # Montar texto final
        texto_reais = converter_numero(reais)

        # Singular ou plural para reais
        if reais == 1:
            texto_reais += " real"
        else:
            texto_reais += " reais"

        # Adicionar centavos se houver
        if centavos > 0:
            texto_centavos = converter_numero(centavos)
            if centavos == 1:
                texto_centavos += " centavo"
            else:
                texto_centavos += " centavos"
            return f"{texto_reais} e {texto_centavos}"

        return texto_reais

    except (ValueError, TypeError, AttributeError):
        return "valor inválido"

def extrair_dados_pdf(caminho_arquivo):
    """
    Extrai informações de PDFs (otimizado para performance)
    Usa sistema de aprendizagem de layouts quando possível
    """
    dados = {
        'tipo_documento': 'PDF',
        'numero': None,
        'data_emissao': None,
        'valor_total': 0.0,
        'cliente_nome': None,
        'cliente_cnpj': None,
        'prestador_nome': None,
        'prestador_cnpj': None,
        'servicos': None,
        'layout_reconhecido': False  # Indica se usou layout aprendido
    }

    try:
        # ==================== TENTATIVA 1: USAR LAYOUT APRENDIDO ====================
        # Tentar identificar layout conhecido
        layout_info = layout_learner.identificar_layout(Path(caminho_arquivo))

        if layout_info and layout_info.get('regex'):
            print(f"  → Layout reconhecido: {layout_info['nome']}")
            dados['layout_reconhecido'] = True

            # Extrair texto do PDF
            with pdfplumber.open(caminho_arquivo) as pdf:
                # Extrair texto da primeira página
                texto = pdf.pages[0].extract_text() or ''

                # Usar os regex otimizados do layout aprendido
                regex_patterns = layout_info['regex']

                # Número da nota
                if 'numero_nota' in regex_patterns:
                    match = re.search(regex_patterns['numero_nota'], texto, re.IGNORECASE)
                    if match:
                        dados['numero'] = match.group(1)

                # Data de emissão
                if 'data_emissao' in regex_patterns:
                    match = re.search(regex_patterns['data_emissao'], texto, re.IGNORECASE)
                    if match:
                        data_str = match.group(1)
                        try:
                            data_obj = datetime.strptime(data_str, '%d/%m/%Y')
                            dados['data_emissao'] = data_obj.strftime('%Y-%m-%d')
                        except ValueError:
                            pass  # Formato de data não reconhecido

                # Valor total
                if 'valor_total' in regex_patterns:
                    match = re.search(regex_patterns['valor_total'], texto, re.IGNORECASE)
                    if match:
                        try:
                            valor_str = match.group(1).replace('.', '').replace(',', '.')
                            dados['valor_total'] = float(valor_str)
                        except (ValueError, AttributeError):
                            pass  # Valor não é número válido

                # CNPJ do cliente/tomador
                if 'cnpj' in regex_patterns:
                    match = re.search(regex_patterns['cnpj'], texto, re.IGNORECASE)
                    if match:
                        dados['cliente_cnpj'] = match.group(1).strip()

                # Nome do cliente/tomador
                if 'tomador_nome' in regex_patterns:
                    match = re.search(regex_patterns['tomador_nome'], texto, re.IGNORECASE)
                    if match:
                        dados['cliente_nome'] = match.group(1).strip()

                # Nome do prestador
                if 'prestador_nome' in regex_patterns:
                    match = re.search(regex_patterns['prestador_nome'], texto, re.IGNORECASE)
                    if match:
                        dados['prestador_nome'] = match.group(1).strip()

                # Discriminação/Serviços
                if 'discriminacao' in regex_patterns:
                    match = re.search(regex_patterns['discriminacao'], texto, re.IGNORECASE)
                    if match:
                        dados['servicos'] = match.group(1).strip()

                # Tipo de documento
                dados['tipo_documento'] = layout_info.get('tipo', 'PDF').upper()

            print(f"OK PDF processado (layout aprendido): {dados['numero']} - R$ {dados['valor_total']:.2f}")
            return dados

        # ==================== TENTATIVA 2: EXTRAÇÃO COM REGEX GENÉRICOS ====================
        print(f"  → Layout desconhecido - usando regex genéricos")

        with pdfplumber.open(caminho_arquivo) as pdf:
            # Limitar a 5 páginas para evitar PDFs muito grandes
            max_pages = min(len(pdf.pages), 5)

            # Extrair texto apenas das páginas necessárias
            textos = []
            for i in range(max_pages):
                texto_pagina = pdf.pages[i].extract_text()
                if texto_pagina:
                    textos.append(texto_pagina)

            texto = '\n'.join(textos)

            # 1. IDENTIFICAR TIPO
            texto_upper = texto.upper()
            if 'LAUDO' in texto_upper or 'TÉCNICO' in texto_upper:
                dados['tipo_documento'] = 'LAUDO'
            elif 'RECIBO' in texto_upper:
                dados['tipo_documento'] = 'RECIBO'
            elif 'ORÇAMENTO' in texto_upper or 'ORCAMENTO' in texto_upper:
                dados['tipo_documento'] = 'ORCAMENTO'

            # 2. EXTRAIR NÚMERO
            numero_match = re.search(r'N[°º]?\s*:?\s*(\d+)', texto)
            if numero_match:
                dados['numero'] = numero_match.group(1)

            # 3. EXTRAIR DATA (procurar múltiplas ocorrências e converter formato)
            datas_encontradas = []
            for match in REGEX_PATTERNS['data'].finditer(texto):
                data_str = match.group(1)  # formato dd/mm/yyyy
                try:
                    # Converter dd/mm/yyyy para yyyy-mm-dd
                    data_obj = datetime.strptime(data_str, '%d/%m/%Y')
                    # Validar se a data é razoável (não muito antiga, não futura)
                    hoje = datetime.now()
                    diferenca_dias = abs((data_obj - hoje).days)
                    # Aceitar datas de até 5 anos no passado ou 1 ano no futuro
                    if diferenca_dias <= (5 * 365):
                        datas_encontradas.append(data_obj)
                except ValueError:
                    continue  # Formato de data inválido

            # Usar a data mais antiga encontrada (geralmente é a data de emissão)
            if datas_encontradas:
                data_mais_antiga = min(datas_encontradas)
                dados['data_emissao'] = data_mais_antiga.strftime('%Y-%m-%d')

            # 4. EXTRAIR VALOR (otimizado)
            valores = []
            for pattern_key in ['valor_bruto', 'valor_total', 'valor_rs']:
                for match in REGEX_PATTERNS[pattern_key].finditer(texto):
                    try:
                        valor_str = match.group(1).replace('.', '').replace(',', '.')
                        valores.append(float(valor_str))
                    except (ValueError, AttributeError):
                        continue  # Valor não convertível
            if valores:
                dados['valor_total'] = max(valores)

            # 5. EXTRAIR DADOS DO CLIENTE
            match = REGEX_PATTERNS['cnpj_tomador'].search(texto)
            if match:
                dados['cliente_cnpj'] = match.group(1).strip()

            match = REGEX_PATTERNS['razao_social'].search(texto)
            if match:
                nome = match.group(1).strip()
                if len(nome) < 100:
                    dados['cliente_nome'] = nome

            # 6. EXTRAIR DADOS DO PRESTADOR
            match = REGEX_PATTERNS['cnpj_prestador'].search(texto)
            if match:
                dados['prestador_cnpj'] = match.group(1).strip()

            match = REGEX_PATTERNS['nome_fantasia'].search(texto)
            if match:
                nome = match.group(1).strip()
                if len(nome) < 100:
                    dados['prestador_nome'] = nome

            # 7. EXTRAIR SERVIÇOS
            match = REGEX_PATTERNS['discriminacao'].search(texto)
            if match:
                servicos = match.group(1).strip()
                servicos = re.sub(r'\d+,\d+', '', servicos)
                servicos = re.sub(r'\s+', ' ', servicos)
                if 10 < len(servicos) < 500:
                    dados['servicos'] = servicos

        print(f"OK PDF processado: {dados['numero']} - R$ {dados['valor_total']:.2f}")
        return dados

    except Exception as e:
        print(f"ERRO Erro ao processar PDF: {e}")
        return dados


def obter_contexto_completo():
    """Busca todos os dados do banco de dados (documentos: laudos, recibos, orçamentos)"""
    conn = get_db()

    docs = conn.execute('''
        SELECT tipo_doc, nome_arquivo, valor_total, data_geracao, cliente_nome
        FROM documentos_gerados
        ORDER BY data_geracao DESC
    ''').fetchall()

    clientes = conn.execute('''
        SELECT nome_fantasia, razao_social, cnpj, cidade, telefone
        FROM clientes_web
        ORDER BY nome_fantasia
    ''').fetchall()

    return {
        'documentos': [dict(d) for d in docs],
        'clientes': [dict(c) for c in clientes]
    }


def calcular_estatisticas(dados):
    """Calcula estatísticas de documentos (laudos, recibos, orçamentos)"""
    docs = dados.get('documentos', [])

    if not docs:
        return {}

    # Contadores por tipo
    laudos = 0
    recibos = 0
    orcamentos = 0
    valor_recibos = 0.0
    por_mes = defaultdict(float)

    for doc in docs:
        tipo = (doc.get('tipo_doc') or '').upper()
        valor = doc.get('valor_total') or 0

        if 'LAUDO' in tipo:
            laudos += 1
        elif 'RECIBO' in tipo:
            recibos += 1
            valor_recibos += valor
        elif 'ORÇAMENTO' in tipo or 'ORCAMENTO' in tipo:
            orcamentos += 1

        data = doc.get('data_geracao')
        if data and len(data) >= 7:
            mes = data[:7]
            por_mes[mes] += valor or 0

    return {
        'total_documentos': len(docs),
        'laudos': laudos,
        'recibos': recibos,
        'orcamentos': orcamentos,
        'valor_recibos': valor_recibos,
        'faturamento_mensal': dict(por_mes)
    }


def consultar_ia_analise(pergunta_usuario, tipo_analise='geral'):
    """
    Consulta Gemini API com análise de dados de documentos
    """
    try:
        dados = obter_contexto_completo()
        stats = calcular_estatisticas(dados)

        # Preparar contexto
        contexto_dados = f"""
=== ESTATÍSTICAS DE DOCUMENTOS ===
Total de Documentos: {stats.get('total_documentos', 0)}
Laudos: {stats.get('laudos', 0)}
Recibos: {stats.get('recibos', 0)}
Orçamentos: {stats.get('orcamentos', 0)}
Valor Total dos Recibos: R$ {stats.get('valor_recibos', 0):.2f}

=== DOCUMENTOS RECENTES (Top 10) ===
"""

        for i, doc in enumerate(dados['documentos'][:10], 1):
            tipo = doc.get('tipo_doc') or 'Documento'
            cliente = doc.get('cliente_nome') or 'Cliente'
            valor = doc.get('valor_total') or 0
            data = doc.get('data_geracao') or 'Sem data'
            contexto_dados += f"{i}. {tipo} - {cliente} - R$ {valor:.2f} ({data})\n"

        # Chamar Gemini API
        api_key = os.environ.get("GEMINI_API_KEY")

        if not api_key or api_key == "SUBSTITUA-PELA-SUA-CHAVE-GEMINI-AQUI":
            return {
                "sucesso": False,
                "erro": "API Key não configurada ou inválida",
                "resposta": "Configure a GEMINI_API_KEY no arquivo .env. Obtenha em: https://makersuite.google.com/app/apikey"
            }

        # Criar modelo (atualizado para Gemini 2.5)
        model = genai.GenerativeModel('gemini-2.5-flash')

        # Preparar prompt completo
        prompt_completo = f"""Você é um assistente de análise de dados de negócios.
Responda perguntas sobre documentos (laudos, recibos, orçamentos) e análises de desempenho.
Seja direto, use dados específicos e forneça insights quando relevante.

{contexto_dados}

PERGUNTA: {pergunta_usuario}"""

        # Gerar resposta
        response = model.generate_content(prompt_completo)

        return {
            "sucesso": True,
            "resposta": response.text,
            "tipo_analise": tipo_analise,
            "modelo": "gemini-2.5-flash"
        }

    except Exception as e:
        traceback.print_exc()
        return {
            "sucesso": False,
            "erro": str(e),
            "resposta": f"Erro ao processar: {str(e)}"
        }


def gerar_insights_automaticos():
    """Gera insights automáticos sobre documentos"""
    pergunta = """Analise os dados de documentos (laudos, recibos, orçamentos) e forneça:
1. 3 principais insights sobre a produção de documentos
2. 2 oportunidades de crescimento no volume de serviços
3. 1 alerta importante

Seja específico e use números reais."""

    return consultar_ia_analise(pergunta, tipo_analise='documentos')


def gerar_analise_completa_faturamento():
    """
    Gera análise completa e detalhada de documentos
    Focado em laudos, recibos e orçamentos com análise profunda
    """
    try:
        conn = get_db()

        # Buscar TODOS os documentos
        todos_docs = conn.execute('''
            SELECT tipo_doc, nome_arquivo, valor_total, data_geracao, cliente_nome
            FROM documentos_gerados
            ORDER BY data_geracao DESC
        ''').fetchall()

        # Separar por tipo
        laudos = [d for d in todos_docs if 'LAUDO' in (d['tipo_doc'] or '').upper()]
        recibos = [d for d in todos_docs if 'RECIBO' in (d['tipo_doc'] or '').upper()]
        orcamentos = [d for d in todos_docs if 'ORÇAMENTO' in (d['tipo_doc'] or '').upper() or 'ORCAMENTO' in (d['tipo_doc'] or '').upper()]

        # Calcular totais
        total_valor_recibos = sum([r['valor_total'] or 0 for r in recibos])

        # Análise por mês e cliente
        por_mes = defaultdict(lambda: {'laudos': 0, 'recibos': 0, 'orcamentos': 0, 'valor': 0})
        por_cliente = defaultdict(lambda: {'docs': 0, 'valor': 0})
        valores_recibos = []

        for doc in todos_docs:
            tipo = (doc['tipo_doc'] or '').upper()
            data = doc['data_geracao']
            valor = doc['valor_total'] or 0
            cliente = doc['cliente_nome'] or 'Não identificado'

            if data and len(data) >= 7:
                mes = data[:7]
                if 'LAUDO' in tipo:
                    por_mes[mes]['laudos'] += 1
                elif 'RECIBO' in tipo:
                    por_mes[mes]['recibos'] += 1
                    por_mes[mes]['valor'] += valor
                    valores_recibos.append(valor)
                elif 'ORÇAMENTO' in tipo or 'ORCAMENTO' in tipo:
                    por_mes[mes]['orcamentos'] += 1

            por_cliente[cliente]['docs'] += 1
            por_cliente[cliente]['valor'] += valor

        # Calcular estatísticas
        media_mensal = sum([m['valor'] for m in por_mes.values()]) / len(por_mes) if por_mes else 0

        # Ticket médio dos recibos
        ticket_medio = sum(valores_recibos) / len(valores_recibos) if valores_recibos else 0
        ticket_maximo = max(valores_recibos) if valores_recibos else 0
        ticket_minimo = min(valores_recibos) if valores_recibos else 0

        # Top 5 clientes por quantidade de documentos
        top_clientes = sorted(por_cliente.items(), key=lambda x: x[1]['docs'], reverse=True)[:5]

        # Preparar relatório
        relatorio = f"""
=== ANÁLISE COMPLETA DE DOCUMENTOS ===

RESUMO:
- Total de Laudos: {len(laudos)}
- Total de Recibos: {len(recibos)} (R$ {total_valor_recibos:,.2f})
- Total de Orçamentos: {len(orcamentos)}
- Total Geral: {len(todos_docs)} documentos

INDICADORES:
- Ticket Médio (Recibos): R$ {ticket_medio:,.2f}
- Maior Recibo: R$ {ticket_maximo:,.2f}
- Menor Recibo: R$ {ticket_minimo:,.2f}
- Média Mensal (Valor): R$ {media_mensal:,.2f}

TOP 5 CLIENTES (por quantidade):
"""
        for i, (cliente, dados) in enumerate(top_clientes, 1):
            relatorio += f"{i}. {cliente}: {dados['docs']} docs - R$ {dados['valor']:,.2f}\n"

        relatorio += f"\nEVOLUÇÃO MENSAL:\n"
        for mes, dados in sorted(por_mes.items(), reverse=True)[:12]:
            relatorio += f"- {mes}: {dados['laudos']} laudos, {dados['recibos']} recibos (R$ {dados['valor']:,.2f}), {dados['orcamentos']} orçamentos\n"

        relatorio += f"\nRECIBOS RECENTES (Top 20):\n"
        for i, recibo in enumerate(recibos[:20], 1):
            relatorio += f"{i}. {recibo['nome_arquivo']} - R$ {recibo['valor_total'] or 0:,.2f} - {recibo['data_geracao']}\n"

        # Chamar Gemini para análise
        api_key = os.environ.get("GEMINI_API_KEY")

        if not api_key or api_key == "SUBSTITUA-PELA-SUA-CHAVE-GEMINI-AQUI":
            return {
                "sucesso": False,
                "erro": "API Key não configurada",
                "dados_brutos": {
                    "total_laudos": len(laudos),
                    "total_recibos": len(recibos),
                    "total_orcamentos": len(orcamentos),
                    "valor_recibos": total_valor_recibos,
                    "relatorio": relatorio
                }
            }

        model = genai.GenerativeModel('gemini-2.5-flash')

        prompt = f"""{relatorio}

Com base nesses dados de documentos (laudos, recibos, orçamentos), forneça uma análise detalhada:

1. **DIAGNÓSTICO DE PRODUTIVIDADE**
   - Volume de documentos gerados
   - Tendência de crescimento ou declínio
   - Proporção entre tipos de documentos

2. **ANÁLISE DE CLIENTES**
   - Clientes mais ativos
   - Diversificação da base
   - Oportunidades de expansão

3. **ANÁLISE DE VALORES (RECIBOS)**
   - Ticket médio e sua variação
   - Distribuição de valores
   - Oportunidades de precificação

4. **TENDÊNCIAS E PADRÕES**
   - Sazonalidade na produção
   - Meses de maior/menor atividade
   - Padrões identificados

5. **RECOMENDAÇÕES**
   - 3 ações para aumentar produtividade
   - 2 oportunidades de crescimento
   - 1 alerta importante

Seja específico e use os números reais do relatório."""

        response = model.generate_content(prompt)

        return {
            "sucesso": True,
            "analise_ia": response.text,
            "dados_brutos": {
                "total_laudos": len(laudos),
                "total_recibos": len(recibos),
                "total_orcamentos": len(orcamentos),
                "valor_recibos": total_valor_recibos,
                "faturamento_mensal": {k: v['valor'] for k, v in sorted(por_mes.items(), reverse=True)},
                "ticket_medio": ticket_medio,
                "top_clientes": {c: d for c, d in top_clientes}
            },
            "relatorio": relatorio
        }

    except Exception as e:
        traceback.print_exc()
        return {
            "sucesso": False,
            "erro": str(e),
            "detalhes": traceback.format_exc()
        }


@app.route('/')
def index():
    # Buscar estatísticas para o dashboard bonito
    stats = {
        'clientes': 0,
        'empresas': 0,
        'documentos': 0
    }

    try:
        db = get_db()

        # Contar clientes
        stats['clientes'] = db.execute('SELECT COUNT(*) FROM clientes_web').fetchone()[0]

        # Contar empresas da Receita Federal
        cnpj_db = sqlite3.connect(str(CNPJ_DB_PATH))
        stats['empresas'] = cnpj_db.execute('SELECT COUNT(*) FROM empresas').fetchone()[0]
        cnpj_db.close()

        # Contar documentos gerados
        stats['documentos'] = db.execute('SELECT COUNT(*) FROM documentos_gerados').fetchone()[0]
    except Exception as e:
        print(f"Erro ao buscar estatísticas: {e}")

    return render_template('index.html', stats=stats)

@app.route('/dashboard')
def dashboard():
    return render_template('dashboard_novo.html')

@app.route('/dashboard-antigo')
def dashboard_antigo():
    return render_template('dashboard.html')

@app.route('/clientes')
def clientes():
    return render_template('clientes.html')

@app.route('/empresas')
def empresas():
    return render_template('empresas.html')

@app.route('/gerar')
def gerar():
    return render_template('gerar.html')


@app.route('/documentos')
def documentos():
    # Redirecionar para dashboard_novo.html que já existe e funciona
    return render_template('dashboard_novo.html')

# =============================================================================
# ROTAS DE CLIENTES - MOVIDAS PARA blueprints/clientes.py
# =============================================================================

@app.route('/api/cnpj/autocomplete', methods=['GET'])
def cnpj_autocomplete():
    """
    Autocomplete de CNPJs da Receita Federal

    Query params:
        campo: cnpj, razao_social, nome_fantasia, logradouro, bairro, cnae_fiscal
        termo: Termo de busca (mínimo 2 caracteres)
        uf: Filtro por UF (opcional)
        limit: Limite de resultados (default: 10, max: 50)

    Exemplo: /api/cnpj/autocomplete?campo=nome_fantasia&termo=Padaria&uf=MG&limit=10
    """
    if not api_cnpj:
        return jsonify({'erro': 'Autocomplete de CNPJs não disponível'}), 503

    try:
        campo = request.args.get('campo')
        termo = request.args.get('termo')

        if not campo or not termo:
            return jsonify({'erro': 'Parâmetros obrigatórios: campo, termo'}), 400

        campos_validos = ['cnpj', 'razao_social', 'nome_fantasia', 'logradouro', 'bairro', 'cnae_fiscal']
        if campo not in campos_validos:
            return jsonify({'erro': f'Campo inválido. Use: {", ".join(campos_validos)}'}), 400

        if len(termo.strip()) < 2:
            return jsonify({'erro': 'Termo deve ter no mínimo 2 caracteres'}), 400

        filtros = {}
        if request.args.get('uf'):
            filtros['uf'] = request.args.get('uf').upper()

        try:
            limit = int(request.args.get('limit', 10))
            limit = max(1, min(limit, 50))
        except ValueError:
            limit = 10

        resultados = api_cnpj.buscar_sugestoes(campo, termo, filtros, limit)

        return jsonify({
            'sucesso': True,
            'total': len(resultados),
            'resultados': resultados
        })

    except ValueError as e:
        return jsonify({'erro': str(e)}), 400
    except Exception as e:
        print(f"[ERRO] Autocomplete CNPJ: {e}")
        traceback.print_exc()
        return jsonify({'erro': 'Erro interno no servidor'}), 500


@app.route('/api/garantias/vencimentos', methods=['GET'])
def api_garantias_vencimentos():
    """
    Retorna clientes com garantia vencida ou próxima ao vencimento
    OTIMIZAÇÃO: Cálculo movido para SQL (10x mais rápido)

    Query params:
        dias_antecedencia: número de dias para alertar antes do vencimento (padrão: 30)
    """
    try:
        conn = get_db()
        dias_antecedencia = request.args.get('dias_antecedencia', 30, type=int)

        # OTIMIZAÇÃO: Fazer todos os cálculos no SQL ao invés de Python (10x mais rápido)
        avisos = conn.execute('''
            SELECT
                id as cliente_id,
                nome_fantasia as cliente_nome,
                razao_social,
                cnpj,
                telefone,
                email,
                cidade,
                data_garantia,
                COALESCE(periodo_garantia_meses, 12) as periodo_meses,
                date(data_garantia, '+' || COALESCE(periodo_garantia_meses, 12) || ' months') as data_vencimento,
                CAST((julianday(date(data_garantia, '+' || COALESCE(periodo_garantia_meses, 12) || ' months')) - julianday('now')) AS INTEGER) as dias_restantes,
                CASE
                    WHEN CAST((julianday(date(data_garantia, '+' || COALESCE(periodo_garantia_meses, 12) || ' months')) - julianday('now')) AS INTEGER) < 0 THEN 'vencida'
                    WHEN CAST((julianday(date(data_garantia, '+' || COALESCE(periodo_garantia_meses, 12) || ' months')) - julianday('now')) AS INTEGER) <= 7 THEN 'vence_esta_semana'
                    ELSE 'proxima_vencer'
                END as status,
                CASE
                    WHEN CAST((julianday(date(data_garantia, '+' || COALESCE(periodo_garantia_meses, 12) || ' months')) - julianday('now')) AS INTEGER) < 0 THEN 'critico'
                    WHEN CAST((julianday(date(data_garantia, '+' || COALESCE(periodo_garantia_meses, 12) || ' months')) - julianday('now')) AS INTEGER) <= 7 THEN 'urgente'
                    ELSE 'atencao'
                END as urgencia
            FROM clientes_web
            WHERE data_garantia IS NOT NULL
                AND data_garantia != ''
                AND CAST((julianday(date(data_garantia, '+' || COALESCE(periodo_garantia_meses, 12) || ' months')) - julianday('now')) AS INTEGER) <= ?
            ORDER BY dias_restantes ASC
        ''', (dias_antecedencia,)).fetchall()

        # Converter para lista de dicts
        avisos_list = [dict(row) for row in avisos]

        # Ordenar por urgência e dias restantes (já vem ordenado do SQL, mas garantir)
        ordem_urgencia = {'critico': 0, 'urgente': 1, 'atencao': 2}
        avisos_list.sort(key=lambda x: (ordem_urgencia[x['urgencia']], x['dias_restantes']))

        return jsonify({
            'sucesso': True,
            'total': len(avisos_list),
            'avisos': avisos_list,
            'resumo': {
                'vencidas': len([a for a in avisos_list if a['status'] == 'vencida']),
                'esta_semana': len([a for a in avisos_list if a['status'] == 'vence_esta_semana']),
                'proximas': len([a for a in avisos_list if a['status'] == 'proxima_vencer'])
            }
        })

    except Exception as e:
        print(f"[ERRO] Garantias vencimentos: {e}")
        traceback.print_exc()
        return jsonify({'erro': str(e)}), 500


@app.route('/api/clientes/<int:cliente_id>/garantia', methods=['PUT'])
def atualizar_garantia_cliente(cliente_id):
    """
    Atualiza data de garantia e período de garantia de um cliente

    Body JSON:
        {
            "data_garantia": "2025-01-14",  // Data do último serviço/garantia
            "periodo_garantia_meses": 12     // Período em meses (padrão: 12)
        }
    """
    try:
        dados = request.json
        data_garantia = dados.get('data_garantia')
        periodo_meses = dados.get('periodo_garantia_meses', 12)

        if not data_garantia:
            return jsonify({'erro': 'Data de garantia é obrigatória'}), 400

        # Validar formato da data
        try:
            datetime.strptime(data_garantia, '%Y-%m-%d')
        except ValueError:
            return jsonify({'erro': 'Formato de data inválido. Use YYYY-MM-DD'}), 400

        conn = get_db()
        conn.execute('''
            UPDATE clientes_web
            SET data_garantia = ?,
                periodo_garantia_meses = ?
            WHERE id = ?
        ''', (data_garantia, periodo_meses, cliente_id))
        conn.commit()

        return jsonify({
            'sucesso': True,
            'mensagem': 'Garantia atualizada com sucesso'
        })

    except Exception as e:
        print(f"[ERRO] Atualizar garantia: {e}")
        traceback.print_exc()
        return jsonify({'erro': str(e)}), 500


@app.route('/api/cnae/formatar', methods=['POST'])
def formatar_cnae_api():
    """
    Formata código CNAE com descrição completa

    Body JSON:
        cnae: Código CNAE (com ou sem formatação)

    Retorna:
        {"cnae_formatado": "81.22-2-00 - Imunização e controle de pragas urbanas"}

    Exemplo: {"cnae": "8122200"} → {"cnae_formatado": "81.22-2-00 - Imunização..."}
    """
    try:
        dados = request.json
        cnae_codigo = dados.get('cnae', '')

        if not cnae_codigo:
            return jsonify({'erro': 'Código CNAE não fornecido'}), 400

        cnae_formatado = formatar_cnae(cnae_codigo)

        return jsonify({
            'sucesso': True,
            'cnae_formatado': cnae_formatado
        })

    except Exception as e:
        print(f"[ERRO] Formatar CNAE: {e}")
        traceback.print_exc()
        return jsonify({'erro': 'Erro ao formatar CNAE'}), 500


@app.route('/api/empresa/sync', methods=['POST'])
def sync_empresa_receita():
    """
    Sincroniza empresa do cnpj_filtrado.db para gestao_documentos.db

    Workflow completo:
    1. Busca empresa no cnpj_filtrado.db
    2. Valida CNAE contra cnaes_permitidos.txt
    3. Transforma dados para formato interno
    4. Insere no gestao_documentos.db (tabela clientes_web)
    5. Retorna ID do cliente para geração de documentos

    Body JSON:
        {
            "cnpj": "12345678000190",
            "validar_cnae": true  // opcional, padrão true
        }

    Resposta:
        {
            "sucesso": true,
            "cliente_id": 123,
            "mensagem": "Empresa sincronizada com sucesso",
            "dados": {...}
        }
    """
    if not data_bridge:
        return jsonify({
            'erro': 'Data Bridge não disponível',
            'detalhes': 'cnpj_filtrado.db ou cnaes_permitidos.txt não encontrado'
        }), 503

    try:
        dados = request.json
        cnpj = dados.get('cnpj', '').strip()
        validar_cnae = dados.get('validar_cnae', True)

        if not cnpj:
            return jsonify({'erro': 'CNPJ é obrigatório'}), 400

        # Limpar CNPJ (apenas dígitos)
        cnpj_limpo = re.sub(r'[^\d]', '', cnpj)

        if len(cnpj_limpo) != 14:
            return jsonify({'erro': 'CNPJ inválido (deve ter 14 dígitos)'}), 400

        # PASSO 1: Buscar empresa na Receita Federal
        print(f"[SYNC] Buscando empresa {cnpj_limpo}...")
        receita_data = data_bridge.buscar_empresa_receita(cnpj_limpo)

        if not receita_data:
            return jsonify({
                'erro': 'Empresa não encontrada',
                'detalhes': f'CNPJ {cnpj} não existe no cnpj_filtrado.db'
            }), 404

        # PASSO 2: Validar CNAE (se solicitado)
        cnae = receita_data.get('cnae_fiscal_principal', '')
        if validar_cnae and not data_bridge.validar_cnae(cnae):
            return jsonify({
                'erro': 'CNAE não permitido',
                'detalhes': f'CNAE {cnae} não está na lista de CNAEs estratégicos',
                'cnae': cnae,
                'empresa': receita_data.get('nome_fantasia') or receita_data.get('razao_social')
            }), 400

        # PASSO 3: Transformar dados
        print(f"[SYNC] Transformando dados da empresa...")
        dados_transformados = data_bridge.transform_empresa_data(receita_data)

        # PASSO 4: Inserir no gestao_documentos.db (usando tabela clientes_web)
        conn = get_db()

        # Verificar se já existe
        cliente_existente = conn.execute(
            "SELECT id FROM clientes_web WHERE cnpj = ?",
            (dados_transformados['cnpj_numerico'],)
        ).fetchone()

        if cliente_existente:
            cliente_id = cliente_existente['id']
            print(f"[SYNC] Cliente já existe (ID: {cliente_id})")
        else:
            # Inserir novo cliente
            endereco_completo = dados_transformados['endereco_completo']

            cursor = conn.execute('''
                INSERT INTO clientes_web (
                    nome_fantasia, razao_social, cnpj, cnae,
                    rua, numero, bairro, cidade, uf, endereco_completo,
                    telefone, data_garantia, periodo_garantia_meses
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (
                dados_transformados['nome_fantasia'],
                dados_transformados['razao_social'],
                dados_transformados['cnpj_numerico'],
                dados_transformados['cnae_numerico'],
                dados_transformados['rua'],
                dados_transformados['numero'],
                dados_transformados['bairro'],
                dados_transformados['cidade'],
                dados_transformados['uf'],
                endereco_completo,
                dados_transformados['telefone'],
                None,  # data_garantia
                12     # periodo_garantia_meses padrão
            ))

            cliente_id = cursor.lastrowid
            conn.commit()

            print(f"[SYNC] Cliente inserido (ID: {cliente_id})")

        # PASSO 5: Retornar dados completos
        return jsonify({
            'sucesso': True,
            'cliente_id': cliente_id,
            'mensagem': 'Empresa sincronizada com sucesso',
            'dados': {
                'id': cliente_id,
                'nome_fantasia': dados_transformados['nome_fantasia'],
                'razao_social': dados_transformados['razao_social'],
                'cnpj': dados_transformados['cnpj'],
                'cnae': dados_transformados['cnae'],
                'endereco': dados_transformados['endereco_completo'],
                'telefone': dados_transformados['telefone'],
                'email': dados_transformados.get('email', ''),
                'categoria_lead': dados_transformados.get('categoria_lead', 'normal'),
                'validado_cnae': dados_transformados.get('validado_por_cnae', False)
            },
            'estatisticas': data_bridge.get_estatisticas()
        })

    except Exception as e:
        print(f"[ERRO] Sync empresa: {e}")
        traceback.print_exc()
        return jsonify({
            'erro': 'Erro ao sincronizar empresa',
            'detalhes': str(e)
        }), 500


@app.route('/api/modelos/listar')
def listar_modelos():
    """
    Lista todos os modelos .docx disponíveis na pasta modelos/
    e detecta as variáveis {{...}} em cada modelo
    """
    try:
        modelos = []

        if not MODELOS_DIR.exists():
            return jsonify({
                'erro': 'Pasta de modelos não encontrada'
            }), 404

        # Listar arquivos .docx na pasta modelos
        for arquivo in MODELOS_DIR.glob('*.docx'):
            if arquivo.name.startswith('~$'):  # Ignorar arquivos temporários
                continue

            try:
                # Extrair variáveis do modelo
                variaveis = extrair_variaveis_modelo(arquivo)

                modelos.append({
                    'nome': arquivo.stem,
                    'arquivo': arquivo.name,
                    'caminho': str(arquivo),
                    'tamanho': arquivo.stat().st_size,
                    'variaveis': variaveis,
                    'total_variaveis': len(variaveis)
                })
            except Exception as e:
                print(f"[ERRO] Erro ao processar modelo {arquivo.name}: {e}")
                continue

        return jsonify({
            'modelos': modelos,
            'total': len(modelos)
        })

    except Exception as e:
        print(f"[ERRO] Listar modelos: {e}")
        traceback.print_exc()
        return jsonify({
            'erro': 'Erro ao listar modelos',
            'detalhes': str(e)
        }), 500


def extrair_variaveis_modelo(caminho_modelo):
    """
    Extrai todas as variáveis {{...}} de um modelo .docx

    Args:
        caminho_modelo: Path do arquivo .docx

    Returns:
        Lista de variáveis únicas encontradas
    """
    variaveis = set()

    try:
        from docx import Document
        import pathlib

        # Converter para Path para garantir encoding correto
        if isinstance(caminho_modelo, str):
            caminho_modelo = pathlib.Path(caminho_modelo)

        print(f"[DEBUG] Tentando ler modelo: {caminho_modelo.name}")

        # Usar o caminho absoluto com encoding UTF-8
        doc = Document(caminho_modelo)
        print(f"[DEBUG] Modelo carregado com sucesso: {caminho_modelo.name}")

        # Buscar variáveis no documento
        padrao = r'\{\{([^}]+)\}\}'

        # Verificar em todos os parágrafos
        for paragrafo in doc.paragraphs:
            matches = re.findall(padrao, paragrafo.text)
            for match in matches:
                variaveis.add(match.strip())

        # Verificar em tabelas
        for tabela in doc.tables:
            for linha in tabela.rows:
                for celula in linha.cells:
                    for paragrafo in celula.paragraphs:
                        matches = re.findall(padrao, paragrafo.text)
                        for match in matches:
                            variaveis.add(match.strip())

        print(f"[DEBUG] Variáveis encontradas em {caminho_modelo.name}: {len(variaveis)} variáveis")
        return sorted(list(variaveis))

    except PermissionError as e:
        print(f"[ERRO] Arquivo bloqueado ou aberto: {caminho_modelo.name} - {e}")
        print(f"[DICA] Feche o arquivo {caminho_modelo.name} no Word ou outro programa e tente novamente")
        return []
    except Exception as e:
        print(f"[ERRO] Extrair variáveis de {caminho_modelo.name}: {e}")
        import traceback
        traceback.print_exc()
        # Retornar lista vazia mas não impedir que o modelo apareça na lista
        return []


@app.route('/api/documento/gerar-inteligente', methods=['POST'])
def gerar_documento_inteligente():
    """
    Gera documento inteligentemente preenchendo variáveis {{...}} automaticamente

    Body JSON:
        {
            "cliente_id": 123,
            "modelo": "LaudoPython2.docx",
            "variaveis_extras": {  // opcional
                "servico": "Dedetização",
                "valor": "R$ 500,00"
            }
        }
    """
    try:
        dados = request.json
        cliente_id = dados.get('cliente_id')
        modelo_nome = dados.get('modelo')
        variaveis_extras = dados.get('variaveis_extras', {})

        if not cliente_id or not modelo_nome:
            return jsonify({
                'erro': 'cliente_id e modelo são obrigatórios'
            }), 400

        # Buscar dados do cliente
        conn = get_db()
        cliente = conn.execute(
            'SELECT * FROM clientes_web WHERE id = ?',
            (cliente_id,)
        ).fetchone()

        if not cliente:
            return jsonify({
                'erro': 'Cliente não encontrado'
            }), 404

        cliente_dict = dict(cliente)

        # Caminho do modelo
        modelo_path = MODELOS_DIR / modelo_nome

        if not modelo_path.exists():
            return jsonify({
                'erro': 'Modelo não encontrado'
            }), 404

        # Preparar context com dados do cliente (passar nome do modelo)
        context = preparar_context_inteligente(cliente_dict, variaveis_extras, modelo_nome)

        # Gerar documento
        doc = DocxTemplate(modelo_path)
        doc.render(context)

        # Salvar documento
        nome_cliente = cliente_dict.get('nome_fantasia') or cliente_dict.get('razao_social', 'Cliente')

        # Obter categoria baseada no CNAE do cliente
        cnae_cliente = cliente_dict.get('cnae', '')
        categoria = obter_categoria_por_cnae(cnae_cliente)

        # Se não encontrar categoria, usar #99 (Outros/Sem CNAE)
        if categoria:
            sufixo = categoria  # Ex: #0, #1, #2, #3
        else:
            sufixo = '#99'  # Empresas sem CNAE ou CNAE não mapeado

        nome_arquivo = f"{nome_cliente}_{modelo_path.stem}_{sufixo}.docx"
        nome_arquivo = re.sub(r'[<>:"/\\|?*]', '_', nome_arquivo)  # Remover caracteres inválidos

        output_path = OUTPUT_DIR / 'documentos' / nome_arquivo
        output_path.parent.mkdir(exist_ok=True, parents=True)

        doc.save(str(output_path))

        print(f"[OK] Documento gerado: {nome_arquivo}")

        return jsonify({
            'sucesso': True,
            'arquivo': nome_arquivo,
            'caminho': str(output_path),
            'cliente': cliente_dict['nome_fantasia'],
            'modelo': modelo_path.stem,
            'variaveis_preenchidas': list(context.keys())
        })

    except Exception as e:
        print(f"[ERRO] Gerar documento inteligente: {e}")
        traceback.print_exc()
        return jsonify({
            'erro': 'Erro ao gerar documento',
            'detalhes': str(e)
        }), 500


def preparar_context_inteligente(cliente, variaveis_extras=None, nome_modelo=''):
    """
    Prepara context para DocxTemplate com dados do cliente
    Mapeia campos do banco para variáveis comuns em templates

    Args:
        cliente: Dicionário com dados do cliente
        variaveis_extras: Variáveis adicionais fornecidas pelo usuário
        nome_modelo: Nome do arquivo de modelo (para lógica condicional)
    """
    if variaveis_extras is None:
        variaveis_extras = {}

    # Formatar CNPJ: 10409228000193 -> 10.409.228/0001-93
    def formatar_cnpj(cnpj):
        if not cnpj:
            return ''
        cnpj_limpo = re.sub(r'\D', '', str(cnpj))
        if len(cnpj_limpo) == 14:
            return f"{cnpj_limpo[:2]}.{cnpj_limpo[2:5]}.{cnpj_limpo[5:8]}/{cnpj_limpo[8:12]}-{cnpj_limpo[12:]}"
        return cnpj

    # Formatar endereço (substituir código IBGE 4445 por Divinópolis)
    def formatar_endereco(texto):
        if not texto:
            return ''
        # Substituir código 4445 por Divinópolis em qualquer parte do endereço
        texto_formatado = str(texto).replace('4445', 'Divinópolis')
        # Limpar vírgulas duplas que podem aparecer
        texto_formatado = texto_formatado.replace(', ,', ',').strip()
        return texto_formatado

    # Formatar cidade
    cidade = cliente.get('cidade', '')
    if cidade == '4445':
        cidade = 'Divinópolis'

    cnpj_formatado = formatar_cnpj(cliente.get('cnpj', ''))
    cnae_formatado = formatar_cnae(cliente.get('cnae', ''))

    # CAMPOS SEPARADOS - nunca espelhar um no outro
    nome_fantasia = cliente.get('nome_fantasia', '') or ''  # Vazio se não existir
    razao_social = cliente.get('razao_social', '') or ''    # Vazio se não existir

    # Detectar tipo de documento pelo nome do modelo
    nome_modelo_lower = nome_modelo.lower()
    eh_laudo = 'laudo' in nome_modelo_lower
    eh_recibo_orcamento = 'recibo' in nome_modelo_lower or 'orcamento' in nome_modelo_lower or 'orçamento' in nome_modelo_lower

    # LÓGICA CONDICIONAL para campos de nome:
    # - LAUDOS: nome_razao_social usa Razão Social, nome_fantasia fica normal
    # - RECIBOS/ORÇAMENTOS: nome_fantasia em CAIXA ALTA, nome_razao_social usa Razão Social
    if eh_laudo:
        # Para Laudos: nome_razao_social = razão social
        nome_razao_social = razao_social if razao_social else nome_fantasia
        nome_fantasia_final = nome_fantasia
    elif eh_recibo_orcamento:
        # Para Recibos/Orçamentos:
        # {{nome_fantasia}} = DEDETIZADORA BORGES (em CAPS)
        # {{nome_razao_social}} = MARIA APARECIDA DE OLIVEIRA BORGES 774.017.436-04
        nome_razao_social = razao_social  # Razão Social completa
        nome_fantasia_final = nome_fantasia.upper() if nome_fantasia else ''  # Nome Fantasia em CAPS
    else:
        # Padrão (outros documentos)
        nome_razao_social = razao_social if razao_social else nome_fantasia
        nome_fantasia_final = nome_fantasia

    # Endereços formatados (substituir 4445 por Divinópolis)
    endereco_completo_formatado = formatar_endereco(cliente.get('endereco_completo', ''))
    rua_formatada = formatar_endereco(cliente.get('rua', ''))
    bairro_formatado = formatar_endereco(cliente.get('bairro', ''))

    # Context base com dados do cliente
    context = {
        # Identificação - CAMPOS SEPARADOS
        'nome': nome_fantasia_final,  # Laudos: normal | Recibos: CAPS
        'nome_fantasia': nome_fantasia_final,  # Laudos: normal | Recibos: CAPS (DEDETIZADORA BORGES)
        'nome_razao_social': nome_razao_social,  # Sempre Razão Social (MARIA APARECIDA...)
        'razao_social': razao_social,  # Razão Social original do banco
        'cliente': nome_fantasia or razao_social,  # Fallback para compatibilidade
        'empresa': nome_fantasia or razao_social,  # Fallback para compatibilidade

        # Documentos (múltiplos aliases) - FORMATADOS
        'cnpj': cnpj_formatado,
        'cnpj_cliente': cnpj_formatado,
        'cnae': cnae_formatado,
        'cnae_cliente': cnae_formatado,

        # Endereço - FORMATADOS (4445 → Divinópolis)
        'rua': rua_formatada,
        'numero': cliente.get('numero', ''),
        'bairro': bairro_formatado,
        'cidade': cidade,
        'uf': cliente.get('uf', ''),
        'cep': cliente.get('cep', ''),
        'endereco': endereco_completo_formatado,
        'endereco_completo': endereco_completo_formatado,
        'endereco_cliente': endereco_completo_formatado,

        # Contato
        'telefone': cliente.get('telefone', ''),
        'email': cliente.get('email', ''),

        # Datas
        'DATA': datetime.now().strftime('%d/%m/%Y'),
        'data': datetime.now().strftime('%d/%m/%Y'),
        'data_hoje': datetime.now().strftime('%d/%m/%Y'),
        'data_atual': datetime.now().strftime('%d/%m/%Y'),
        'mes': datetime.now().strftime('%B'),
        'ano': datetime.now().strftime('%Y'),
        'hora': datetime.now().strftime('%H:%M'),
    }

    # ORÇAMENTOS: Aplicar formatação especial de data (por extenso)
    eh_orcamento = 'orcamento' in nome_modelo_lower or 'orçamento' in nome_modelo_lower
    if eh_orcamento:
        # Converter DATA para formato extenso (ex: "09 de janeiro de 2026")
        data_obj = datetime.now()
        context['DATA'] = formatar_data_extenso(data_obj)
        print(f"[DEBUG ORÇAMENTO] Data por extenso: {context['DATA']}")

    # Adicionar variáveis extras (fornecidas pelo usuário)
    context.update(variaveis_extras)

    # ORÇAMENTOS: Converter preço para extenso se houver valor de preço
    if eh_orcamento:
        print(f"[DEBUG ORÇAMENTO] Variáveis extras recebidas: {variaveis_extras}")
        # Tentar pegar o preço de diferentes fontes
        preco_valor = variaveis_extras.get('preco') or variaveis_extras.get('preço') or variaveis_extras.get('valor')
        if preco_valor:
            preco_extenso_calculado = numero_por_extenso(preco_valor)
            context['preco_extenso'] = preco_extenso_calculado
            context['preço'] = preco_valor
            print(f"[DEBUG ORÇAMENTO] Preço: {preco_valor} -> Por extenso: {preco_extenso_calculado}")
        else:
            print(f"[DEBUG ORÇAMENTO] AVISO: Nenhum valor de preço encontrado!")

    # Mapear servico -> serviço (com acento) para compatibilidade com templates
    if 'servico' in context:
        context['serviço'] = context['servico']
    if 'servico1' in context:
        context['serviço1'] = context['servico1']
    if 'servico2' in context:
        context['serviço2'] = context['servico2']
    if 'servico3' in context:
        context['serviço3'] = context['servico3']

    # Mapear campos de orçamento (sem acento -> com acento)
    if 'servico_categoria' in context:
        context['serviço_categoria'] = context['servico_categoria']
        if eh_orcamento:
            print(f"[DEBUG ORÇAMENTO] servico_categoria mapeado: {context['servico_categoria']}")
    if 'dependencias_servico' in context:
        context['dependências_serviço'] = context['dependencias_servico']
        if eh_orcamento:
            print(f"[DEBUG ORÇAMENTO] dependencias_servico mapeado: {context['dependencias_servico']}")

    # Debug final do contexto de orçamento
    if eh_orcamento:
        print(f"[DEBUG ORÇAMENTO] Contexto final - serviço_categoria: {context.get('serviço_categoria', 'NÃO ENCONTRADO')}")
        print(f"[DEBUG ORÇAMENTO] Contexto final - dependências_serviço: {context.get('dependências_serviço', 'NÃO ENCONTRADO')}")
        print(f"[DEBUG ORÇAMENTO] Contexto final - preço: {context.get('preço', 'NÃO ENCONTRADO')}")
        print(f"[DEBUG ORÇAMENTO] Contexto final - preco_extenso: {context.get('preco_extenso', 'NÃO ENCONTRADO')}")

    # Formatar valores monetários para sempre ter ,00 (ex: 100 -> 100,00)
    for valor_key in ['valor', 'valor1', 'valor2', 'valor3', 'preco', 'preço']:
        if valor_key in context and context[valor_key]:
            try:
                # Converter para float e formatar com 2 casas decimais
                valor_limpo = str(context[valor_key]).replace('R$', '').replace(',', '.').strip()
                valor_float = float(valor_limpo)
                valor_formatado = f"{valor_float:.2f}".replace('.', ',')
                context[valor_key] = valor_formatado
                # Se for preco, garantir que preço também tenha o valor formatado
                if valor_key == 'preco':
                    context['preço'] = valor_formatado
            except (ValueError, AttributeError):
                pass

    # Formatar numero_meses se presente
    if 'numero_meses' in context and context['numero_meses']:
        context['numero_meses'] = formatar_numero_meses(context['numero_meses'])

    # CÁLCULOS AUTOMÁTICOS baseados nas variáveis extras

    # 1. Calcular Proximo_servico baseado em DATA + numero_meses
    if 'numero_meses' in variaveis_extras and variaveis_extras['numero_meses']:
        try:
            # Extrair número de meses (pode vir como "6", "06", "6 meses", etc.)
            numero_meses_str = str(variaveis_extras['numero_meses'])
            meses = int(''.join(filter(str.isdigit, numero_meses_str)) or '6')

            # Usar DATA fornecida pelo usuário, ou data atual se não fornecida
            if 'DATA' in variaveis_extras and variaveis_extras['DATA']:
                # Parsear data no formato DD/MM/YYYY
                data_str = variaveis_extras['DATA']
                dia, mes_str, ano = data_str.split('/')
                data_base = datetime(int(ano), int(mes_str), int(dia))
            else:
                data_base = datetime.now()

            from dateutil.relativedelta import relativedelta
            data_proxima = data_base + relativedelta(months=meses)
            context['Proximo_servico'] = data_proxima.strftime('%d/%m/%Y')
        except (ValueError, ImportError) as e:
            # Se não conseguir calcular ou não tiver dateutil, deixa vazio
            print(f"[AVISO] Erro ao calcular Proximo_servico: {e}")
            context['Proximo_servico'] = context.get('Proximo_servico', '')

    # 2. Calcular Proximo_servico_caixa (sempre 6 meses)
    if 'Proximo_servico_caixa' not in context or not context['Proximo_servico_caixa']:
        try:
            from dateutil.relativedelta import relativedelta
            data_proxima_caixa = datetime.now() + relativedelta(months=6)
            context['Proximo_servico_caixa'] = data_proxima_caixa.strftime('%d/%m/%Y')
        except ImportError:
            # Cálculo manual se não tiver dateutil
            import calendar
            hoje = datetime.now()
            mes_futuro = (hoje.month + 6 - 1) % 12 + 1
            ano_futuro = hoje.year + (hoje.month + 6 - 1) // 12
            dia_futuro = min(hoje.day, calendar.monthrange(ano_futuro, mes_futuro)[1])
            context['Proximo_servico_caixa'] = f"{dia_futuro:02d}/{mes_futuro:02d}/{ano_futuro}"

    # 3. Calcular valortotal nos recibos (quant * valor)
    if 'quant' in variaveis_extras and 'valor' in variaveis_extras:
        try:
            quant = float(variaveis_extras['quant'].replace(',', '.'))
            valor = float(variaveis_extras['valor'].replace('R$', '').replace(',', '.').strip())
            context['valortotal'] = f"R$ {quant * valor:.2f}".replace('.', ',')
        except (ValueError, AttributeError):
            pass

    # Repetir para quant1, quant2, quant3
    for i in ['1', '2', '3']:
        quant_key = f'quant{i}'
        valor_key = f'valor{i}'
        total_key = f'valortotal{i}'

        if quant_key in variaveis_extras and valor_key in variaveis_extras:
            try:
                quant = float(str(variaveis_extras[quant_key]).replace(',', '.'))
                valor = float(str(variaveis_extras[valor_key]).replace('R$', '').replace(',', '.').strip())
                context[total_key] = f"R$ {quant * valor:.2f}".replace('.', ',')
            except (ValueError, AttributeError):
                pass

    # 4. Calcular valorsomatotal (soma de todos os valortotal)
    if any(key in context for key in ['valortotal', 'valortotal1', 'valortotal2', 'valortotal3']):
        try:
            soma = 0
            for key in ['valortotal', 'valortotal1', 'valortotal2', 'valortotal3']:
                if key in context and context[key]:
                    valor_str = str(context[key]).replace('R$', '').replace(',', '.').strip()
                    soma += float(valor_str)
            context['valorsomatotal'] = f"R$ {soma:.2f}".replace('.', ',')
        except (ValueError, AttributeError):
            pass

    return context


@app.route('/api/cnae/descricao/<cnae_codigo>')
def buscar_descricao_cnae(cnae_codigo):
    """
    Busca a descrição de um CNAE no arquivo cnaes_permitidos.txt

    Exemplo: /api/cnae/descricao/4711302
    Retorna: {"cnae": "47.11-3-02", "descricao": "Comércio varejista..."}
    """
    try:
        # Limpar e formatar CNAE
        cnae_limpo = re.sub(r'[^\d]', '', cnae_codigo)

        # Formatar para XX.XX-X-XX
        if len(cnae_limpo) == 7:
            cnae_formatado = f"{cnae_limpo[0:2]}.{cnae_limpo[2:4]}-{cnae_limpo[4]}-{cnae_limpo[5:7]}"
        else:
            cnae_formatado = cnae_codigo

        # Buscar no arquivo cnaes_permitidos.txt
        if not CNAES_FILE.exists():
            return jsonify({
                'erro': 'Arquivo cnaes_permitidos.txt não encontrado'
            }), 404

        with open(CNAES_FILE, 'r', encoding='utf-8') as f:
            for linha in f:
                linha = linha.strip()

                # Ignorar comentários e linhas vazias
                if not linha or linha.startswith('#'):
                    continue

                # Procurar pela linha que começa com o CNAE
                if linha.startswith(cnae_formatado):
                    # Extrair descrição (após o hífen)
                    partes = linha.split(' - ', 1)
                    if len(partes) == 2:
                        return jsonify({
                            'cnae': partes[0].strip(),
                            'descricao': partes[1].strip(),
                            'descricao_completa': f"{partes[0].strip()} - {partes[1].strip()}"
                        })

        # Se não encontrou
        return jsonify({
            'cnae': cnae_formatado,
            'descricao': 'Descrição não encontrada',
            'descricao_completa': cnae_formatado
        })

    except Exception as e:
        print(f"[ERRO] Buscar descrição CNAE: {e}")
        return jsonify({
            'erro': 'Erro ao buscar descrição do CNAE',
            'detalhes': str(e)
        }), 500


# =============================================================================
# ROTAS DE DOCUMENTOS - MOVIDAS PARA blueprints/documentos.py
# /api/documentos-gerados, /api/documentos/<nome>, /api/documentos/renomear, /api/documentos/mover
# =============================================================================

@app.route('/api/orcamentos')
def api_orcamentos():
    """Lista orçamentos da tabela documentos_gerados"""
    conn = get_db()
    if not tabela_existe('documentos_gerados'):
        return jsonify([])
    orcamentos = conn.execute("""
        SELECT id, tipo_doc, nome_arquivo, caminho_arquivo,
               cliente_nome, razao_social, cliente_cnpj, valor_total,
               data_geracao as data_criacao
        FROM documentos_gerados
        WHERE tipo_doc = 'ORCAMENTO'
        ORDER BY data_geracao DESC
    """).fetchall()
    return jsonify([dict(row) for row in orcamentos])

@app.route('/api/orcamentos/<int:orc_id>/aprovar', methods=['POST'])
def aprovar_orcamento(orc_id):
    """Aprova um orçamento alterando seu status"""
    try:
        conn = get_db()

        # Verificar se a coluna status_aprovacao existe, se não, criar
        cursor = conn.execute("PRAGMA table_info(documentos_gerados)")
        colunas = [col[1] for col in cursor.fetchall()]

        if 'status_aprovacao' not in colunas:
            conn.execute("ALTER TABLE documentos_gerados ADD COLUMN status_aprovacao TEXT DEFAULT 'pendente'")
            conn.commit()

        # Atualizar status do orçamento
        conn.execute("""
            UPDATE documentos_gerados
            SET status_aprovacao = 'aprovado'
            WHERE id = ? AND tipo_doc = 'ORCAMENTO'
        """, (orc_id,))
        conn.commit()

        return jsonify({
            'success': True,
            'message': 'Orçamento aprovado com sucesso',
            'orcamento_id': orc_id
        })

    except Exception as e:
        print(f"Erro ao aprovar orçamento: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/orcamentos/<int:orc_id>/rejeitar', methods=['POST'])
def rejeitar_orcamento(orc_id):
    """Rejeita um orçamento alterando seu status"""
    try:
        conn = get_db()

        # Verificar se a coluna status_aprovacao existe, se não, criar
        cursor = conn.execute("PRAGMA table_info(documentos_gerados)")
        colunas = [col[1] for col in cursor.fetchall()]

        if 'status_aprovacao' not in colunas:
            conn.execute("ALTER TABLE documentos_gerados ADD COLUMN status_aprovacao TEXT DEFAULT 'pendente'")
            conn.commit()

        # Atualizar status do orçamento
        conn.execute("""
            UPDATE documentos_gerados
            SET status_aprovacao = 'rejeitado'
            WHERE id = ? AND tipo_doc = 'ORCAMENTO'
        """, (orc_id,))
        conn.commit()

        return jsonify({
            'success': True,
            'message': 'Orçamento rejeitado com sucesso',
            'orcamento_id': orc_id
        })

    except Exception as e:
        print(f"Erro ao rejeitar orçamento: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@lru_cache(maxsize=1)
def carregar_config_diretorios():
    """Carrega configuração com cache (suporta formato antigo e novo)"""
    try:
        if CONFIG_FILE.exists():
            with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
                config = json.load(f)

            # Se já está no formato novo, retorna direto
            if 'download' in config or 'upload' in config:
                return config

            # Converter formato antigo para novo (retrocompatibilidade)
            return {
                'download': {
                    'laudos': config.get('laudos', ''),
                    'recibos': config.get('recibos', ''),
                    'orcamentos': config.get('orcamentos', '')
                },
                'upload': {
                    'laudos': config.get('laudos', ''),
                    'recibos': config.get('recibos', ''),
                    'orcamentos': config.get('orcamentos', '')
                }
            }

        # Configuração padrão
        return {
            'download': {'laudos': '', 'recibos': '', 'orcamentos': ''},
            'upload': {'laudos': '', 'recibos': '', 'orcamentos': ''}
        }
    except Exception as e:
        print(f"[ERRO] Erro ao carregar config: {e}")
        return {
            'download': {'laudos': '', 'recibos': '', 'orcamentos': ''},
            'upload': {'laudos': '', 'recibos': '', 'orcamentos': ''}
        }

def obter_diretorio_download(tipo_doc):
    """
    Retorna o diretório de download configurado para o tipo de documento
    Se não configurado, retorna OUTPUT_DIR padrão
    tipo_doc: 'LAUDO', 'RECIBO', 'ORCAMENTO'
    """
    try:
        config_file = BASE_DIR / 'config_diretorios_duplos.json'

        # Se não houver configuração nova, usar sistema antigo
        if not config_file.exists():
            config = carregar_config_diretorios()
            download_config = config.get('download', {})

            # Mapear tipo de documento para chave de configuração
            mapa = {
                'LAUDO': 'laudos',
                'RECIBO': 'recibos',
                'ORCAMENTO': 'orcamentos'
            }

            chave = mapa.get(tipo_doc, 'laudos')
            diretorio_config = download_config.get(chave, '').strip()

            if diretorio_config:
                diretorio = Path(diretorio_config)
                diretorio.mkdir(parents=True, exist_ok=True)
                return diretorio

            return OUTPUT_DIR

        # Usar novo sistema de diretórios duplos
        with open(config_file, 'r', encoding='utf-8') as f:
            config = json.load(f)

        dir_principal = config.get('principal')

        if not dir_principal:
            return OUTPUT_DIR

        # Mapear tipo de documento para nome de pasta
        mapa_pastas = {
            'LAUDO': 'Laudos',
            'RECIBO': 'Recibos',
            'ORCAMENTO': 'Orcamentos',
            'BOLETO': 'Boletos'
        }

        pasta = mapa_pastas.get(tipo_doc, 'Laudos')
        diretorio = Path(dir_principal) / pasta
        diretorio.mkdir(parents=True, exist_ok=True)

        return diretorio

    except Exception as e:
        print(f"[ERRO] ao obter diretorio de download: {e}")
        return OUTPUT_DIR

@app.route('/api/config/diretorios', methods=['GET', 'POST'])
def api_config_diretorios():
    """Gerencia configuração de diretórios personalizados"""
    if request.method == 'GET':
        config = carregar_config_diretorios()
        return jsonify(config)

    if request.method == 'POST':
        try:
            dados = request.json
            with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
                json.dump(dados, f, indent=4, ensure_ascii=False)
            # Limpar cache após salvar
            carregar_config_diretorios.cache_clear()
            return jsonify({"message": "Configuração salva!"})
        except Exception as e:
            return jsonify({"error": str(e)}), 500

def listar_arquivos_diretorio(diretorio, tipo, origem, tag_id=None):
    """Helper otimizado para listar arquivos"""
    arquivos = []
    if not diretorio.exists() or not diretorio.is_dir():
        return arquivos

    for arquivo in diretorio.iterdir():
        if arquivo.is_file():
            stat = arquivo.stat()
            arquivo_info = {
                'nome': arquivo.name,
                'caminho': str(arquivo),
                'tamanho': stat.st_size,
                'data_modificacao': datetime.fromtimestamp(stat.st_mtime).isoformat(),
                'extensao': arquivo.suffix,
                'tipo': tipo,
                'origem': origem
            }

            # Adicionar tag_id se configurado para aplicação automática no frontend
            if tag_id:
                arquivo_info['tag_automatica_id'] = tag_id

            arquivos.append(arquivo_info)
    return arquivos

@app.route('/api/arquivos')
def api_arquivos():
    """Lista arquivos de todos os diretórios configurados com paginação (otimizado)"""
    try:
        # Parâmetros de paginação
        page = request.args.get('page', 1, type=int)
        per_page = request.args.get('per_page', 50, type=int)

        # Validar parâmetros
        page = max(1, page)  # Mínimo página 1
        per_page = min(max(1, per_page), 200)  # Entre 1 e 200 itens por página

        # Listar arquivos do output padrão e da pasta documentos
        arquivos = []

        # Listar da pasta documentos (onde ficam os documentos gerados)
        documentos_dir = OUTPUT_DIR / 'documentos'
        if documentos_dir.exists():
            arquivos.extend(listar_arquivos_diretorio(documentos_dir, 'documentos', 'Sistema'))

        # Também listar diretamente do OUTPUT_DIR (para compatibilidade)
        arquivos.extend(listar_arquivos_diretorio(OUTPUT_DIR, 'output', 'Sistema'))

        # Diretórios personalizados (novo formato: download/upload)
        config = carregar_config_diretorios()
        download_config = config.get('download', {})
        upload_config = config.get('upload', {})
        tags_config = config.get('tags', {})

        # Adicionar arquivos dos diretórios de DOWNLOAD (onde salvamos)
        diretorios_download = [
            (download_config.get('laudos'), 'download_laudo', 'Laudos', None),
            (download_config.get('recibos'), 'download_recibo', 'Recibos', None),
            (download_config.get('orcamentos'), 'download_orcamento', 'Orçamentos', None)
        ]

        # Adicionar arquivos dos diretórios de UPLOAD (onde lemos) COM TAGS AUTOMÁTICAS
        diretorios_upload = [
            (upload_config.get('laudos'), 'upload_laudo', 'Laudos (Upload)', tags_config.get('laudos')),
            (upload_config.get('recibos'), 'upload_recibo', 'Recibos (Upload)', tags_config.get('recibos')),
            (upload_config.get('orcamentos'), 'upload_orcamento', 'Orçamentos (Upload)', tags_config.get('orcamentos'))
        ]

        for caminho, tipo, origem, tag_id in diretorios_download + diretorios_upload:
            if caminho and caminho.strip():
                try:
                    # Converter tag_id para int se for string numérica, senão None
                    tag_id_int = None
                    if tag_id:
                        try:
                            tag_id_int = int(tag_id) if tag_id else None
                        except (ValueError, TypeError):
                            pass  # tag_id não é numérico
                    arquivos.extend(listar_arquivos_diretorio(Path(caminho), tipo, origem, tag_id_int))
                except (OSError, PermissionError):
                    pass  # Ignora diretórios inacessíveis

        # Ordenar por data de modificação (mais recente primeiro)
        arquivos.sort(key=lambda x: x['data_modificacao'], reverse=True)

        # Aplicar paginação
        total = len(arquivos)
        start_idx = (page - 1) * per_page
        end_idx = start_idx + per_page
        arquivos_paginados = arquivos[start_idx:end_idx]

        # Calcular total de páginas
        total_pages = (total + per_page - 1) // per_page  # Arredonda para cima

        return jsonify({
            'arquivos': arquivos_paginados,
            'pagination': {
                'page': page,
                'per_page': per_page,
                'total': total,
                'total_pages': total_pages,
                'has_next': page < total_pages,
                'has_prev': page > 1
            }
        })

    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

@app.route('/api/sincronizar-arquivos', methods=['POST'])
def api_sincronizar_arquivos():
    """Sincroniza todos os arquivos das pastas configuradas com o banco de dados"""
    try:
        stats = {
            'total_encontrados': 0,
            'ja_existentes': 0,
            'adicionados': 0,
            'erros': 0,
            'detalhes': []
        }

        # Carregar configuração de diretórios
        config = carregar_config_diretorios()
        download_config = config.get('download', {})

        # Lista de diretórios para sincronizar
        diretorios_sync = [
            (download_config.get('laudos'), 'LAUDO'),
            (download_config.get('recibos'), 'RECIBO'),
            (download_config.get('orcamentos'), 'ORCAMENTO')
        ]

        conn = get_db()

        for caminho_dir, tipo_doc in diretorios_sync:
            if not caminho_dir or not caminho_dir.strip():
                continue

            dir_path = Path(caminho_dir)
            if not dir_path.exists() or not dir_path.is_dir():
                continue

            # Listar todos os arquivos do diretório
            for arquivo in dir_path.glob('**/*'):
                if not arquivo.is_file():
                    continue

                # Filtrar apenas arquivos suportados
                extensoes_suportadas = ['.pdf', '.docx', '.doc', '.xml']
                if arquivo.suffix.lower() not in extensoes_suportadas:
                    continue

                stats['total_encontrados'] += 1
                nome_arquivo = arquivo.name
                caminho_completo = str(arquivo.absolute())

                # Verificar se já existe no banco
                existente = conn.execute(
                    'SELECT id FROM documentos_gerados WHERE nome_arquivo = ?',
                    (nome_arquivo,)
                ).fetchone()

                if existente:
                    stats['ja_existentes'] += 1
                    continue

                # Processar novo arquivo
                try:
                    # Extrair dados baseado na extensão
                    dados_extraidos = {}

                    if arquivo.suffix.lower() == '.pdf':
                        dados_extraidos = extrair_dados_pdf(str(arquivo))
                    elif arquivo.suffix.lower() in ['.docx', '.doc']:
                        # Para arquivos Word, apenas dados básicos
                        dados_extraidos = {
                            'tipo_documento': tipo_doc,
                            'cliente_nome': '',
                            'cliente_cnpj': '',
                            'valor_total': 0.0
                        }
                    elif arquivo.suffix.lower() == '.xml':
                        dados_extraidos = extrair_dados_xml(str(arquivo))

                    # Inserir no banco
                    cursor = conn.execute('''INSERT INTO documentos_gerados
                        (tipo_doc, nome_arquivo, caminho_arquivo, cliente_nome, razao_social, cliente_cnpj, valor_total, data_geracao)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?)''',
                        (
                            tipo_doc,
                            nome_arquivo,
                            caminho_completo,
                            dados_extraidos.get('cliente_nome', ''),
                            dados_extraidos.get('razao_social', ''),
                            dados_extraidos.get('cliente_cnpj', ''),
                            dados_extraidos.get('valor_total', 0.0),
                            datetime.now()
                        ))

                    documento_id = cursor.lastrowid
                    conn.commit()

                    stats['adicionados'] += 1
                    stats['detalhes'].append({
                        'arquivo': nome_arquivo,
                        'tipo': tipo_doc,
                        'status': 'adicionado',
                        'id': documento_id
                    })

                except Exception as e:
                    stats['erros'] += 1
                    stats['detalhes'].append({
                        'arquivo': nome_arquivo,
                        'tipo': tipo_doc,
                        'status': 'erro',
                        'mensagem': str(e)
                    })
                    continue

        return jsonify({
            'success': True,
            'message': f'Sincronização concluída: {stats["adicionados"]} arquivos adicionados',
            'stats': stats
        })

    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

@app.route('/api/arquivo/excluir', methods=['POST'])
def api_excluir_arquivo():
    """Exclui um arquivo do sistema"""
    try:
        dados = request.json
        caminho_arquivo = dados.get('caminho')

        if not caminho_arquivo:
            return jsonify({"error": "Caminho do arquivo não fornecido"}), 400

        arquivo_path = Path(caminho_arquivo)

        # Verificar se o arquivo existe
        if not arquivo_path.exists():
            return jsonify({"error": "Arquivo não encontrado"}), 404

        # Verificar se é um arquivo (não diretório)
        if not arquivo_path.is_file():
            return jsonify({"error": "O caminho não é um arquivo válido"}), 400

        # Deletar o arquivo
        arquivo_path.unlink()

        return jsonify({
            "message": "Arquivo excluído com sucesso!",
            "arquivo": arquivo_path.name
        })

    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": f"Erro ao excluir arquivo: {str(e)}"}), 500

def encontrar_arquivo_em_diretorios(filename):
    """Procura arquivo em todos os diretórios configurados"""
    # Primeiro tenta no OUTPUT_DIR padrão
    caminho = OUTPUT_DIR / filename
    if caminho.exists():
        return caminho

    # Tenta em output/documentos (geração inteligente)
    caminho_docs = OUTPUT_DIR / 'documentos' / filename
    if caminho_docs.exists():
        return caminho_docs

    # Buscar em todos os diretórios configurados
    config = carregar_config_diretorios()
    download_config = config.get('download', {})
    upload_config = config.get('upload', {})

    todos_diretorios = []
    for secao in [download_config, upload_config]:
        for chave in ['laudos', 'recibos', 'orcamentos']:
            dir_config = secao.get(chave, '').strip()
            if dir_config:
                todos_diretorios.append(Path(dir_config))

    # Procurar em cada diretório
    for diretorio in todos_diretorios:
        if diretorio.exists():
            caminho = diretorio / filename
            if caminho.exists():
                return caminho

    return None

@app.route('/download/<path:filename>')
@app.route('/api/download/<path:filename>')
def download_arquivo(filename):
    """Permite download de arquivo de qualquer diretório configurado"""
    try:
        caminho_arquivo = encontrar_arquivo_em_diretorios(filename)

        if not caminho_arquivo:
            return jsonify({"error": "Arquivo não encontrado"}), 404

        return send_file(str(caminho_arquivo), as_attachment=True, download_name=filename)

    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/visualizar/<path:filename>')
def visualizar_arquivo(filename):
    """Permite visualizar arquivo no navegador (sem forçar download)"""
    try:
        caminho_arquivo = encontrar_arquivo_em_diretorios(filename)

        if not caminho_arquivo:
            return jsonify({"error": "Arquivo não encontrado"}), 404

        # Enviar arquivo sem forçar download (as_attachment=False)
        return send_file(str(caminho_arquivo), as_attachment=False)

    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/visualizar-documento/<path:filename>')
def visualizar_documento_html(filename):
    """Converte documento DOCX para PDF e exibe no navegador (com cache)"""
    try:
        import time
        import os

        caminho_arquivo = encontrar_arquivo_em_diretorios(filename)

        if not caminho_arquivo:
            return jsonify({"error": "Arquivo não encontrado"}), 404

        # Verificar se é um arquivo DOCX
        if not filename.lower().endswith('.docx'):
            return jsonify({"error": "Apenas arquivos DOCX podem ser visualizados"}), 400

        # Criar diretório de cache para PDFs
        cache_dir = OUTPUT_DIR / 'cache_pdf'
        cache_dir.mkdir(exist_ok=True)

        pdf_filename = filename.replace('.docx', '.pdf')
        pdf_path = cache_dir / pdf_filename

        # Verificar se o PDF já existe no cache E está atualizado
        converter_novamente = True
        if pdf_path.exists():
            # Verificar se o DOCX foi modificado após o PDF ser criado
            docx_mtime = os.path.getmtime(str(caminho_arquivo))
            pdf_mtime = os.path.getmtime(str(pdf_path))
            if pdf_mtime >= docx_mtime:
                print(f"[CACHE] Usando PDF em cache: {pdf_filename}")
                converter_novamente = False

        # Converter DOCX para PDF se necessário
        if converter_novamente:
            try:
                # Tentar importar docx2pdf
                from docx2pdf import convert

                print(f"[INFO] Convertendo {filename} para PDF...")
                start = time.time()
                convert(str(caminho_arquivo), str(pdf_path))

                # Aguardar conversão finalizar (com timeout reduzido)
                timeout = 15
                start_time = time.time()
                while not pdf_path.exists() and (time.time() - start_time) < timeout:
                    time.sleep(0.2)

                if not pdf_path.exists():
                    raise Exception("PDF não foi criado após conversão")

                elapsed = time.time() - start
                print(f"[OK] PDF criado em {elapsed:.2f}s: {pdf_filename}")

            except ImportError:
                print(f"[AVISO] docx2pdf não instalado. Fazendo download direto do DOCX.")
                # Se docx2pdf não estiver instalado, fazer download ao invés de visualizar
                return send_file(str(caminho_arquivo), as_attachment=True, download_name=filename)

            except Exception as conv_error:
                print(f"[ERRO] Ao converter DOCX para PDF: {conv_error}")
                traceback.print_exc()

                # Fallback: fazer download do DOCX ao invés de tentar visualizar
                print(f"[FALLBACK] Iniciando download do arquivo DOCX original")
                return send_file(str(caminho_arquivo), as_attachment=True, download_name=filename)

        # Enviar o PDF para visualização no navegador
        return send_file(
            str(pdf_path),
            mimetype='application/pdf',
            as_attachment=False,
            download_name=pdf_filename
        )

    except Exception as e:
        print(f"[ERRO] Ao visualizar documento: {e}")
        traceback.print_exc()
        return jsonify({"error": f"Erro ao visualizar documento: {str(e)}"}), 500



@app.route('/api/compartilhar/<path:filename>')
def compartilhar_arquivo(filename):
    """Gera link de compartilhamento"""
    try:
        caminho_arquivo = OUTPUT_DIR / filename

        if not caminho_arquivo.exists():
            return jsonify({"error": "Arquivo não encontrado"}), 404

        # Gerar link simples (em produção, use tokens temporários)
        link = f"http://localhost:5000/api/download/{filename}"

        return jsonify({
            "link": link,
            "nome": filename,
            "expira": "Link permanente"
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500

# /api/documentos - MOVIDA PARA blueprints/documentos.py

@app.route('/api/upload', methods=['POST'])
def upload_arquivo():
    try:
        if 'files[]' not in request.files:
            return jsonify({"error": "Nenhum arquivo enviado"}), 400

        files = request.files.getlist('files[]')
        tipo_documento = request.form.get('tipo_documento', 'outro')  # Pegar tipo do formulário
        respeitar_limite = request.form.get('respeitar_limite', 'false') == 'true'
        uploaded = []

        print(f"\n[UPLOAD] Tipo de documento selecionado: {tipo_documento}")
        print(f"[UPLOAD] Total de arquivos: {len(files)}")
        print(f"[UPLOAD] Respeitar limite API: {respeitar_limite}")

        for file in files:
            if file.filename == '':
                continue

            # Sanitizar nome do arquivo
            filename = file.filename.replace('/', '-').replace('\\', '-')
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            nome_final = f"{timestamp}_{filename}"

            # Mapear tipo_documento para tipo_doc
            tipo_doc_map = {
                'boleto': 'BOLETO',
                'laudo': 'LAUDO',
                'recibo': 'RECIBO',
                'orcamento': 'ORCAMENTO',
                'outro': 'UPLOAD'
            }
            tipo_doc = tipo_doc_map.get(tipo_documento, 'UPLOAD')

            # Salvar temporariamente
            temp_path = OUTPUT_DIR / 'temp' / nome_final
            temp_path.parent.mkdir(parents=True, exist_ok=True)
            file.save(str(temp_path))
            print(f"[UPLOAD] Arquivo temporário salvo: {temp_path}")
            print(f"[UPLOAD] Tipo de documento: {tipo_doc}")

            # Salvar no diretório principal configurado
            if tipo_doc != 'UPLOAD':
                print(f"[UPLOAD] Chamando salvar_upload_simples...")
                caminho = salvar_upload_simples(str(temp_path), nome_final, tipo_doc)
                if caminho:
                    print(f"[UPLOAD] Arquivo salvo em: {caminho}")
                else:
                    print(f"[UPLOAD] ERRO: salvar_upload_simples retornou None")
                    # Remover temp e continuar com próximo arquivo
                    if temp_path.exists():
                        temp_path.unlink(missing_ok=True)
                    continue
            else:
                # Para 'outro', salvar em OUTPUT_DIR
                caminho = OUTPUT_DIR / nome_final
                shutil.copy2(str(temp_path), str(caminho))
                caminho = str(caminho)
                print(f"[UPLOAD] Arquivo 'outro' salvo em: {caminho}")

            # Remover arquivo temporário
            if temp_path.exists():
                temp_path.unlink(missing_ok=True)
                print(f"[UPLOAD] Arquivo temporário removido")

            # Registrar no banco
            conn = get_db()
            cursor = conn.execute('''INSERT INTO documentos_gerados
                (tipo_doc, nome_arquivo, caminho_arquivo, data_geracao)
                VALUES (?, ?, ?, ?)''',
                (tipo_doc, filename, str(caminho), datetime.now()))
            documento_id = cursor.lastrowid
            conn.commit()

            print(f"[UPLOAD] Arquivo salvo: {filename} | ID: {documento_id} | Tipo: {tipo_doc}")

            # Aplicar tags automaticamente
            tags_sugeridas = sugerir_tags_automaticas(documento_id, 'documento', filename, '')
            for tag_sugerida in tags_sugeridas:
                adicionar_tag_documento(
                    documento_id,
                    'documento',
                    tag_sugerida['tag_id'],
                    tag_sugerida['confianca'],
                    manual=False
                )

            # PROCESSAMENTO ESPECÍFICO PARA BOLETOS (função dedicada)
            if tipo_documento == 'boleto':
                print(f"[UPLOAD] Iniciando processamento de BOLETO...")
                try:
                    dados_boleto = processar_boleto(str(caminho), filename, documento_id)
                    if dados_boleto:
                        print(f"[UPLOAD] OKOK Boleto processado com sucesso!")
                    else:
                        print(f"[UPLOAD] AVISO Falha ao processar boleto (não é um boleto válido)")
                except Exception as e:
                    print(f"[UPLOAD] ERRO Erro ao processar boleto: {e}")
                    traceback.print_exc()

            uploaded.append({
                'nome': filename,
                'caminho': str(caminho),
                'tamanho': Path(caminho).stat().st_size
            })

        return jsonify({
            "message": f"{len(uploaded)} arquivo(s) enviado(s)",
            "files": uploaded
        })

    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

def extrair_valor_pdf(caminho_arquivo):
    """Tenta extrair valor de PDF"""
    dados = extrair_dados_pdf(caminho_arquivo)
    return dados['valor_total']

@app.route('/api/empresas')
def api_empresas():
    if not tabela_existe('empresas'):
        return jsonify([])

    nome = request.args.get('nome', '')
    cnpj = request.args.get('cnpj', '')

    query = """SELECT
        cnpj_completo as cnpj,
        nome_fantasia,
        nome_fantasia as razao_social,
        cnae_fiscal_principal,
        tipo_logradouro,
        logradouro,
        numero,
        bairro,
        cep,
        uf,
        codigo_municipio as municipio
        FROM empresas
        WHERE situacao_cadastral = '02'"""
    params = []

    if nome:
        query += " AND nome_fantasia LIKE ?"
        params.append(f'%{nome}%')

    if cnpj:
        query += " AND cnpj_completo LIKE ?"
        params.append(f'%{cnpj}%')

    query += " ORDER BY nome_fantasia LIMIT 50"

    empresas = get_db().execute(query, params).fetchall()

    # Processar cada empresa para formatar CNAE e verificar se está permitido
    resultado = []
    for row in empresas:
        emp = dict(row)
        cnae_original = emp.get('cnae_fiscal_principal', '')

        # Formatar CNAE e adicionar descrição
        emp['cnae_fiscal'] = formatar_cnae(cnae_original)
        emp['cnae_permitido'] = cnae_esta_permitido(cnae_original)

        # Converter código de município para nome
        codigo_mun = emp.get('municipio', '')
        emp['municipio'] = obter_nome_municipio(codigo_mun)

        # Remover campo temporário
        if 'cnae_fiscal_principal' in emp:
            del emp['cnae_fiscal_principal']

        resultado.append(emp)

    return jsonify(resultado)

# =============================================================================
# ROTAS DE TAGS - MOVIDAS PARA blueprints/tags.py
# ROTAS DE BOLETOS - MOVIDAS PARA blueprints/boletos.py
# ROTAS DE GARANTIAS - MOVIDAS PARA blueprints/clientes.py
# =============================================================================

def formatar_numero_meses(valor) -> str:
    """
    Formata o número de meses para o formato correto: "X mês" ou "X meses".

    Args:
        valor: Pode ser string ("6", "06", "6 Meses", "06 Meses") ou int (6)

    Returns:
        str: Formato padronizado "1 mês", "2 meses", "6 meses", etc.

    Examples:
        >>> formatar_numero_meses("6")
        '6 meses'
        >>> formatar_numero_meses("06 Meses")
        '6 meses'
        >>> formatar_numero_meses("1")
        '1 mês'
        >>> formatar_numero_meses(12)
        '12 meses'
    """
    if not valor:
        return '6 meses'  # Valor padrão

    try:
        # Extrair apenas o número da string
        if isinstance(valor, str):
            # Remove tudo que não é dígito
            numero_str = ''.join(filter(str.isdigit, valor))
            if not numero_str:
                return '6 meses'
            numero = int(numero_str)
        else:
            numero = int(valor)

        # Formatar com singular/plural correto
        if numero == 1:
            return '1 mês'
        else:
            return f'{numero} meses'

    except (ValueError, TypeError):
        return '6 meses'  # Fallback seguro

def preparar_contexto_nomes(nome_fantasia: str, razao_social: str, tipo_documento: str) -> dict:
    """
    Prepara campos de nome e razão social baseado no tipo de documento.

    Regras:
    - Laudos: nome_fantasia em formato normal
    - Recibos/Orçamentos: nome_fantasia em CAIXA ALTA
    - razao_social: sempre sem fallback (se vazio, mantém vazio)

    Args:
        nome_fantasia: Nome fantasia do cliente (pode ser None/vazio)
        razao_social: Razão social do cliente (pode ser None/vazio)
        tipo_documento: 'laudo', 'recibo' ou 'orcamento'

    Returns:
        dict: {'nome_fantasia': str, 'nome_razao_social': str}
    """
    # Normalizar inputs (None -> string vazia)
    nome_fantasia = nome_fantasia or ''
    razao_social = razao_social or ''

    if tipo_documento == 'laudo':
        # Laudos: nome fantasia em formato normal
        return {
            'nome_fantasia': nome_fantasia if nome_fantasia else 'Cliente',
            'nome_razao_social': razao_social
        }
    elif tipo_documento in ['recibo', 'orcamento']:
        # Recibos/Orçamentos: nome fantasia em CAIXA ALTA
        return {
            'nome_fantasia': nome_fantasia.upper() if nome_fantasia else 'Cliente',
            'nome_razao_social': razao_social
        }
    else:
        # Padrão
        return {
            'nome_fantasia': nome_fantasia if nome_fantasia else 'Cliente',
            'nome_razao_social': razao_social
        }

@app.route('/gerar-laudo', methods=['POST'])
def gerar_laudo():
    try:
        dados = request.json
        print(f"[DEBUG LAUDO] Data recebida: '{dados.get('data')}'")
        data_obj, data_str = formatar_data(dados.get('data'))
        print(f"[DEBUG LAUDO] Data formatada: '{data_str}'")

        tipo = dados.get('tipo_laudo', 'padrao')
        template_path = TEMPLATES.get(f'laudo_{tipo}', TEMPLATES['laudo_padrao'])

        if not template_path.exists():
            return jsonify({"error": "Template nao encontrado"}), 404

        # Preparar campos de nome usando função helper
        nomes = preparar_contexto_nomes(
            dados.get('nome'),
            dados.get('razao_social'),
            'laudo'
        )

        contexto = {
            **nomes,  # Expande nome_fantasia e nome_razao_social
            'endereco_cliente': dados.get('endereco', ''),
            'cnpj_cliente': formatar_cnpj(dados.get('cnpj', '')),
            'cnae_cliente': formatar_cnae(dados.get('cnae', '')),
            'DATA': data_str,
        }
        print(f"[DEBUG LAUDO] Contexto DATA: '{contexto['DATA']}'")

        pragas = dados.get('pragas', {})
        mapa = {k: " " for k in ['B', 'R', 'C', 'P', 'V', 'M', 'F', 'O', 'W']}
        check_map = {
            'barata': 'B', 'rato': 'R', 'camundongo': 'C', 'cupim': 'W',
            'pulga': 'P', 'escorpiao': 'V', 'mosca': 'M', 'formiga': 'F', 'outros': 'O'
        }
        for praga, code in check_map.items():
            if pragas.get(praga):
                mapa[code] = "X"
        contexto.update(mapa)

        garantia = dados.get('garantia', '6')
        garantia_formatada = formatar_numero_meses(garantia)
        try:
            # Extrair número de meses da garantia formatada
            meses = int(garantia_formatada.split()[0])
        except (ValueError, IndexError, AttributeError):
            meses = 6
        contexto['Proximo_servico'] = (data_obj + relativedelta(months=meses)).strftime('%d/%m/%Y')
        contexto['numero_meses'] = garantia_formatada

        # Proximo_servico_caixa sempre 6 meses após a data
        contexto['Proximo_servico_caixa'] = (data_obj + relativedelta(months=6)).strftime('%d/%m/%Y')

        doc = DocxTemplate(str(template_path))
        doc.render(contexto)

        nome_cliente = dados.get('nome', 'Cliente')
        nome_arquivo = gerar_nome_arquivo(nome_cliente, data_obj, 'LAUDO')

        # Salvar temporariamente
        temp_path = OUTPUT_DIR / 'temp' / nome_arquivo
        temp_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(temp_path))
        print(f"[DEBUG] Arquivo temporário salvo em: {temp_path}")

        # Salvar nos diretórios configurados (principal + backup)
        print(f"[DEBUG] Chamando salvar_arquivo_duplo para LAUDO...")
        resultado = salvar_arquivo_duplo(str(temp_path), nome_arquivo, 'LAUDO')
        caminho_principal = resultado['caminho_principal']
        print(f"[DEBUG] Arquivo salvo em: {caminho_principal}")
        print(f"[DEBUG] Todos os caminhos: {resultado['caminhos_todos']}")

        # Remover arquivo temporário
        temp_path.unlink(missing_ok=True)
        print(f"[DEBUG] Arquivo temporário removido")

        conn = get_db()
        cursor = conn.execute('''INSERT INTO documentos_gerados
            (tipo_doc, nome_arquivo, caminho_arquivo, cliente_nome, razao_social, cliente_cnpj, valor_total, data_geracao)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)''',
            ('LAUDO', nome_arquivo, caminho_principal,
             dados.get('nome', ''),
             dados.get('razao_social', ''),
             dados.get('cnpj', ''),
             0.0,  # Laudos não têm valor
             datetime.now()))
        documento_id = cursor.lastrowid
        conn.commit()

        # Aplicar tag de Laudo automaticamente
        tags_sugeridas = sugerir_tags_automaticas(documento_id, 'documento', nome_arquivo, '')
        for tag_sugerida in tags_sugeridas:
            adicionar_tag_documento(
                documento_id,
                'documento',
                tag_sugerida['tag_id'],
                tag_sugerida['confianca'],
                manual=False
            )

        return send_file(str(caminho_principal), as_attachment=True, download_name=nome_arquivo)

    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

@app.route('/gerar-recibo', methods=['POST'])
def gerar_recibo():
    try:
        dados = request.json
        print(f"[DEBUG RECIBO] Data recebida: '{dados.get('data')}'")
        data_obj, data_str = formatar_data(dados.get('data'))
        print(f"[DEBUG RECIBO] Data formatada: '{data_str}'")

        template_path = TEMPLATES['recibo']
        if not template_path.exists():
            return jsonify({"error": "Template de recibo nao encontrado"}), 404

        # Linha 0
        quant = float(dados.get('quant', 0) or 0)
        valor = float(dados.get('valor', 0) or 0)
        valortotal = quant * valor

        # Linha 1
        quant1 = float(dados.get('quant1', 0) or 0)
        valor1 = float(dados.get('valor1', 0) or 0)
        valortotal1 = quant1 * valor1

        # Linha 2
        quant2 = float(dados.get('quant2', 0) or 0)
        valor2 = float(dados.get('valor2', 0) or 0)
        valortotal2 = quant2 * valor2

        # Linha 3
        quant3 = float(dados.get('quant3', 0) or 0)
        valor3 = float(dados.get('valor3', 0) or 0)
        valortotal3 = quant3 * valor3

        # Soma total
        valorsomatotal = valortotal + valortotal1 + valortotal2 + valortotal3

        # Preparar campos de nome usando função helper
        nomes = preparar_contexto_nomes(
            dados.get('nome'),
            dados.get('razao_social'),
            'recibo'
        )

        contexto = {
            'cliente_nome': dados.get('nome', 'Cliente'),
            **nomes,  # Expande nome_fantasia e nome_razao_social
            'endereco_cliente': dados.get('endereco', ''),
            'cnpj_cliente': formatar_cnpj(dados.get('cnpj', '')),
            'DATA': data_str,  # Corrigido: template usa {{DATA}} em maiúsculo
            'numero_meses': formatar_numero_meses(dados.get('numero_meses', '6')),

            # Linha 0
            'quant': str(int(quant)) if quant else '',
            'serviço': dados.get('servico', ''),
            'valor': f"{valor:.2f}".replace('.', ','),
            'valortotal': f"{valortotal:.2f}".replace('.', ','),

            # Linha 1
            'quant1': str(int(quant1)) if quant1 else '',
            'serviço1': dados.get('servico1', ''),
            'valor1': f"{valor1:.2f}".replace('.', ','),
            'valortotal1': f"{valortotal1:.2f}".replace('.', ','),

            # Linha 2
            'quant2': str(int(quant2)) if quant2 else '',
            'serviço2': dados.get('servico2', ''),
            'valor2': f"{valor2:.2f}".replace('.', ','),
            'valortotal2': f"{valortotal2:.2f}".replace('.', ','),

            # Linha 3
            'quant3': str(int(quant3)) if quant3 else '',
            'serviço3': dados.get('servico3', ''),
            'valor3': f"{valor3:.2f}".replace('.', ','),
            'valortotal3': f"{valortotal3:.2f}".replace('.', ','),

            # Soma total
            'valorsomatotal': f"{valorsomatotal:.2f}".replace('.', ',') if valorsomatotal else '0,00',
        }

        doc = DocxTemplate(str(template_path))
        doc.render(contexto)

        nome_cliente = dados.get('nome', 'Cliente')
        nome_arquivo = gerar_nome_arquivo(nome_cliente, data_obj, 'RECIBO')

        # Salvar temporariamente
        temp_path = OUTPUT_DIR / 'temp' / nome_arquivo
        temp_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(temp_path))

        # Salvar nos diretórios configurados (principal + backup)
        resultado = salvar_arquivo_duplo(str(temp_path), nome_arquivo, 'RECIBO')
        caminho_principal = resultado['caminho_principal']

        # Remover arquivo temporário
        temp_path.unlink(missing_ok=True)

        conn = get_db()
        cursor = conn.execute('''INSERT INTO documentos_gerados
            (tipo_doc, nome_arquivo, caminho_arquivo, cliente_nome, razao_social, cliente_cnpj, valor_total, data_geracao)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)''',
            ('RECIBO', nome_arquivo, caminho_principal,
             dados.get('nome', ''),
             dados.get('razao_social', ''),
             dados.get('cnpj', ''),
             valorsomatotal,
             datetime.now()))
        documento_id = cursor.lastrowid

        # Também salvar na tabela recibos
        descricao_servicos = []
        if dados.get('servico'):
            descricao_servicos.append(dados.get('servico'))
        if dados.get('servico1'):
            descricao_servicos.append(dados.get('servico1'))
        if dados.get('servico3'):
            descricao_servicos.append(dados.get('servico3'))
        if dados.get('servico4'):
            descricao_servicos.append(dados.get('servico4'))

        descricao = ', '.join(descricao_servicos) if descricao_servicos else 'Serviços de dedetização'

        conn.execute('''INSERT INTO recibos
            (cliente_nome, numero_recibo, descricao, valor_total, data_emissao, arquivo_caminho)
            VALUES (?, ?, ?, ?, ?, ?)''',
            (dados.get('nome', ''),
             nome_arquivo,
             descricao,
             valorsomatotal,
             data_obj.strftime('%Y-%m-%d'),
             caminho_principal))

        conn.commit()

        # Aplicar tag de Recibo automaticamente
        tags_sugeridas = sugerir_tags_automaticas(documento_id, 'documento', nome_arquivo, '')
        for tag_sugerida in tags_sugeridas:
            adicionar_tag_documento(
                documento_id,
                'documento',
                tag_sugerida['tag_id'],
                tag_sugerida['confianca'],
                manual=False
            )

        return send_file(str(caminho_principal), as_attachment=True, download_name=nome_arquivo)

    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

@app.route('/gerar-orcamento', methods=['POST'])
def gerar_orcamento():
    try:
        dados = request.json

        # Para orçamentos, a data já vem formatada por extenso do frontend
        # Ex: "09 de janeiro de 2026"
        data_extenso = dados.get('DATA', '')
        print(f"[DEBUG ORCAMENTO] Data por extenso recebida: '{data_extenso}'")

        template_path = TEMPLATES['orcamento']
        if not template_path.exists():
            return jsonify({"error": "Template de orcamento nao encontrado"}), 404

        # Preparar campos de nome usando função helper
        nomes = preparar_contexto_nomes(
            dados.get('nome'),
            dados.get('razao_social'),
            'orcamento'
        )

        # Valores vêm direto do formulário (sem conversão automática)
        valor_formatado = dados.get('valor_formatado', '')
        valor_escrito = dados.get('escrito', '')

        print(f"[DEBUG ORÇAMENTO] Valor formatado recebido: {valor_formatado}")
        print(f"[DEBUG ORÇAMENTO] Valor escrito recebido: {valor_escrito}")

        contexto = {
            'cliente_nome': dados.get('nome', 'Cliente'),
            **nomes,  # Expande nome_fantasia e nome_razao_social
            'endereco_cliente': dados.get('endereco', ''),
            'cliente_endereco': dados.get('endereco', ''),  # Manter compatibilidade
            'cliente_cnpj': formatar_cnpj(dados.get('cnpj', '')),
            'cnpj_cliente': formatar_cnpj(dados.get('cnpj', '')),  # Manter compatibilidade
            'servicos': dados.get('servicos', 'Servicos de dedetizacao'),
            'valor_total': dados.get('valor_total', '0,00'),
            'DATA': data_extenso,  # Data formatada por extenso vem do frontend (ex: 09 de janeiro de 2026)
            'validade': dados.get('validade', '30 dias'),
            'servico_categoria': dados.get('servico_categoria', ''),
            'dependencias_servico': dados.get('dependencias_servico', ''),
            # Valores diretos dos campos do formulário
            'valor_formatado': valor_formatado,  # Ex: R$ 100,00
            'escrito': valor_escrito,  # Ex: cem reais
        }

        doc = DocxTemplate(str(template_path))
        doc.render(contexto)

        nome_cliente = dados.get('nome', 'Cliente')
        nome_arquivo = gerar_nome_arquivo(nome_cliente, data_obj, 'ORCAMENTO')

        # Salvar temporariamente
        temp_path = OUTPUT_DIR / 'temp' / nome_arquivo
        temp_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(temp_path))

        # Salvar nos diretórios configurados (principal + backup)
        resultado = salvar_arquivo_duplo(str(temp_path), nome_arquivo, 'ORCAMENTO')
        caminho_principal = resultado['caminho_principal']

        # Remover arquivo temporário
        temp_path.unlink(missing_ok=True)

        # Converter preço para float (pode vir como string com vírgula)
        preco_str = dados.get('preco', '0,00')
        try:
            valor_orcamento = float(preco_str.replace('.', '').replace(',', '.'))
        except (ValueError, AttributeError):
            valor_orcamento = 0.0

        conn = get_db()
        cursor = conn.execute('''INSERT INTO documentos_gerados
            (tipo_doc, nome_arquivo, caminho_arquivo, cliente_nome, razao_social, cliente_cnpj, valor_total, data_geracao)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)''',
            ('ORCAMENTO', nome_arquivo, caminho_principal,
             dados.get('nome', ''),
             dados.get('razao_social', ''),
             dados.get('cnpj', ''),
             valor_orcamento,
             datetime.now()))
        documento_id = cursor.lastrowid
        conn.commit()

        # Aplicar tag de Orçamento automaticamente
        tags_sugeridas = sugerir_tags_automaticas(documento_id, 'documento', nome_arquivo, '')
        for tag_sugerida in tags_sugeridas:
            adicionar_tag_documento(
                documento_id,
                'documento',
                tag_sugerida['tag_id'],
                tag_sugerida['confianca'],
                manual=False
            )

        return send_file(str(caminho_principal), as_attachment=True, download_name=nome_arquivo)

    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

@app.route('/api/analise-financeira', methods=['GET'])
def analise_financeira():
    """Análise focada em Laudos, Orçamentos e Recibos"""
    try:
        data_inicio = request.args.get('data_inicio')
        data_fim = request.args.get('data_fim')

        conn = get_db()

        # Contar documentos gerados por tipo (LAUDO, RECIBO, ORÇAMENTO)
        query_docs = "SELECT tipo_doc, COUNT(*) as total, SUM(COALESCE(valor_total, 0)) as valor_total FROM documentos_gerados"
        params_docs = []

        if data_inicio and data_fim:
            query_docs += " WHERE date(data_geracao) BETWEEN ? AND ?"
            params_docs = [data_inicio, data_fim]

        query_docs += " GROUP BY tipo_doc"

        docs = conn.execute(query_docs, params_docs).fetchall()
        documentos = {}
        valores_por_tipo = {}
        for doc in docs:
            documentos[doc['tipo_doc']] = doc['total']
            valores_por_tipo[doc['tipo_doc']] = doc['valor_total'] or 0

        # Valor total dos recibos
        valor_recibos = valores_por_tipo.get('RECIBO', 0)

        # Total de laudos
        total_laudos = documentos.get('LAUDO', 0)

        # Total de orçamentos
        total_orcamentos = documentos.get('ORCAMENTO', 0) + documentos.get('ORÇAMENTO', 0)

        # Total de recibos
        total_recibos = documentos.get('RECIBO', 0)

        # Faturamento total = soma dos valores dos recibos (documentos com valor)
        faturamento_total = valor_recibos

        # Documentos por mês (últimos 12 meses) - Laudos, Orçamentos e Recibos
        query_mensal = """
        SELECT
            strftime('%Y-%m', data_geracao) as mes,
            tipo_doc,
            COUNT(*) as total_docs,
            SUM(COALESCE(valor_total, 0)) as total_valor
        FROM documentos_gerados
        WHERE tipo_doc IN ('LAUDO', 'RECIBO', 'ORCAMENTO', 'ORÇAMENTO')
        GROUP BY mes, tipo_doc
        ORDER BY mes DESC
        LIMIT 36
        """

        docs_mensal = conn.execute(query_mensal).fetchall()

        # Agrupar por mês
        mensal_agrupado = {}
        for row in docs_mensal:
            mes = row['mes']
            if mes not in mensal_agrupado:
                mensal_agrupado[mes] = {'mes': mes, 'laudos': 0, 'recibos': 0, 'orcamentos': 0, 'valor_recibos': 0}
            tipo = row['tipo_doc']
            if tipo == 'LAUDO':
                mensal_agrupado[mes]['laudos'] = row['total_docs']
            elif tipo == 'RECIBO':
                mensal_agrupado[mes]['recibos'] = row['total_docs']
                mensal_agrupado[mes]['valor_recibos'] = row['total_valor']
            elif tipo in ('ORCAMENTO', 'ORÇAMENTO'):
                mensal_agrupado[mes]['orcamentos'] += row['total_docs']

        # Converter para lista ordenada
        faturamento_mensal = sorted(mensal_agrupado.values(), key=lambda x: x['mes'], reverse=True)[:12]

        # Top clientes por quantidade de documentos
        query_top_clientes = """
        SELECT
            cliente_nome,
            COUNT(*) as total_docs,
            SUM(COALESCE(valor_total, 0)) as total_valor
        FROM documentos_gerados
        WHERE cliente_nome IS NOT NULL AND cliente_nome != ''
        GROUP BY cliente_nome
        ORDER BY total_docs DESC, total_valor DESC
        LIMIT 10
        """
        top_clientes = conn.execute(query_top_clientes).fetchall()

        return jsonify({
            "documentos": documentos,
            "laudos": {
                "total": total_laudos
            },
            "recibos": {
                "total": total_recibos,
                "valor_total": valor_recibos
            },
            "orcamentos": {
                "total": total_orcamentos
            },
            "faturamento_total": faturamento_total,
            "faturamento_mensal": faturamento_mensal,
            "top_clientes": [dict(row) for row in top_clientes]
        })

    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

@app.route('/api/upload-pdf', methods=['POST'])
def upload_pdf():
    """Upload e processamento de múltiplos PDFs (salva em documentos_gerados)"""
    try:
        if 'arquivos' not in request.files:
            return jsonify({"error": "Nenhum arquivo enviado"}), 400

        arquivos = request.files.getlist('arquivos')
        if len(arquivos) == 0:
            return jsonify({"error": "Lista vazia"}), 400

        resultados = []
        conn = get_db()

        for arquivo in arquivos:
            if not arquivo.filename.lower().endswith('.pdf'):
                resultados.append({
                    "arquivo": arquivo.filename,
                    "status": "erro",
                    "mensagem": "Não é PDF"
                })
                continue

            try:
                # Salvar arquivo
                filename = secure_filename(arquivo.filename)
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                nome_arquivo = f"{timestamp}_{filename}"
                caminho_arquivo = UPLOAD_DIR / nome_arquivo
                arquivo.save(str(caminho_arquivo))

                # Extrair dados
                dados = extrair_dados_pdf(str(caminho_arquivo))

                # Determinar tipo de documento
                tipo_doc = dados.get('tipo_documento', 'UPLOAD')

                # Salvar em documentos_gerados
                cursor = conn.execute('''
                    INSERT INTO documentos_gerados
                    (tipo_doc, nome_arquivo, caminho_arquivo, valor_total, cliente_nome, data_geracao)
                    VALUES (?, ?, ?, ?, ?, ?)
                ''', (
                    tipo_doc,
                    nome_arquivo,
                    str(caminho_arquivo),
                    dados.get('valor_total', 0),
                    dados.get('cliente_nome'),
                    datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                ))
                arquivo_id = cursor.lastrowid
                conn.commit()

                # Sugerir e aplicar tags automaticamente
                conteudo_texto = dados.get('texto_completo', '')
                tags_sugeridas = sugerir_tags_automaticas(arquivo_id, 'documento', nome_arquivo, conteudo_texto)

                # Aplicar tags sugeridas automaticamente
                for tag_sugerida in tags_sugeridas:
                    adicionar_tag_documento(
                        arquivo_id,
                        'documento',
                        tag_sugerida['tag_id'],
                        tag_sugerida['confianca'],
                        manual=False
                    )

                resultados.append({
                    "arquivo": filename,
                    "status": "sucesso",
                    "dados": dados,
                    "tags_sugeridas": tags_sugeridas,
                    "arquivo_id": arquivo_id
                })

            except Exception as e:
                resultados.append({
                    "arquivo": arquivo.filename,
                    "status": "erro",
                    "mensagem": str(e)
                })

        sucessos = len([r for r in resultados if r['status'] == 'sucesso'])
        erros = len([r for r in resultados if r['status'] == 'erro'])

        return jsonify({
            "mensagem": f"{sucessos} sucesso(s), {erros} erro(s)",
            "resultados": resultados,
            "total": len(arquivos),
            "sucessos": sucessos,
            "erros": erros
        })

    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500


@app.route('/api/pdfs-processados', methods=['GET'])
def listar_pdfs_processados():
    """Lista todos os documentos PDF processados"""
    try:
        conn = get_db()
        docs = conn.execute('''
            SELECT * FROM documentos_gerados
            WHERE nome_arquivo LIKE '%.pdf'
            ORDER BY data_geracao DESC
        ''').fetchall()
        return jsonify([dict(row) for row in docs])
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/api/ia-consulta', methods=['POST'])
def ia_consulta():
    """Endpoint para consultas à IA"""
    try:
        dados = request.json
        pergunta = dados.get('pergunta', '').strip()
        tipo_analise = dados.get('tipo_analise', 'geral')

        if not pergunta or len(pergunta) < 3:
            return jsonify({"erro": "Pergunta muito curta"}), 400

        if len(pergunta) > 1000:
            return jsonify({"erro": "Pergunta muito longa"}), 400

        resultado = consultar_ia_analise(pergunta, tipo_analise)
        return jsonify(resultado)

    except Exception as e:
        traceback.print_exc()
        return jsonify({
            "sucesso": False,
            "erro": str(e),
            "resposta": "Erro ao processar"
        }), 500


@app.route('/api/ia-insights', methods=['GET'])
def ia_insights():
    """Insights automáticos"""
    try:
        resultado = gerar_insights_automaticos()
        return jsonify(resultado)
    except Exception as e:
        return jsonify({"erro": str(e)}), 500


@app.route('/api/ia-status', methods=['GET'])
def ia_status():
    """Verifica status da IA"""
    try:
        api_key = os.environ.get("GEMINI_API_KEY")

        if not api_key:
            return jsonify({
                "configurada": False,
                "mensagem": "API Key não configurada. Obtenha em: https://makersuite.google.com/app/apikey"
            })

        if api_key == "SUBSTITUA-PELA-SUA-CHAVE-GEMINI-AQUI":
            return jsonify({
                "configurada": False,
                "mensagem": "API Key precisa ser configurada no arquivo .env"
            })

        return jsonify({
            "configurada": True,
            "mensagem": "IA pronta com Gemini 2.5 Flash",
            "modelo": "gemini-2.5-flash"
        })

    except Exception as e:
        return jsonify({
            "configurada": False,
            "erro": str(e)
        })


@app.route('/api/ia-analise-completa', methods=['GET'])
def ia_analise_completa():
    """Retorna análise completa de faturamento com IA"""
    try:
        resultado = gerar_analise_completa_faturamento()
        return jsonify(resultado)
    except Exception as e:
        traceback.print_exc()
        return jsonify({
            "sucesso": False,
            "erro": str(e),
            "detalhes": traceback.format_exc()
        }), 500


# ==================== GESTÃO DE BANCO DE DADOS ====================

@app.route('/api/database/tables', methods=['GET'])
def listar_tabelas():
    """Lista todas as tabelas do banco de dados"""
    try:
        conn = get_db()
        tabelas = conn.execute("""
            SELECT name FROM sqlite_master
            WHERE type='table' AND name NOT LIKE 'sqlite_%'
            ORDER BY name
        """).fetchall()

        resultado = []
        for tabela in tabelas:
            nome_tabela = tabela['name']
            # Contar registros
            count = conn.execute(f"SELECT COUNT(*) as total FROM {nome_tabela}").fetchone()['total']
            resultado.append({
                'nome': nome_tabela,
                'registros': count
            })

        return jsonify(resultado)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/database/table/<table_name>', methods=['GET'])
def listar_registros_tabela(table_name):
    """Lista todos os registros de uma tabela específica"""
    try:
        # Validar nome da tabela (segurança)
        tabelas_permitidas = ['clientes_web', 'documentos_gerados', 'boletos', 'recibos']
        if table_name not in tabelas_permitidas:
            return jsonify({"error": "Tabela não permitida"}), 403

        conn = get_db()
        limit = request.args.get('limit', 100, type=int)
        offset = request.args.get('offset', 0, type=int)

        registros = conn.execute(f"SELECT * FROM {table_name} LIMIT ? OFFSET ?",
                                (limit, offset)).fetchall()

        # Pegar informações das colunas
        colunas_info = conn.execute(f"PRAGMA table_info({table_name})").fetchall()
        colunas = [col['name'] for col in colunas_info]

        return jsonify({
            'colunas': colunas,
            'registros': [dict(row) for row in registros],
            'total': conn.execute(f"SELECT COUNT(*) as total FROM {table_name}").fetchone()['total']
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/database/table/<table_name>/record/<int:record_id>', methods=['PUT'])
def editar_registro(table_name, record_id):
    """Edita um registro específico"""
    try:
        tabelas_permitidas = ['clientes_web', 'documentos_gerados', 'boletos', 'recibos']
        if table_name not in tabelas_permitidas:
            return jsonify({"error": "Tabela não permitida"}), 403

        dados = request.json
        conn = get_db()

        # Construir UPDATE dinamicamente
        campos = [f"{k} = ?" for k in dados.keys() if k != 'id']
        valores = [v for k, v in dados.items() if k != 'id']
        valores.append(record_id)

        query = f"UPDATE {table_name} SET {', '.join(campos)} WHERE id = ?"
        conn.execute(query, valores)
        conn.commit()

        return jsonify({"message": "Registro atualizado com sucesso"})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/database/table/<table_name>/record/<int:record_id>', methods=['DELETE'])
def deletar_registro(table_name, record_id):
    """Deleta um registro específico"""
    try:
        tabelas_permitidas = ['clientes_web', 'documentos_gerados', 'boletos', 'recibos']
        if table_name not in tabelas_permitidas:
            return jsonify({"error": "Tabela não permitida"}), 403

        conn = get_db()
        conn.execute(f"DELETE FROM {table_name} WHERE id = ?", (record_id,))
        conn.commit()

        return jsonify({"message": "Registro deletado com sucesso"})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/database/table/<table_name>/clear', methods=['POST'])
def limpar_tabela(table_name):
    """Limpa todos os registros de uma tabela (mantém estrutura)"""
    try:
        tabelas_permitidas = ['clientes_web', 'documentos_gerados', 'boletos', 'recibos']
        if table_name not in tabelas_permitidas:
            return jsonify({"error": "Tabela não permitida"}), 403

        conn = get_db()
        conn.execute(f"DELETE FROM {table_name}")
        # Resetar autoincrement
        conn.execute(f"DELETE FROM sqlite_sequence WHERE name = '{table_name}'")
        conn.commit()

        return jsonify({"message": f"Tabela {table_name} limpa com sucesso"})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/database/limpar-laudos', methods=['POST'])
def limpar_laudos():
    """Limpa apenas os laudos da tabela documentos_gerados"""
    try:
        conn = get_db()
        # Contar quantos serão deletados
        count = conn.execute("SELECT COUNT(*) as total FROM documentos_gerados WHERE tipo_doc = 'LAUDO'").fetchone()['total']

        # Deletar apenas laudos
        conn.execute("DELETE FROM documentos_gerados WHERE tipo_doc = 'LAUDO'")
        conn.commit()

        return jsonify({
            "message": f"Laudos limpos com sucesso",
            "deletados": count
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# ==================== CONFIGURAÇÕES DE TEMA ====================

@app.route('/api/config/tema', methods=['GET', 'POST'])
def config_tema():
    """Salva/carrega configurações de tema"""
    config_file = Path('config_tema.json')

    if request.method == 'GET':
        if config_file.exists():
            with open(config_file, 'r', encoding='utf-8') as f:
                return jsonify(json.load(f))
        else:
            return jsonify({
                'modo': 'claro',
                'cor_primaria': '#4f46e5',
                'cor_secundaria': '#10b981',
                'cor_destaque': '#f59e0b'
            })

    else:  # POST
        try:
            dados = request.json
            with open(config_file, 'w', encoding='utf-8') as f:
                json.dump(dados, f, ensure_ascii=False, indent=2)
            return jsonify({"message": "Tema salvo com sucesso"})
        except Exception as e:
            return jsonify({"error": str(e)}), 500

# ==================== TREINAMENTO DO SISTEMA ====================

@app.route('/api/treinar-sistema', methods=['POST'])
def treinar_sistema():
    """Treina o sistema de aprendizagem com os arquivos da pasta training_samples"""
    try:
        print("\n" + "="*60)
        print("INICIANDO TREINAMENTO DO SISTEMA")
        print("="*60)

        # Verificar se tabelas existem
        layout_learner.criar_tabelas()
        print("[OK] Tabelas do banco verificadas/criadas")

        # Treinar com Documentos
        print("\n[1/2] Treinando com Documentos...")
        print(f"    Diretório: {layout_learner.nfse_dir}")
        resultado_docs = layout_learner.treinar_com_diretorio(layout_learner.nfse_dir)
        print(f"    Resultado Docs: {resultado_docs}")

        # Treinar com Boletos
        print("\n[2/2] Treinando com Boletos...")
        print(f"    Diretório: {layout_learner.boleto_dir}")
        resultado_boletos = layout_learner.treinar_com_diretorio(layout_learner.boleto_dir)
        print(f"    Resultado Boletos: {resultado_boletos}")

        # Recarregar layouts
        layout_learner.carregar_layouts()
        print(f"[OK] Layouts recarregados: {len(layout_learner.layouts_conhecidos)} layouts em cache")

        print("\n" + "="*60)
        print("TREINAMENTO CONCLUÍDO!")
        print("="*60 + "\n")

        return jsonify({
            "success": True,
            "message": "Sistema treinado com sucesso!",
            "documentos": {
                "total_arquivos": resultado_docs.get('total_arquivos', 0),
                "processados": resultado_docs.get('processados', 0),
                "erros": resultado_docs.get('erros', 0),
                "layouts_novos": resultado_docs.get('layouts_novos', 0),
                "layouts_existentes": resultado_docs.get('layouts_existentes', 0)
            },
            "boletos": {
                "total_arquivos": resultado_boletos.get('total_arquivos', 0),
                "processados": resultado_boletos.get('processados', 0),
                "erros": resultado_boletos.get('erros', 0),
                "layouts_novos": resultado_boletos.get('layouts_novos', 0),
                "layouts_existentes": resultado_boletos.get('layouts_existentes', 0)
            }
        })

    except Exception as e:
        traceback.print_exc()
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500

@app.route('/api/layouts-conhecidos', methods=['GET'])
def listar_layouts():
    """Lista todos os layouts aprendidos com detalhes"""
    try:
        conn = get_db()

        # Buscar layouts
        layouts = conn.execute("""
            SELECT id, tipo_doc, nome_layout, emissor, palavras_chave,
                   regex_patterns, posicoes_campos, confianca, num_amostras,
                   data_criacao, ultima_atualizacao
            FROM document_layouts
            ORDER BY tipo_doc, data_criacao DESC
        """).fetchall()

        # Buscar amostras de cada layout
        resultado = []
        for layout in layouts:
            layout_dict = dict(layout)

            # Buscar amostras deste layout
            amostras = conn.execute("""
                SELECT arquivo_nome, sucesso, data_adicao
                FROM training_samples
                WHERE layout_id = ?
                ORDER BY data_adicao DESC
            """, (layout['id'],)).fetchall()

            layout_dict['amostras'] = [dict(a) for a in amostras]

            # Parse JSON fields
            if layout_dict['palavras_chave']:
                try:
                    layout_dict['palavras_chave'] = json.loads(layout_dict['palavras_chave'])
                except (json.JSONDecodeError, TypeError):
                    pass  # Manter como string se não for JSON válido
            if layout_dict['regex_patterns']:
                try:
                    layout_dict['regex_patterns'] = json.loads(layout_dict['regex_patterns'])
                except (json.JSONDecodeError, TypeError):
                    pass  # Manter como string se não for JSON válido
            if layout_dict['posicoes_campos']:
                try:
                    layout_dict['posicoes_campos'] = json.loads(layout_dict['posicoes_campos'])
                except (json.JSONDecodeError, TypeError):
                    pass  # Manter como string se não for JSON válido

            resultado.append(layout_dict)

        return jsonify({
            "success": True,
            "total_layouts": len(resultado),
            "layouts": resultado
        })
    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

# ==================== ANÁLISE GEMINI AUTOMÁTICA ====================

def extrair_dados_com_regex(texto):
    """Extrai dados usando regex patterns pré-compilados (otimizado)"""
    dados = {
        'data_emissao': None,
        'data_vencimento': None,
        'valor_total': None,
        'numero_documento': None,
        'cnpj': None,
        'nome_cliente': None,
        'tributos': {}
    }

    # Data de emissão - usar patterns pré-compilados
    for pattern in COMPILED_PATTERNS['data_emissao']:
        match = pattern.search(texto)
        if match:
            dados['data_emissao'] = match.group(1).replace('-', '/')
            break

    # Data de vencimento - usar patterns pré-compilados
    for pattern in COMPILED_PATTERNS['vencimento']:
        match = pattern.search(texto)
        if match:
            dados['data_vencimento'] = match.group(1).replace('-', '/')
            break

    # Valor total - usar patterns pré-compilados
    for pattern in COMPILED_PATTERNS['valor']:
        match = pattern.search(texto)
        if match:
            valor_str = match.group(1).replace('.', '').replace(',', '.')
            try:
                dados['valor_total'] = float(valor_str)
                break
            except (ValueError, AttributeError):
                pass  # Valor não convertível

    # Número do documento - usar patterns pré-compilados
    for pattern in COMPILED_PATTERNS['numero']:
        match = pattern.search(texto)
        if match:
            dados['numero_documento'] = match.group(1)
            break

    # CNPJ - usar patterns pré-compilados
    for pattern in COMPILED_PATTERNS['cnpj']:
        match = pattern.search(texto)
        if match:
            dados['cnpj'] = match.group(1)
            break

    # Nome do Cliente / Tomador - usar patterns pré-compilados
    for pattern in COMPILED_PATTERNS['nome']:
        match = pattern.search(texto)
        if match:
            nome = match.group(1).strip()
            if len(nome) > 5:  # Pelo menos 5 caracteres
                dados['nome_cliente'] = nome
                break

    # Endereço - usar patterns pré-compilados
    for pattern in COMPILED_PATTERNS['endereco']:
        match = pattern.search(texto)
        if match:
            endereco = match.group(1).strip()
            # Limpar possíveis caracteres extras
            endereco = re.sub(r'\s*-\s*$', '', endereco)  # Remove traço no final
            if len(endereco) > 10:  # Endereço razoável
                dados['endereco'] = endereco
                break

    # Bairro - usar patterns pré-compilados
    for pattern in COMPILED_PATTERNS['bairro']:
        match = pattern.search(texto)
        if match:
            bairro = match.group(1).strip()
            if len(bairro) > 2:
                dados['bairro'] = bairro
                break

    # CEP - usar patterns pré-compilados
    for pattern in COMPILED_PATTERNS['cep']:
        match = pattern.search(texto)
        if match:
            cep = match.group(1)
            # Formatar CEP se necessário
            if '-' not in cep and len(cep) == 8:
                cep = f"{cep[:5]}-{cep[5:]}"
            dados['cep'] = cep
            break

    # Tributos - usar patterns pré-compilados
    for tributo, patterns in COMPILED_PATTERNS['tributos'].items():
        for pattern in patterns:
            match = pattern.search(texto)
            if match:
                valor_str = match.group(1).replace('.', '').replace(',', '.')
                try:
                    dados['tributos'][tributo] = float(valor_str)
                    break
                except (ValueError, AttributeError):
                    pass  # Valor não convertível

    return dados

def analisar_arquivo_com_gemini(caminho_arquivo, nome_arquivo):
    """Analisa arquivo com múltiplas estratégias (REGEX, OCR, IA)"""
    try:
        conteudo = ""
        dados_extraidos = {}

        # ESTRATÉGIA 1: Extrair texto do PDF com pdfplumber
        if caminho_arquivo.endswith('.pdf'):
            try:
                import pdfplumber
                with pdfplumber.open(caminho_arquivo) as pdf:
                    for page in pdf.pages[:10]:  # Aumentar para 10 páginas
                        texto_pagina = page.extract_text() or ""
                        conteudo += texto_pagina + "\n"
            except Exception as e:
                print(f"Erro ao extrair PDF: {e}")
                return None
        else:
            try:
                with open(caminho_arquivo, 'r', encoding='utf-8') as f:
                    conteudo = f.read()[:20000]
            except (FileNotFoundError, PermissionError, UnicodeDecodeError) as e:
                print(f"[AVISO] Erro ao ler arquivo {caminho_arquivo}: {e}")
                return None

        if not conteudo or len(conteudo) < 50:
            return None

        # ESTRATÉGIA 3: REGEX para extrair dados estruturados (MAIS CONFIÁVEL)
        dados_regex = extrair_dados_com_regex(conteudo)

        # Se regex encontrou dados suficientes, usar eles
        if dados_regex.get('data_emissao') and dados_regex.get('valor_total'):
            print(f"OK Dados extraídos via REGEX: {dados_regex}")

            # Tentar identificar tipo de documento (MELHORADO)
            tipo_doc = 'outro'
            conteudo_lower = conteudo.lower()

            # Detectar LAUDO
            palavras_laudo = ['laudo', 'técnico', 'vistoria', 'inspeção', 'diagnóstico']
            matches_laudo = sum(1 for palavra in palavras_laudo if palavra in conteudo_lower)

            # Detectar RECIBO
            palavras_recibo = ['recibo', 'recebi de', 'valor recebido']
            matches_recibo = sum(1 for palavra in palavras_recibo if palavra in conteudo_lower)

            # Detectar ORÇAMENTO
            palavras_orcamento = ['orçamento', 'orcamento', 'proposta', 'cotação']
            matches_orcamento = sum(1 for palavra in palavras_orcamento if palavra in conteudo_lower)

            # Determinar tipo baseado em pontuação
            if matches_laudo >= 1:
                tipo_doc = 'laudo'
                print(f"[DETECÇÃO] OK Identificado como LAUDO ({matches_laudo} palavras-chave)")
            elif matches_recibo >= 1:
                tipo_doc = 'recibo'
                print(f"[DETECÇÃO] OK Identificado como RECIBO ({matches_recibo} palavras-chave)")
            elif matches_orcamento >= 1:
                tipo_doc = 'orcamento'
                print(f"[DETECÇÃO] OK Identificado como ORÇAMENTO ({matches_orcamento} palavras-chave)")
            else:
                print(f"[DETECÇÃO] AVISO Não identificado - marcado como 'outro'")

            return {
                'tipo_documento': tipo_doc,
                'dados_cliente': {
                    'nome': dados_regex.get('nome_cliente', ''),
                    'cnpj': dados_regex.get('cnpj', ''),
                    'endereco': dados_regex.get('endereco', ''),
                    'telefone': ''
                },
                'valores': {
                    'total': dados_regex.get('valor_total'),
                    'data_emissao': dados_regex.get('data_emissao'),
                    'data_vencimento': dados_regex.get('data_vencimento'),
                    'numero_documento': dados_regex.get('numero_documento', ''),
                    'tributos': dados_regex.get('tributos', {})
                },
                'informacoes_chave': ['Extraído via REGEX'],
                'resumo': f'{tipo_doc} - {dados_regex.get("data_emissao", "")}'
            }

        # ESTRATÉGIA 4: GEMINI como fallback (apenas se regex falhou)
        if not os.environ.get("GEMINI_API_KEY"):
            return None

        import google.generativeai as genai
        genai.configure(api_key=os.environ.get("GEMINI_API_KEY"))
        model = genai.GenerativeModel('gemini-2.5-flash')

        # Prompt para extração estruturada
        prompt = f"""
Analise o documento a seguir e extraia as informações mais importantes em formato JSON.

DOCUMENTO: {nome_arquivo}

CONTEÚDO:
{conteudo[:3000]}

Retorne APENAS um objeto JSON válido com esta estrutura:
{{
  "tipo_documento": "boleto|recibo|orcamento|laudo|contrato|outro",
  "dados_cliente": {{
    "nome": "nome ou razão social encontrada",
    "cnpj": "CNPJ se encontrado",
    "endereco": "endereço se encontrado",
    "telefone": "telefone se encontrado"
  }},
  "valores": {{
    "total": valor_numerico_ou_null,
    "data_emissao": "data de emissão no formato DD/MM/YYYY ou null",
    "data_vencimento": "data de vencimento no formato DD/MM/YYYY ou null (se for boleto)",
    "numero_documento": "número do documento/nota/boleto se houver",
    "tributos": {{
      "iss": valor_ou_null,
      "pis": valor_ou_null,
      "cofins": valor_ou_null,
      "inss": valor_ou_null,
      "ir": valor_ou_null
    }}
  }},
  "informacoes_chave": ["lista", "de", "informações", "importantes"],
  "resumo": "resumo breve do documento em 1-2 frases"
}}

IMPORTANTE:
- Extraia SEMPRE a data que está no documento, não use a data de hoje
- Se for boleto, procure por "Data de Vencimento" e "Data de Emissão"
- Se for recibo/orçamento/laudo, procure pela data do documento
- Formato de data SEMPRE DD/MM/YYYY

Retorne APENAS o JSON, sem explicações adicionais.
"""

        response = model.generate_content(prompt)
        resultado_texto = response.text.strip()

        # Limpar markdown se houver
        if resultado_texto.startswith('```'):
            resultado_texto = resultado_texto.split('```')[1]
            if resultado_texto.startswith('json'):
                resultado_texto = resultado_texto[4:]

        # Parse JSON
        dados_extraidos = json.loads(resultado_texto)
        return dados_extraidos

    except Exception as e:
        print(f"Erro na análise Gemini: {e}")
        return None

def consultar_cnpj_receita(cnpj):
    """Consulta dados do CNPJ na Receita Federal (API ReceitaWS)"""
    try:
        import requests

        # Limpar CNPJ (apenas números)
        cnpj_limpo = ''.join(filter(str.isdigit, cnpj))

        if len(cnpj_limpo) != 14:
            print(f"[CNPJ] CNPJ inválido: {cnpj}")
            return None

        print(f"[CNPJ] Consultando CNPJ na Receita Federal: {cnpj_limpo}")

        # API gratuita da ReceitaWS
        url = f"https://www.receitaws.com.br/v1/cnpj/{cnpj_limpo}"

        response = requests.get(url, timeout=10)

        if response.status_code == 200:
            dados = response.json()

            if dados.get('status') == 'ERROR':
                print(f"[CNPJ] Erro na consulta: {dados.get('message')}")
                return None

            # Montar endereço completo
            print(f"[CNPJ] Montando endereço a partir dos dados da API...")
            print(f"[CNPJ]   Logradouro: {dados.get('logradouro', 'N/A')}")
            print(f"[CNPJ]   Número: {dados.get('numero', 'N/A')}")
            print(f"[CNPJ]   Complemento: {dados.get('complemento', 'N/A')}")
            print(f"[CNPJ]   Bairro: {dados.get('bairro', 'N/A')}")
            print(f"[CNPJ]   Município: {dados.get('municipio', 'N/A')}")
            print(f"[CNPJ]   UF: {dados.get('uf', 'N/A')}")
            print(f"[CNPJ]   CEP: {dados.get('cep', 'N/A')}")

            endereco_completo = ""
            if dados.get('logradouro'):
                endereco_completo = dados['logradouro']
                if dados.get('numero'):
                    endereco_completo += f", {dados['numero']}"
                if dados.get('complemento') and dados['complemento'].strip():
                    endereco_completo += f", {dados['complemento']}"
                if dados.get('bairro'):
                    endereco_completo += f", {dados['bairro']}"
                if dados.get('municipio'):
                    endereco_completo += f", {dados['municipio']}"
                if dados.get('uf'):
                    endereco_completo += f", {dados['uf']}"
                if dados.get('cep'):
                    endereco_completo += f", CEP: {dados['cep']}"

            print(f"[CNPJ] Endereço montado: {endereco_completo}")

            resultado = {
                'nome': dados.get('nome', ''),
                'razao_social': dados.get('nome', ''),
                'nome_fantasia': dados.get('fantasia', ''),
                'cnpj': dados.get('cnpj', ''),
                'endereco_completo': endereco_completo,
                'logradouro': dados.get('logradouro', ''),
                'numero': dados.get('numero', ''),
                'complemento': dados.get('complemento', ''),
                'bairro': dados.get('bairro', ''),
                'municipio': dados.get('municipio', ''),
                'uf': dados.get('uf', ''),
                'cep': dados.get('cep', ''),
                'telefone': dados.get('telefone', ''),
                'email': dados.get('email', ''),
                'situacao': dados.get('situacao', ''),
                'data_abertura': dados.get('abertura', ''),
            }

            print(f"[CNPJ] OK Dados encontrados na Receita Federal!")
            print(f"[CNPJ]   Razão Social: {resultado['razao_social']}")
            print(f"[CNPJ]   Nome Fantasia: {resultado['nome_fantasia']}")
            print(f"[CNPJ]   Endereço: {endereco_completo[:100]}")
            print(f"[CNPJ]   Situação: {resultado['situacao']}")

            return resultado
        else:
            print(f"[CNPJ] Erro HTTP {response.status_code}")
            return None

    except Exception as e:
        print(f"[CNPJ] Erro ao consultar CNPJ: {e}")
        return None

def processar_boleto(caminho_arquivo, arquivo_nome, documento_id):
    """
    Processa ESPECIFICAMENTE boletos - Função dedicada e otimizada

    Fluxo:
    1. Extrai texto do PDF (OCR com pdfplumber)
    2. Detecta se é realmente um boleto (palavras-chave)
    3. Extrai dados: valor, vencimento, código de barras, beneficiário
    4. Salva na tabela boletos
    5. Atualiza documentos_gerados

    Returns:
        dict: Dados extraídos ou None se falhar
    """
    try:
        print(f"\n{'='*70}")
        print(f"[BOLETO] Processando boleto: {arquivo_nome}")
        print(f"{'='*70}")

        # ========== ETAPA 1: EXTRAIR TEXTO DO PDF ==========
        print(f"[BOLETO] [1/5] Extraindo texto do PDF...")
        texto_completo = ""

        try:
            import pdfplumber
            with pdfplumber.open(caminho_arquivo) as pdf:
                for pagina in pdf.pages:
                    texto = pagina.extract_text()
                    if texto:
                        texto_completo += texto + "\n"
            print(f"[BOLETO] OK Texto extraído: {len(texto_completo)} caracteres")
        except Exception as e:
            print(f"[BOLETO] ERRO Erro ao extrair texto: {e}")
            return None

        if not texto_completo or len(texto_completo) < 50:
            print(f"[BOLETO] ERRO Texto insuficiente para análise")
            return None

        texto_lower = texto_completo.lower()

        # ========== ETAPA 2: VALIDAR SE É BOLETO ==========
        print(f"[BOLETO] [2/5] Validando se é realmente um boleto...")

        palavras_chave_boleto = [
            'boleto', 'vencimento', 'nosso número', 'nosso numero',
            'código de barras', 'codigo de barras', 'linha digitável', 'linha digitavel',
            'beneficiário', 'beneficiario', 'pagador', 'sacado', 'cedente',
            'banco do brasil', 'itaú', 'itau', 'bradesco', 'santander',
            'caixa econômica', 'caixa economica', 'sicoob', 'sicredi',
            'valor do documento', 'valor cobrado', 'agência', 'agencia',
            'conta corrente', 'data de vencimento', 'data do vencimento'
        ]

        matches = sum(1 for palavra in palavras_chave_boleto if palavra in texto_lower)

        if matches < 3:  # Precisa de pelo menos 3 palavras-chave
            print(f"[BOLETO] ERRO Não parece ser um boleto (apenas {matches} palavras-chave encontradas)")
            print(f"[BOLETO] ℹ Palavras encontradas: {[p for p in palavras_chave_boleto if p in texto_lower]}")
            return None

        print(f"[BOLETO] OK Confirmado como boleto ({matches} palavras-chave encontradas)")

        # ========== ETAPA 3: EXTRAIR DADOS DO NOME DO ARQUIVO ==========
        print(f"[BOLETO] [3/5] Extraindo dados do nome do arquivo...")

        import re
        dados_boleto = {}

        # Remover timestamp do nome do arquivo se houver (formato: YYYYMMDD_HHMMSS_)
        nome_limpo = re.sub(r'^\d{8}_\d{6}_', '', arquivo_nome)
        # Remover extensão .pdf
        nome_limpo = nome_limpo.replace('.pdf', '').replace('.PDF', '')

        print(f"[BOLETO]   Nome original: {arquivo_nome}")
        print(f"[BOLETO]   Nome limpo: {nome_limpo}")

        # Padrão: NOME - VALOR - DATA
        # Exemplos: "ELASA - 950,00 - 26-09-2025" ou "EMPRESA TESTE - 1500.50 - 15/12/2024"
        padrao_completo = r'^(.+?)\s*[-–]\s*([\d\.,]+)\s*[-–]\s*(\d{2})[-/](\d{2})[-/](\d{4})'
        match = re.search(padrao_completo, nome_limpo, re.IGNORECASE)

        if match:
            # Extrair do nome do arquivo
            dados_boleto['cliente_nome'] = match.group(1).strip().upper()

            # Extrair valor
            valor_str = match.group(2).replace('.', '').replace(',', '.')
            try:
                dados_boleto['valor'] = float(valor_str)
            except (ValueError, AttributeError):
                dados_boleto['valor'] = 0.0

            # Extrair data de vencimento (DD-MM-YYYY ou DD/MM/YYYY)
            dia = match.group(3)
            mes = match.group(4)
            ano = match.group(5)
            dados_boleto['data_vencimento'] = f"{dia}/{mes}/{ano}"
            dados_boleto['data_emissao'] = datetime.now().strftime('%d/%m/%Y')

            print(f"[BOLETO]   OK Cliente (do arquivo): {dados_boleto['cliente_nome']}")
            print(f"[BOLETO]   OK Valor (do arquivo): R$ {dados_boleto['valor']:.2f}")
            print(f"[BOLETO]   OK Vencimento (do arquivo): {dados_boleto['data_vencimento']}")
            print(f"[BOLETO]   OK Emissão: {dados_boleto['data_emissao']} (data atual)")
        else:
            print(f"[BOLETO]   AVISO Padrão não encontrado no nome do arquivo")
            print(f"[BOLETO]   ℹ Esperado: NOME - VALOR - DD-MM-AAAA")

            # Tentar extrair apenas o nome (tudo antes do primeiro -)
            nome_simples = nome_limpo.split('-')[0].strip() if '-' in nome_limpo else nome_limpo
            dados_boleto['cliente_nome'] = nome_simples.upper() if nome_simples else 'CLIENTE NÃO IDENTIFICADO'
            dados_boleto['valor'] = 0.0
            dados_boleto['data_vencimento'] = datetime.now().strftime('%d/%m/%Y')
            dados_boleto['data_emissao'] = datetime.now().strftime('%d/%m/%Y')

            print(f"[BOLETO]   → Usando: Cliente='{dados_boleto['cliente_nome']}', Valor=R$ 0,00")

        # Campos adicionais (padrão)
        dados_boleto['endereco'] = ''
        dados_boleto['numero_documento'] = ''
        dados_boleto['codigo_barras'] = None

        # Gemini desativado - dados vêm do nome do arquivo
        print(f"[BOLETO]   ℹ Usando dados do nome do arquivo (Gemini desativado)")

        # ========== ETAPA 4: SALVAR NA TABELA BOLETOS ==========
        print(f"[BOLETO] [4/5] Salvando na tabela boletos...")

        conn = get_db()
        cursor = conn.execute('''
            INSERT INTO boletos (
                cliente_nome, descricao, valor, data_emissao, data_vencimento,
                numero_documento, codigo_barras, status, arquivo_caminho, data_criacao
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (
            dados_boleto['cliente_nome'],
            f"Boleto de {dados_boleto['cliente_nome']} - Venc: {dados_boleto['data_vencimento']}",
            dados_boleto['valor'],
            dados_boleto['data_emissao'],
            dados_boleto['data_vencimento'],
            dados_boleto['numero_documento'],
            dados_boleto['codigo_barras'],
            'pendente',
            caminho_arquivo,
            datetime.now()
        ))
        boleto_id = cursor.lastrowid
        conn.commit()

        print(f"[BOLETO] OKOK Boleto salvo com ID: {boleto_id}")

        # ========== ETAPA 5: ATUALIZAR DOCUMENTOS_GERADOS ==========
        print(f"[BOLETO] [5/5] Atualizando documentos_gerados...")

        # Converter data de vencimento para datetime
        try:
            partes = dados_boleto['data_vencimento'].split('/')
            data_obj = datetime(int(partes[2]), int(partes[1]), int(partes[0]))
        except (ValueError, IndexError, AttributeError):
            data_obj = datetime.now()

        conn.execute('''
            UPDATE documentos_gerados
            SET tipo_doc = 'BOLETO',
                cliente_nome = ?,
                valor_total = ?,
                data_geracao = ?
            WHERE id = ?
        ''', (
            dados_boleto['cliente_nome'],
            dados_boleto['valor'],
            data_obj,
            documento_id
        ))
        conn.commit()

        print(f"[BOLETO] OK documentos_gerados atualizado")
        print(f"{'='*70}")
        print(f"[BOLETO] OKOKOK PROCESSAMENTO CONCLUÍDO COM SUCESSO!")
        print(f"{'='*70}\n")

        return dados_boleto

    except Exception as e:
        print(f"[BOLETO] ERROERROERRO ERRO NO PROCESSAMENTO: {e}")
        traceback.print_exc()
        return None

def salvar_dados_gemini_no_banco(dados_extraidos, arquivo_id, arquivo_nome, respeitar_limite=False):
    """Salva dados extraídos pelo Gemini no banco de dados"""
    try:
        if not dados_extraidos:
            return

        conn = get_db()
        tipo_doc = dados_extraidos.get('tipo_documento', 'outro')
        valores = dados_extraidos.get('valores', {})

        print(f"\n[SALVAR] Iniciando salvamento no banco...")
        print(f"[SALVAR] Tipo do documento extraído: {tipo_doc}")
        print(f"[SALVAR] Valores extraídos: {valores}")
        print(f"[SALVAR] Tem valor total? {bool(valores.get('total'))}")

        # Se for cliente identificado, tentar cadastrar (sem duplicar)
        # REGRA: Cadastro automático SÓ com CNPJ + Nome + Endereço
        if dados_extraidos.get('dados_cliente'):
            cliente = dados_extraidos['dados_cliente']
            nome = cliente.get('nome', '').strip()
            cnpj = cliente.get('cnpj', '').strip()
            endereco = cliente.get('endereco', '').strip()

            # CONSULTAR CNPJ NA RECEITA FEDERAL PARA COMPLEMENTAR DADOS
            if cnpj:
                # Se deve respeitar limite, adicionar delay
                if respeitar_limite:
                    import time
                    # Verificar se já houve consultas recentes
                    if not hasattr(salvar_dados_gemini_no_banco, '_ultimas_consultas'):
                        salvar_dados_gemini_no_banco._ultimas_consultas = []

                    agora = time.time()
                    # Remover consultas mais antigas que 60 segundos
                    salvar_dados_gemini_no_banco._ultimas_consultas = [
                        t for t in salvar_dados_gemini_no_banco._ultimas_consultas
                        if agora - t < 60
                    ]

                    # Se já tem 3 consultas no último minuto, esperar
                    if len(salvar_dados_gemini_no_banco._ultimas_consultas) >= 3:
                        tempo_espera = 60 - (agora - salvar_dados_gemini_no_banco._ultimas_consultas[0])
                        if tempo_espera > 0:
                            print(f"[CNPJ] Aguardando {tempo_espera:.0f}s para respeitar limite de 3/min...")
                            time.sleep(tempo_espera + 1)  # +1 segundo de margem

                    # Registrar esta consulta
                    salvar_dados_gemini_no_banco._ultimas_consultas.append(time.time())

                # Tentar consultar, mas não falhar se der erro
                try:
                    dados_receita = consultar_cnpj_receita(cnpj)
                except Exception as e:
                    print(f"[CNPJ] ERRO ao consultar Receita: {e}")
                    dados_receita = None
                if dados_receita:
                    # Usar dados da Receita Federal (SEMPRE PRIORITÁRIOS)
                    print(f"[CLIENTE] OK Dados da Receita Federal recebidos!")

                    # SEMPRE usar dados da Receita quando disponíveis (mais confiáveis)
                    nome_receita = dados_receita.get('nome_fantasia') or dados_receita.get('razao_social', '')
                    razao_receita = dados_receita.get('razao_social', '')
                    endereco_receita = dados_receita.get('endereco_completo', '')
                    cnpj_receita = dados_receita.get('cnpj', cnpj)
                    telefone_receita = dados_receita.get('telefone', '')
                    email_receita = dados_receita.get('email', '')

                    print(f"[CLIENTE] Dados extraídos da Receita:")
                    print(f"[CLIENTE]   Nome Fantasia: {nome_receita}")
                    print(f"[CLIENTE]   Razão Social: {razao_receita}")
                    print(f"[CLIENTE]   Endereço: {endereco_receita}")
                    print(f"[CLIENTE]   Telefone: {telefone_receita}")
                    print(f"[CLIENTE]   Email: {email_receita}")

                    # PRIORIZAR dados da Receita, usar extraído do PDF apenas como fallback
                    nome = nome_receita if nome_receita else nome
                    endereco = endereco_receita if endereco_receita else endereco

                    # Atualizar dados do cliente com informações da Receita (prioritárias)
                    cliente.update({
                        'nome': nome,
                        'razao_social': razao_receita if razao_receita else nome,
                        'nome_fantasia': nome_receita if nome_receita else nome,
                        'cnpj': cnpj_receita,
                        'endereco': endereco,
                        'telefone': telefone_receita if telefone_receita else cliente.get('telefone', ''),
                        'email': email_receita if email_receita else cliente.get('email', ''),
                    })

                    print(f"[CLIENTE] Dados finais após Receita:")
                    print(f"[CLIENTE]   Nome: {nome}")
                    print(f"[CLIENTE]   Endereço: {endereco[:100] if endereco else 'N/A'}")
                else:
                    print(f"[CLIENTE] AVISO Consulta à Receita falhou ou não retornou dados")

            # Validar se tem TODOS os dados obrigatórios para cadastro automático
            tem_dados_completos = bool(nome and cnpj and endereco)

            if not tem_dados_completos:
                print(f"[CLIENTE] Dados incompletos para cadastro automático")
                print(f"[CLIENTE]   Nome: {'OK' if nome else 'ERRO'} {nome[:50] if nome else 'N/A'}")
                print(f"[CLIENTE]   CNPJ: {'OK' if cnpj else 'ERRO'} {cnpj if cnpj else 'N/A'}")
                print(f"[CLIENTE]   Endereço: {'OK' if endereco else 'ERRO'} {endereco[:50] if endereco else 'N/A'}")
                print(f"[CLIENTE] AVISO Cliente NÃO será cadastrado automaticamente (faltam dados)")
            else:
                # Verificar se cliente já existe (por CNPJ primeiro, depois por nome)
                existe = None

                if cnpj:
                    # Remover pontuação do CNPJ para comparação
                    cnpj_limpo = ''.join(filter(str.isdigit, cnpj))
                    if cnpj_limpo:
                        existe = conn.execute(
                            "SELECT id, nome_fantasia FROM clientes_web WHERE REPLACE(REPLACE(REPLACE(cnpj, '.', ''), '/', ''), '-', '') = ?",
                            (cnpj_limpo,)
                        ).fetchone()

                # Se não encontrou por CNPJ, tentar por nome exato
                if not existe and nome:
                    existe = conn.execute(
                        "SELECT id, nome_fantasia FROM clientes_web WHERE nome_fantasia = ? OR razao_social = ?",
                        (nome, nome)
                    ).fetchone()

                if existe:
                    print(f"[CLIENTE] Cliente já existe: {existe[1]} (ID: {existe[0]})")
                else:
                    # Cadastrar novo cliente com todos os dados obrigatórios
                    print(f"[CLIENTE] OK Dados completos encontrados!")
                    print(f"[CLIENTE]   Nome: {nome}")
                    print(f"[CLIENTE]   CNPJ: {cnpj}")
                    print(f"[CLIENTE]   Endereço: {endereco[:100]}")

                    # Valores que serão inseridos no banco
                    razao_social_final = cliente.get('razao_social', nome)
                    telefone_final = cliente.get('telefone', '')
                    email_final = cliente.get('email', '')

                    print(f"[CLIENTE] >>> VALORES SENDO INSERIDOS NO BANCO:")
                    print(f"[CLIENTE] >>> nome_fantasia: {nome}")
                    print(f"[CLIENTE] >>> razao_social: {razao_social_final}")
                    print(f"[CLIENTE] >>> cnpj: {cnpj}")
                    print(f"[CLIENTE] >>> endereco_completo: {endereco}")
                    print(f"[CLIENTE] >>> telefone: {telefone_final}")
                    print(f"[CLIENTE] >>> email: {email_final}")

                    conn.execute('''INSERT INTO clientes_web
                        (nome_fantasia, razao_social, cnpj, endereco_completo, telefone, email, data_cadastro)
                        VALUES (?, ?, ?, ?, ?, ?, ?)''',
                        (nome,
                         razao_social_final,
                         cnpj,
                         endereco,
                         telefone_final,
                         email_final,
                         datetime.now()))
                    conn.commit()

                    # Verificar o que foi realmente salvo
                    cliente_inserido = conn.execute(
                        "SELECT id, nome_fantasia, endereco_completo FROM clientes_web WHERE cnpj = ? ORDER BY id DESC LIMIT 1",
                        (cnpj,)
                    ).fetchone()

                    if cliente_inserido:
                        print(f"[CLIENTE] OKOKOK Cliente cadastrado com sucesso! ID: {cliente_inserido[0]}")
                        print(f"[CLIENTE] >>> VERIFICAÇÃO NO BANCO:")
                        print(f"[CLIENTE] >>> Nome salvo: {cliente_inserido[1]}")
                        print(f"[CLIENTE] >>> Endereço salvo: {cliente_inserido[2]}")
                    else:
                        print(f"[CLIENTE] AVISO️ Erro ao verificar cliente inserido!")

        # Log do tipo detectado
        print(f"[SALVAR] Tipo detectado: {tipo_doc}, tem_total: {bool(valores.get('total'))}")

        # Atualizar documento_gerados com a data correta do documento
        data_documento = valores.get('data_emissao') or datetime.now().strftime('%d/%m/%Y')
        try:
            # Converter DD/MM/YYYY para objeto datetime
            if '/' in data_documento:
                partes = data_documento.split('/')
                data_obj = datetime(int(partes[2]), int(partes[1]), int(partes[0]))
            else:
                data_obj = datetime.now()
        except (ValueError, IndexError, AttributeError):
            data_obj = datetime.now()

        # Preparar nome do cliente
        cliente_nome_para_doc = dados_extraidos.get('dados_cliente', {}).get('nome', '')

        # Mapear tipo_doc para formato documentos_gerados
        tipo_doc_map = {
            'recibo': 'RECIBO',
            'laudo': 'LAUDO',
            'orcamento': 'ORÇAMENTO',
            'outro': 'UPLOAD'
        }
        tipo_doc_final = tipo_doc_map.get(tipo_doc, tipo_doc.upper())

        conn.execute('''UPDATE documentos_gerados
                       SET data_geracao = ?, valor_total = ?, tipo_doc = ?, cliente_nome = ?
                       WHERE id = ?''',
                    (data_obj, valores.get('total', 0.0), tipo_doc_final, cliente_nome_para_doc, arquivo_id))
        conn.commit()
        print(f"[SALVAR] OK Registro atualizado em documentos_gerados: tipo='{tipo_doc_final}', cliente='{cliente_nome_para_doc}'")

        # Salvar resumo e informações em tabela de análises
        conn.execute('''CREATE TABLE IF NOT EXISTS analises_gemini (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            arquivo_id INTEGER,
            arquivo_nome TEXT,
            tipo_documento TEXT,
            dados_extraidos TEXT,
            resumo TEXT,
            data_analise TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )''')

        conn.execute('''INSERT INTO analises_gemini
            (arquivo_id, arquivo_nome, tipo_documento, dados_extraidos, resumo)
            VALUES (?, ?, ?, ?, ?)''',
            (arquivo_id, arquivo_nome,
             tipo_doc,
             json.dumps(dados_extraidos, ensure_ascii=False),
             dados_extraidos.get('resumo', '')))
        conn.commit()

    except Exception as e:
        print(f"Erro ao salvar dados Gemini: {e}")
        traceback.print_exc()

with app.app_context():
    criar_tabelas()
    # Registrar tabelas legadas no ORM (transição para SQLAlchemy)
    try:
        registrar_tabelas_legadas()
        print("[OK] Serviços ORM para tabelas legadas inicializados")
    except Exception as e:
        print(f"[AVISO] Erro ao registrar tabelas legadas: {e}")
    # Criar tabelas do sistema de aprendizagem
    layout_learner.criar_tabelas()
    layout_learner.carregar_layouts()
    # Carregar CNAEs ao iniciar
    carregar_cnaes_permitidos()
    # Limpar arquivos temporários
    limpar_arquivos_temporarios()

# ==================== ENDPOINTS DE CONFIGURAÇÃO DE DIRETÓRIOS ====================

@app.route('/api/selecionar-diretorio', methods=['POST'])
def selecionar_diretorio():
    """Abre diálogo para selecionar diretório"""
    try:
        import tkinter as tk
        from tkinter import filedialog

        # Criar janela oculta
        root = tk.Tk()
        root.withdraw()
        root.attributes('-topmost', True)

        # Abrir diálogo de seleção de pasta
        caminho = filedialog.askdirectory(
            title="Selecione o diretório",
            mustexist=True
        )

        root.destroy()

        if caminho:
            return jsonify({'success': True, 'caminho': caminho})
        else:
            return jsonify({'success': False, 'message': 'Nenhum diretório selecionado'})

    except Exception as e:
        print(f"Erro ao selecionar diretório: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/salvar-diretorios', methods=['POST'])
def salvar_diretorios():
    """Salva configuração de diretórios duplos"""
    try:
        data = request.json
        dir_principal = data.get('principal')
        dir_backup = data.get('backup')
        criar_pastas = data.get('criar_pastas', True)  # Por padrão, cria pastas

        if not dir_principal or not dir_backup:
            return jsonify({'success': False, 'message': 'Ambos os diretórios são obrigatórios'})

        # Verificar se os diretórios existem
        if not Path(dir_principal).exists():
            return jsonify({'success': False, 'message': 'Diretório principal não existe'})

        if not Path(dir_backup).exists():
            return jsonify({'success': False, 'message': 'Diretório de backup não existe'})

        estrutura_criada = []

        # Criar estrutura de pastas APENAS se a opção estiver marcada
        if criar_pastas:
            pastas = ['Laudos', 'Recibos', 'Orcamentos', 'Boletos']

            for pasta in pastas:
                # Criar no diretório principal
                caminho_principal = Path(dir_principal) / pasta
                caminho_principal.mkdir(parents=True, exist_ok=True)

                # Criar no diretório de backup
                caminho_backup = Path(dir_backup) / pasta
                caminho_backup.mkdir(parents=True, exist_ok=True)

                estrutura_criada.append(pasta)

        # Salvar configuração
        config = {
            'principal': str(dir_principal),
            'backup': str(dir_backup),
            'ativo': True,
            'criar_pastas_auto': criar_pastas,
            'data_configuracao': datetime.now().isoformat()
        }

        config_file = BASE_DIR / 'config_diretorios_duplos.json'
        with open(config_file, 'w', encoding='utf-8') as f:
            json.dump(config, f, ensure_ascii=False, indent=2)

        response_data = {
            'success': True,
            'message': 'Diretórios configurados com sucesso'
        }

        if estrutura_criada:
            response_data['estrutura_criada'] = estrutura_criada

        return jsonify(response_data)

    except Exception as e:
        print(f"Erro ao salvar diretórios: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/testar-diretorios', methods=['POST'])
def testar_diretorios():
    """Testa se os diretórios estão acessíveis e cria arquivo de teste"""
    try:
        data = request.json
        dir_principal = data.get('principal')
        dir_backup = data.get('backup')

        if not dir_principal or not dir_backup:
            return jsonify({'success': False, 'message': 'Diretórios não configurados'})

        # Testar escrita no diretório principal
        teste_principal = Path(dir_principal) / 'teste_acesso.txt'
        teste_principal.write_text(f'Teste de acesso - {datetime.now()}', encoding='utf-8')

        # Testar escrita no diretório de backup
        teste_backup = Path(dir_backup) / 'teste_acesso.txt'
        teste_backup.write_text(f'Teste de acesso - {datetime.now()}', encoding='utf-8')

        # Remover arquivos de teste
        teste_principal.unlink()
        teste_backup.unlink()

        return jsonify({
            'success': True,
            'message': 'Ambos os diretórios estão acessíveis e funcionando',
            'arquivos_teste_criados': True
        })

    except Exception as e:
        print(f"Erro ao testar diretórios: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/obter-diretorios')
def obter_diretorios():
    """Retorna os diretórios configurados"""
    try:
        config_file = BASE_DIR / 'config_diretorios_duplos.json'

        if config_file.exists():
            with open(config_file, 'r', encoding='utf-8') as f:
                config = json.load(f)

            return jsonify({
                'success': True,
                'diretorios': {
                    'principal': config.get('principal'),
                    'backup': config.get('backup')
                }
            })
        else:
            return jsonify({'success': True, 'diretorios': {}})

    except Exception as e:
        print(f"Erro ao obter diretórios: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

def salvar_upload_simples(conteudo, nome_arquivo, tipo_documento):
    """
    Salva um arquivo de upload APENAS no diretório principal configurado
    (uploads são manuais, então não fazemos backup automático)

    Args:
        conteudo: bytes do arquivo OU caminho do arquivo já salvo
        nome_arquivo: nome do arquivo
        tipo_documento: 'LAUDO', 'RECIBO', 'ORCAMENTO', 'BOLETO'

    Returns:
        str: caminho onde o arquivo foi salvo
    """
    try:
        print(f"[SALVAR] Iniciando salvar_upload_simples()")
        print(f"[SALVAR] - nome_arquivo: {nome_arquivo}")
        print(f"[SALVAR] - tipo_documento: {tipo_documento}")
        print(f"[SALVAR] - tipo de conteudo: {type(conteudo).__name__}")

        # Mapear tipo de documento para nome de pasta
        mapa_pastas = {
            'LAUDO': 'Laudos',
            'RECIBO': 'Recibos',
            'ORCAMENTO': 'Orcamentos',
            'BOLETO': 'Boletos'
        }

        pasta = mapa_pastas.get(tipo_documento, 'Laudos')
        print(f"[SALVAR] - pasta destino: {pasta}")

        config_file = BASE_DIR / 'config_diretorios_duplos.json'
        print(f"[SALVAR] - config_file: {config_file}")
        print(f"[SALVAR] - config existe: {config_file.exists()}")

        # Se não houver configuração, usar diretório padrão
        if not config_file.exists():
            print(f"[SALVAR] Sem config - usando diretório padrão")
            caminho_padrao = OUTPUT_DIR / pasta / nome_arquivo
            caminho_padrao.parent.mkdir(parents=True, exist_ok=True)

            if isinstance(conteudo, bytes):
                caminho_padrao.write_bytes(conteudo)
            elif isinstance(conteudo, str):
                shutil.copy2(conteudo, caminho_padrao)

            print(f"[SALVAR] OK Arquivo salvo (sem config): {caminho_padrao}")
            return str(caminho_padrao)

        # Carregar configuração
        with open(config_file, 'r', encoding='utf-8') as f:
            config = json.load(f)

        print(f"[SALVAR] Config carregada: {config}")
        dir_principal = config.get('principal')
        print(f"[SALVAR] Diretorio principal: {dir_principal}")

        # Salvar APENAS no diretório principal (uploads são manuais)
        if dir_principal:
            caminho_principal = Path(dir_principal) / pasta / nome_arquivo
            print(f"[SALVAR] Caminho completo: {caminho_principal}")
            print(f"[SALVAR] Criando diretorio: {caminho_principal.parent}")
            caminho_principal.parent.mkdir(parents=True, exist_ok=True)

            if isinstance(conteudo, bytes):
                print(f"[SALVAR] Salvando bytes...")
                caminho_principal.write_bytes(conteudo)
            elif isinstance(conteudo, str):
                print(f"[SALVAR] Copiando de: {conteudo}")
                shutil.copy2(conteudo, caminho_principal)

            print(f"[SALVAR] OK Arquivo salvo (config): {caminho_principal}")
            return str(caminho_principal)

        # Fallback para diretório padrão
        print(f"[SALVAR] Sem dir_principal - usando fallback")
        caminho_padrao = OUTPUT_DIR / pasta / nome_arquivo
        caminho_padrao.parent.mkdir(parents=True, exist_ok=True)

        if isinstance(conteudo, bytes):
            caminho_padrao.write_bytes(conteudo)
        elif isinstance(conteudo, str):
                shutil.copy2(conteudo, caminho_padrao)

        print(f"[SALVAR] OK Arquivo salvo (fallback): {caminho_padrao}")
        return str(caminho_padrao)

    except Exception as e:
        print(f"[ERRO] Erro ao salvar upload: {e}")
        traceback.print_exc()

        # Fallback para diretório padrão
        try:
            caminho_padrao = OUTPUT_DIR / pasta / nome_arquivo
            caminho_padrao.parent.mkdir(parents=True, exist_ok=True)

            if isinstance(conteudo, bytes):
                caminho_padrao.write_bytes(conteudo)
            elif isinstance(conteudo, str):
                shutil.copy2(conteudo, caminho_padrao)

            print(f"[ERRO] Arquivo salvo em fallback: {caminho_padrao}")
            return str(caminho_padrao)
        except Exception as e2:
            print(f"[ERRO] Falha crítica no fallback: {e2}")
            traceback.print_exc()
            return None

def salvar_arquivo_duplo(conteudo, nome_arquivo, tipo_documento):
    """
    Salva um arquivo nos diretórios configurados (principal + backup)

    Args:
        conteudo: bytes do arquivo OU caminho do arquivo já salvo
        nome_arquivo: nome do arquivo (ex: laudo_123.pdf)
        tipo_documento: 'LAUDO', 'RECIBO', 'ORCAMENTO', 'BOLETO'

    Returns:
        dict com {'caminho_principal': str, 'caminhos_todos': [str, str, ...]}
    """
    try:
        # Mapear tipo de documento para nome de pasta
        mapa_pastas = {
            'LAUDO': 'Laudos',
            'RECIBO': 'Recibos',
            'ORCAMENTO': 'Orcamentos',
            'BOLETO': 'Boletos'
        }

        pasta = mapa_pastas.get(tipo_documento, 'Laudos')

        config_file = BASE_DIR / 'config_diretorios_duplos.json'

        # Se não houver configuração, salvar apenas no diretório padrão
        if not config_file.exists():
            caminho_padrao = OUTPUT_DIR / pasta / nome_arquivo
            caminho_padrao.parent.mkdir(parents=True, exist_ok=True)

            if isinstance(conteudo, bytes):
                caminho_padrao.write_bytes(conteudo)
            elif isinstance(conteudo, str):
                # Copiar arquivo
                        shutil.copy2(conteudo, caminho_padrao)

            return {
                'caminho_principal': str(caminho_padrao),
                'caminhos_todos': [str(caminho_padrao)]
            }

        # Carregar configuração
        with open(config_file, 'r', encoding='utf-8') as f:
            config = json.load(f)

        dir_principal = config.get('principal')
        dir_backup = config.get('backup')

        caminhos_salvos = []
        caminho_principal_str = None

        # Salvar no diretório principal
        if dir_principal:
            caminho_principal = Path(dir_principal) / pasta / nome_arquivo
            caminho_principal.parent.mkdir(parents=True, exist_ok=True)

            if isinstance(conteudo, bytes):
                caminho_principal.write_bytes(conteudo)
            elif isinstance(conteudo, str):
                shutil.copy2(conteudo, caminho_principal)

            caminho_principal_str = str(caminho_principal)
            caminhos_salvos.append(caminho_principal_str)

        # Salvar no diretório de backup
        if dir_backup:
            caminho_backup = Path(dir_backup) / pasta / nome_arquivo
            caminho_backup.parent.mkdir(parents=True, exist_ok=True)

            if isinstance(conteudo, bytes):
                caminho_backup.write_bytes(conteudo)
            elif isinstance(conteudo, str):
                shutil.copy2(conteudo, caminho_backup)

            caminhos_salvos.append(str(caminho_backup))

        # Se não salvou em nenhum configurado, usar padrão
        if not caminho_principal_str:
            caminho_padrao = OUTPUT_DIR / pasta / nome_arquivo
            caminho_padrao.parent.mkdir(parents=True, exist_ok=True)

            if isinstance(conteudo, bytes):
                caminho_padrao.write_bytes(conteudo)
            elif isinstance(conteudo, str):
                shutil.copy2(conteudo, caminho_padrao)

            caminho_principal_str = str(caminho_padrao)
            caminhos_salvos.append(caminho_principal_str)

        return {
            'caminho_principal': caminho_principal_str,
            'caminhos_todos': caminhos_salvos
        }

    except Exception as e:
        print(f"Erro ao salvar arquivo duplo: {e}")
        traceback.print_exc()

        # Fallback para diretório padrão
        try:
            caminho_padrao = OUTPUT_DIR / pasta / nome_arquivo
            caminho_padrao.parent.mkdir(parents=True, exist_ok=True)

            if isinstance(conteudo, bytes):
                caminho_padrao.write_bytes(conteudo)
            elif isinstance(conteudo, str):
                shutil.copy2(conteudo, caminho_padrao)

            return {
                'caminho_principal': str(caminho_padrao),
                'caminhos_todos': [str(caminho_padrao)]
            }
        except Exception as e2:
            raise Exception(f"Erro crítico ao salvar arquivo: {e2}")

@app.route('/prospeccao')
def prospeccao():
    """Página de prospecção de clientes (estilo Econodata)"""
    return render_template('prospeccao.html')

@app.route('/api/abrir-arquivo', methods=['POST'])
def abrir_arquivo():
    """Abre um arquivo no visualizador padrão do sistema"""
    try:
        import subprocess
        import platform

        data = request.json
        caminho = data.get('caminho')

        if not caminho:
            return jsonify({'success': False, 'message': 'Caminho não fornecido'})

        # Converter para Path
        arquivo = Path(caminho)

        if not arquivo.exists():
            return jsonify({'success': False, 'message': 'Arquivo não encontrado'})

        # Abrir arquivo de acordo com o sistema operacional
        sistema = platform.system()

        if sistema == 'Windows':
            os.startfile(str(arquivo))
        elif sistema == 'Darwin':  # macOS
            subprocess.run(['open', str(arquivo)])
        else:  # Linux
            subprocess.run(['xdg-open', str(arquivo)])

        return jsonify({'success': True, 'message': 'Arquivo aberto'})

    except Exception as e:
        print(f"Erro ao abrir arquivo: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

if __name__ == '__main__':
    print("\n" + "="*70)
    print("  SISTEMA DE GESTAO DE DOCUMENTOS - ULTRA OTIMIZADO")
    print("  ACESSO DIRETO - SEM LOGIN - SEM SENHA")
    print("="*70)
    print("\n  OTIMIZACOES ATIVAS:")
    print("     - SQLite com WAL mode e cache de 64MB")
    print("     - Cache LRU para CNAEs e configuracoes")
    print("     - Regex patterns compilados")
    print("     - Processamento PDF otimizado")
    print("     - Queries SQL indexadas")
    print("     - Limpeza automatica de arquivos temporarios")
    print("     - Gemini 2.5 Flash (IA mais rapida)")
    print("\n  Acesse: http://localhost:5000")
    print("  Pressione CTRL+C para parar")
    print("="*70 + "\n")

    # Configurações otimizadas para produção
    app.run(
        debug=True,
        host='0.0.0.0',
        port=5000,
        threaded=True,  # Habilitar threading
        use_reloader=False  # Desabilitar reloader em produção
    )

# Reload 1767837493.9213495